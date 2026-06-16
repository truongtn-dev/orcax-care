"""Update WDP301 document: text Momo->SePay + replace payment-related diagram images."""
from __future__ import annotations

import re
import shutil
import zipfile
from io import BytesIO
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "WDP301-SE1816-GROUP4_Document.docx"
BACKUP_PATH = ROOT / "docs" / "WDP301-SE1816-GROUP4_Document.backup.docx"
ASSETS_DIR = ROOT / "docs" / "_generated_diagrams"

TEXT_REPLACEMENTS = [
    ("Pay via Momo", "Pay via SePay"),
    ("PayOS (UC-7.1) or Momo (UC-7.2)", "PayOS (UC-7.1) or SePay (UC-7.2)"),
    ("Wallet / PayOS (UC-7.1) / Momo (UC-7.2)", "Wallet / PayOS (UC-7.1) / SePay (UC-7.2)"),
    ("Wallet, PayOS, or Momo", "Wallet, PayOS, or SePay"),
    ("PayOS, Momo, or medical wallet", "PayOS, SePay, or medical wallet"),
    ("PayOS, Momo (refund)", "PayOS, SePay (refund)"),
    ("PayOS or Momo", "PayOS or SePay"),
    ("PayOS/Momo", "PayOS/SePay"),
    ("PayOS / Momo", "PayOS / SePay"),
    ("wallet/PayOS/Momo", "wallet/PayOS/SePay"),
    ("Wallet / PayOS / Momo", "Wallet / PayOS / SePay"),
    ("Payment (PayOS / Momo / Wallet)", "Payment (PayOS / SePay / Wallet)"),
    ("Process PayOS / Momo callback", "Process PayOS / SePay callback"),
    ("Enum(payos, momo, wallet)", "Enum(payos, sepay, wallet)"),
    ("Payment gateways (PayOS/Momo)", "Payment gateways (PayOS/SePay)"),
    ("gateway choice PayOS/Momo", "gateway choice PayOS/SePay"),
    ("Secondary Actors: PayOS, Momo", "Secondary Actors: PayOS, SePay"),
    ("PayOS, Momo", "PayOS, SePay"),
    ("through Momo", "through SePay"),
    ("via Momo", "via SePay"),
    ("Momo gateways", "SePay gateways"),
    ("Momo.", "SePay."),
    ("Momo,", "SePay,"),
    ("Momo ", "SePay "),
    ("Momo", "SePay"),
    ("momo", "sepay"),
    (
        "5. System redirects to the gateway; patient completes the payment.",
        "5. System opens the in-app checkout page with a VietQR code (PayOS or SePay); patient completes payment in their banking app.",
    ),
    (
        "6. Gateway callback verifies the payment; system credits the wallet and records the transaction.",
        "6. PayOS webhook or SePay IPN/polling verifies the payment; system credits the wallet and records the transaction.",
    ),
    (
        "Amount input (min 50,000 VND) and gateway choice PayOS/SePay.",
        "Amount input (min 10,000 VND) and gateway choice PayOS/SePay.",
    ),
    (
        "wallet and payment (PayOS/SePay)",
        "wallet and payment (PayOS/SePay)",
    ),
]

TARGET_FIGURES = {2, 10, 12, 19, 24, 47, 61, 62, 71, 72}


def apply_text_replacements(text: str) -> str:
    out = text
    for old, new in TEXT_REPLACEMENTS:
        out = out.replace(old, new)
    return out


def replace_doc_text(doc: Document) -> int:
    count = 0

    def fix(text: str) -> str:
        nonlocal count
        new = apply_text_replacements(text)
        if new != text:
            count += 1
        return new

    for p in doc.paragraphs:
        if p.text:
            new = fix(p.text)
            if new != p.text:
                p.text = new

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    new = fix(cell.text)
                    if new != cell.text:
                        cell.text = new
    return count


def save_fig(path: Path, fig) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=160, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def draw_box(ax, x, y, w, h, text, fc="#e0f2fe", ec="#0891b2", fontsize=9):
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.08",
        linewidth=1.4,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, wrap=True)


def draw_arrow(ax, x1, y1, x2, y2, text=None):
    arr = FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=12, linewidth=1.2, color="#334155")
    ax.add_patch(arr)
    if text:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.15, text, ha="center", fontsize=8, color="#475569")


def gen_figure_2_usecase(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.text(6, 7.6, "Patient Use Case Diagram (excerpt)", ha="center", fontsize=14, fontweight="bold")

    draw_box(ax, 0.5, 3.2, 1.4, 1.2, "Patient", fc="#cffafe")
    draw_box(ax, 3, 5.8, 3.2, 0.9, "Book Appointment\n(UC-7)", fc="#dbeafe")
    draw_box(ax, 7.2, 6.5, 2.2, 0.7, "Pay via PayOS\n(UC-7.1)", fc="#ecfeff", fontsize=8)
    draw_box(ax, 7.2, 5.5, 2.2, 0.7, "Pay via SePay\n(UC-7.2)", fc="#ecfeff", fontsize=8)
    draw_box(ax, 7.2, 4.5, 2.6, 0.7, "Apply Insurance\nDiscount (UC-7.3)", fc="#ecfeff", fontsize=8)

    draw_box(ax, 3, 2.8, 3.4, 0.9, "View Wallet Balance\n(UC-12)", fc="#dbeafe")
    draw_box(ax, 7.2, 2.6, 2.4, 0.7, "Top up via PayOS", fc="#ecfeff", fontsize=8)
    draw_box(ax, 7.2, 1.7, 2.4, 0.7, "Top up via SePay", fc="#ecfeff", fontsize=8)

    draw_box(ax, 10.2, 6.2, 1.5, 0.8, "PayOS", fc="#fef3c7", ec="#d97706")
    draw_box(ax, 10.2, 4.8, 1.5, 0.8, "SePay", fc="#fef3c7", ec="#d97706")

    draw_arrow(ax, 1.9, 3.8, 3, 6.2)
    draw_arrow(ax, 6.2, 6.2, 7.2, 6.85)
    draw_arrow(ax, 6.2, 6.0, 7.2, 5.85)
    draw_arrow(ax, 6.2, 5.7, 7.2, 4.85)
    draw_arrow(ax, 1.9, 3.5, 3, 3.2)
    draw_arrow(ax, 6.4, 3.2, 7.2, 2.95)
    draw_arrow(ax, 6.4, 3.0, 7.2, 2.05)
    draw_arrow(ax, 9.4, 6.85, 10.2, 6.6, "<<include>>")
    draw_arrow(ax, 9.4, 5.85, 10.2, 5.2, "<<include>>")
    save_fig(path, fig)


def gen_figure_10_erd(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 5))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.text(5.5, 4.6, "ERD excerpt — Wallet & Payment", ha="center", fontsize=14, fontweight="bold")

    draw_box(ax, 0.5, 1.5, 2.8, 2.2, "wallets\n────────────\n_id: ObjectId\nuserId: ObjectId\nbalance: Number", fc="#ecfeff")
    draw_box(ax, 4, 1.2, 3.4, 2.8, "wallet_transactions\n────────────\n_id: ObjectId\nuserId: ObjectId\ntype: Enum(topup,payment,refund)\nprovider: payos | sepay\namount: Number\nstatus: Enum(...)", fc="#dbeafe")
    draw_box(ax, 8.2, 1.5, 2.5, 2.2, "payments\n────────────\n_id: ObjectId\nappointmentId\nmethod: Enum(payos,sepay,wallet)\namount: Number\nstatus: Enum(...)", fc="#ecfeff")

    draw_arrow(ax, 3.3, 2.6, 4, 2.6, "1..*")
    draw_arrow(ax, 7.4, 2.6, 8.2, 2.6, "optional")
    save_fig(path, fig)


def gen_figure_12_package(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.text(5.5, 5.6, "Code Package Diagram — Payment modules", ha="center", fontsize=14, fontweight="bold")

    draw_box(ax, 0.4, 0.5, 10.2, 4.8, "", fc="#f8fafc", ec="#64748b")
    draw_box(ax, 0.8, 3.8, 2.5, 1.2, "client/pages\nPatientWalletPage\nPatientWalletCheckoutPage", fc="#cffafe")
    draw_box(ax, 3.6, 3.8, 2.5, 1.2, "client/services\npatientApi.js", fc="#cffafe")
    draw_box(ax, 6.4, 3.8, 3.5, 1.2, "server/controllers\nPatientWalletController\nPayosPaymentController\nSepayPaymentController", fc="#dbeafe")
    draw_box(ax, 0.8, 1.5, 4.8, 1.8, "server/services\npatientWallet.service.js\npayos.service.js\nsepay.service.js\nwallet.service.js", fc="#dbeafe")
    draw_box(ax, 6.0, 1.5, 3.9, 1.8, "server/models\nWalletModel\nWalletTransactionModel", fc="#ecfeff")
    draw_box(ax, 0.8, 0.7, 9.1, 0.6, "External: PayOS API  ·  SePay API", fc="#fef3c7", ec="#d97706", fontsize=10)

    draw_arrow(ax, 3.3, 4.4, 3.6, 4.4)
    draw_arrow(ax, 6.1, 4.4, 6.4, 4.4)
    draw_arrow(ax, 5.3, 3.8, 3.2, 3.3)
    draw_arrow(ax, 7.8, 3.8, 7.8, 3.3)
    save_fig(path, fig)


def gen_figure_19_ui_booking(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.text(5, 5.6, "UI — Book Appointment (Payment popup)", ha="center", fontsize=14, fontweight="bold")

    draw_box(ax, 1, 0.8, 8, 4.5, "", fc="#ffffff", ec="#cbd5e1")
    ax.text(5, 4.7, "Confirm booking & payment", ha="center", fontsize=12, fontweight="bold")
    ax.text(5, 4.1, "Consultation fee: 350,000 VND  (insurance discount applied)", ha="center", fontsize=10, color="#475569")

    for i, label in enumerate(["Medical Wallet", "PayOS", "SePay"]):
        y = 3.2 - i * 0.9
        fc = "#ecfeff" if i > 0 else "#f1f5f9"
        draw_box(ax, 2.2, y, 5.6, 0.65, f"○  {label}", fc=fc, fontsize=11)
    draw_box(ax, 6.5, 1.2, 2.2, 0.7, "Confirm", fc="#0891b2", ec="#0e7490")
    ax.text(7.6, 1.55, "Confirm", ha="center", va="center", color="white", fontsize=11, fontweight="bold")
    save_fig(path, fig)


def gen_figure_24_ui_wallet(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(10, 6.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    ax.text(5, 6.1, "UI — My Wallet (/patient/wallet)", ha="center", fontsize=14, fontweight="bold")

    draw_box(ax, 0.8, 4.5, 8.4, 1.2, "Balance: 1,250,000 VND", fc="#0891b2", ec="#0e7490")
    ax.text(5, 5.1, "Balance: 1,250,000 VND", ha="center", color="white", fontsize=16, fontweight="bold")

    ax.text(1.2, 4.0, "Top up", fontsize=12, fontweight="bold")
    draw_box(ax, 1.2, 3.0, 3.5, 0.75, "PayOS\nVietQR in-app checkout", fc="#ecfeff", fontsize=9)
    draw_box(ax, 5.0, 3.0, 3.5, 0.75, "SePay\nVietQR in-app checkout", fc="#ecfeff", fontsize=9)
    ax.text(1.2, 2.5, "Amount (min 10,000 VND):  [ 100,000 ]", fontsize=10)
    draw_box(ax, 6.8, 1.8, 2.0, 0.6, "Continue", fc="#0891b2", ec="#0e7490")
    ax.text(7.8, 2.1, "Continue", ha="center", color="white", fontsize=10, fontweight="bold")

    ax.text(1.2, 1.2, "Recent transactions", fontsize=11, fontweight="bold")
    draw_box(ax, 1.2, 0.4, 7.6, 0.65, "Top-up · PayOS · +100,000 · success", fc="#f8fafc", ec="#cbd5e1", fontsize=9)
    save_fig(path, fig)


def gen_figure_61_class_uc7(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.text(6, 6.6, "Class Diagram — UC-7: Book Appointment", ha="center", fontsize=14, fontweight="bold")

    classes = [
        (0.5, 4.8, "BookAppointmentPage"),
        (0.5, 3.2, "AppointmentApiClient"),
        (0.5, 1.6, "AppointmentController"),
        (4.0, 4.8, "AppointmentService"),
        (4.0, 3.2, "WalletService"),
        (4.0, 1.6, "PaymentModel"),
        (7.5, 4.8, "PayosService"),
        (7.5, 3.2, "SepayService"),
        (7.5, 1.6, "WalletModel"),
        (10.5, 4.8, "<<external>>\nPayOS"),
        (10.5, 3.2, "<<external>>\nSePay"),
    ]
    for x, y, name in classes:
        draw_box(ax, x, y, 2.3, 1.0, name, fc="#ecfeff" if "external" not in name else "#fef3c7", ec="#0891b2" if "external" not in name else "#d97706", fontsize=8)

    draw_arrow(ax, 2.8, 5.3, 4.0, 5.3)
    draw_arrow(ax, 2.8, 3.7, 4.0, 5.0)
    draw_arrow(ax, 2.8, 2.1, 4.0, 5.0)
    draw_arrow(ax, 6.3, 5.3, 7.5, 5.3)
    draw_arrow(ax, 6.3, 5.0, 7.5, 3.7)
    draw_arrow(ax, 6.3, 4.5, 4.0, 3.7)
    draw_arrow(ax, 6.3, 2.1, 7.5, 2.1)
    draw_arrow(ax, 9.8, 5.3, 10.5, 5.3)
    draw_arrow(ax, 9.8, 3.7, 10.5, 3.7)
    save_fig(path, fig)


def gen_sequence_diagram(path: Path, title: str, participants: list[str], steps: list[tuple], height: float = 8) -> None:
    n = len(participants)
    fig, ax = plt.subplots(figsize=(13, height))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, height)
    ax.axis("off")
    ax.text(6.5, height - 0.35, title, ha="center", fontsize=14, fontweight="bold")

    xs = [1.0 + i * (11.0 / max(n - 1, 1)) for i in range(n)]
    top = height - 1.2
    bottom = 1.0

    for x, name in zip(xs, participants):
        draw_box(ax, x - 0.55, top, 1.1, 0.55, name, fc="#ecfeff", fontsize=7)
        ax.plot([x, x], [top, bottom], linestyle="--", color="#94a3b8", linewidth=1)

    y = top - 0.8
    for i, step in enumerate(steps):
        frm, to, label = step
        x1, x2 = xs[frm], xs[to]
        dy = 0.55
        y -= dy
        color = "#334155"
        if x1 == x2:
            arr = FancyArrowPatch((x1 + 0.1, y), (x1 + 0.9, y), arrowstyle="-|>", mutation_scale=10, color=color)
        else:
            arr = FancyArrowPatch((x1, y), (x2, y), arrowstyle="-|>", mutation_scale=10, color=color)
        ax.add_patch(arr)
        ax.text((x1 + x2) / 2, y + 0.08, label, ha="center", fontsize=7, color="#475569")

    save_fig(path, fig)


def gen_figure_62_sequence_uc7(path: Path) -> None:
    participants = ["Patient", "BookPage", "Controller", "ApptService", "WalletSvc", "PayOS", "SePay", "DB"]
    steps = [
        (0, 1, "select slot + payment method"),
        (1, 2, "POST /appointments/book"),
        (2, 3, "bookAppointment()"),
        (3, 4, "alt: deduct wallet"),
        (3, 5, "alt: PayOS verify"),
        (3, 6, "alt: SePay verify"),
        (3, 7, "save appointment + payment"),
        (3, 1, "confirmation"),
        (1, 0, "booking success"),
    ]
    gen_sequence_diagram(path, "Sequence Diagram — UC-7: Book Appointment", participants, steps, height=7.5)


def gen_figure_71_class_uc12(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.text(6, 6.6, "Class Diagram — UC-12: View Wallet Balance", ha="center", fontsize=14, fontweight="bold")

    classes = [
        (0.5, 4.8, "WalletPage"),
        (0.5, 3.2, "WalletApiClient"),
        (0.5, 1.6, "WalletController"),
        (4.0, 4.8, "PatientWalletService"),
        (4.0, 3.2, "PayosService"),
        (4.0, 1.6, "SepayService"),
        (7.5, 4.8, "WalletModel"),
        (7.5, 3.2, "WalletTransactionModel"),
        (10.5, 4.8, "<<external>>\nPayOS"),
        (10.5, 3.2, "<<external>>\nSePay"),
    ]
    for x, y, name in classes:
        draw_box(ax, x, y, 2.3, 1.0, name, fc="#ecfeff" if "external" not in name else "#fef3c7", ec="#0891b2" if "external" not in name else "#d97706", fontsize=8)

    draw_arrow(ax, 2.8, 5.3, 4.0, 5.3)
    draw_arrow(ax, 2.8, 3.7, 4.0, 5.0)
    draw_arrow(ax, 2.8, 2.1, 4.0, 5.0)
    draw_arrow(ax, 6.3, 5.3, 7.5, 5.3)
    draw_arrow(ax, 6.3, 5.0, 4.0, 3.7)
    draw_arrow(ax, 6.3, 4.8, 4.0, 2.1)
    draw_arrow(ax, 6.3, 5.0, 7.5, 3.7)
    draw_arrow(ax, 9.8, 3.7, 10.5, 5.3)
    draw_arrow(ax, 9.8, 2.1, 10.5, 3.7)
    save_fig(path, fig)


def gen_figure_72_sequence_uc12(path: Path) -> None:
    participants = ["Patient", "WalletPage", "Controller", "WalletService", "PayOS/Sepay", "Gateway API", "DB"]
    steps = [
        (0, 1, "open My Wallet"),
        (1, 2, "GET /api/patient/wallet"),
        (2, 3, "getPatientWallet()"),
        (3, 6, "load balance + txns"),
        (0, 1, "Top up + select PayOS/SePay"),
        (1, 2, "POST /wallet/topups/{provider}"),
        (2, 3, "createTopup()"),
        (3, 4, "init checkout + VietQR"),
        (4, 5, "create payment / QR"),
        (3, 6, "pending transaction"),
        (1, 0, "show checkout page (VietQR)"),
        (0, 5, "pay via banking app"),
        (5, 2, "webhook / IPN"),
        (2, 3, "verify + completeTopup"),
        (3, 6, "credit wallet"),
        (1, 0, "receipt + updated balance"),
    ]
    gen_sequence_diagram(path, "Sequence Diagram — UC-12: View Wallet Balance (Top-up)", participants, steps, height=9)


GENERATORS = {
    2: gen_figure_2_usecase,
    10: gen_figure_10_erd,
    12: gen_figure_12_package,
    19: gen_figure_19_ui_booking,
    24: gen_figure_24_ui_wallet,
    47: gen_figure_12_package,  # same package diagram style
    61: gen_figure_61_class_uc7,
    62: gen_figure_62_sequence_uc7,
    71: gen_figure_71_class_uc12,
    72: gen_figure_72_sequence_uc12,
}


def build_figure_images() -> dict[int, Path]:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    paths: dict[int, Path] = {}
    for fig_num, gen in GENERATORS.items():
        out = ASSETS_DIR / f"figure_{fig_num:02d}.png"
        gen(out)
        paths[fig_num] = out
    return paths


def map_figures_to_embeds(doc: Document) -> dict[int, str]:
    body = doc.element.body
    last_embed: str | None = None
    mapping: dict[int, str] = {}

    for child in body:
        tag = child.tag.split("}")[-1]
        if tag != "p":
            continue
        blips = child.findall(".//" + qn("a:blip"))
        for b in blips:
            last_embed = b.get(qn("r:embed"))
        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        m = re.match(r"Figure\s+(\d+):", text)
        if m:
            fig_num = int(m.group(1))
            if fig_num in TARGET_FIGURES and last_embed and fig_num not in mapping:
                mapping[fig_num] = last_embed
            # keep updating last_embed for duplicate captions — prefer design section later
            if fig_num in TARGET_FIGURES and last_embed:
                mapping[fig_num] = last_embed

    return mapping


def replace_embedded_images(docx_path: Path, fig_to_embed: dict[int, str], fig_images: dict[int, Path]) -> None:
    with zipfile.ZipFile(docx_path, "r") as z:
        rels_xml = z.read("word/_rels/document.xml.rels").decode("utf-8")

    embed_to_media: dict[str, str] = {}
    for embed in fig_to_embed.values():
        m = re.search(rf'Id="{re.escape(embed)}"[^>]+Target="([^"]+)"', rels_xml)
        if not m:
            m = re.search(rf'Target="([^"]+)"[^>]+Id="{re.escape(embed)}"', rels_xml)
        if m:
            embed_to_media[embed] = "word/" + m.group(1).lstrip("/")

    media_replacements: dict[str, bytes] = {}
    for fig_num, embed in fig_to_embed.items():
        media_path = embed_to_media.get(embed)
        if media_path and fig_num in fig_images:
            media_replacements[media_path] = fig_images[fig_num].read_bytes()

    with zipfile.ZipFile(docx_path, "r") as zin:
        buf = BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = media_replacements.get(item.filename, zin.read(item.filename))
                zout.writestr(item, data)
        docx_path.write_bytes(buf.getvalue())


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"Missing {DOC_PATH}")

    shutil.copy2(DOC_PATH, BACKUP_PATH)
    print(f"Backup: {BACKUP_PATH}")

    fig_images = build_figure_images()
    print(f"Generated {len(fig_images)} diagram PNGs in {ASSETS_DIR}")

    doc = Document(DOC_PATH)
    changed = replace_doc_text(doc)
    print(f"Updated text blocks: {changed}")

    fig_to_embed = map_figures_to_embeds(doc)
    print("Figure -> embed mapping:")
    for k in sorted(fig_to_embed):
        if k in TARGET_FIGURES:
            print(f"  Figure {k}: {fig_to_embed[k]}")

    doc.save(DOC_PATH)

    replace_embedded_images(DOC_PATH, fig_to_embed, fig_images)
    print(f"Saved: {DOC_PATH}")

    # verify
    doc2 = Document(DOC_PATH)
    momo = sepay = 0
    for p in doc2.paragraphs:
        momo += p.text.count("Momo") + p.text.count("momo")
        sepay += p.text.count("SePay") + p.text.count("sepay")
    for t in doc2.tables:
        for row in t.rows:
            for cell in row.cells:
                momo += cell.text.count("Momo") + cell.text.count("momo")
                sepay += cell.text.count("SePay") + cell.text.count("sepay")
    print(f"Verify text — momo refs: {momo}, sepay refs: {sepay}")


if __name__ == "__main__":
    main()
