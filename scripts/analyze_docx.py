from docx import Document
from pathlib import Path

files = [
    r"c:\Users\nguye\Downloads\Iteration 1.docx",
    r"c:\Users\nguye\Downloads\Iteration 1 (1).docx",
    r"c:\Users\nguye\Downloads\Iteration 1 (2).docx",
    r"c:\Users\nguye\Downloads\Iteration 1 (3).docx",
    r"e:\SU26\WDP301\WDP301-SE1816-GROUP4_RDSDocument.docx",
]

for f in files:
    p = Path(f)
    if not p.exists():
        print(f"MISSING: {f}")
        continue
    doc = Document(f)
    print(f"\n=== {p.name} ({len(doc.paragraphs)} paras, {len(doc.inline_shapes)} shapes, {len(doc.tables)} tables) ===")
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t:
            style = para.style.name if para.style else ""
            print(f"  [{i}] ({style}) {t[:150]}")
