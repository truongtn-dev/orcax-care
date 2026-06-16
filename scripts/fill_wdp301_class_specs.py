"""Fill WDP301 class specification placeholder tables."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from class_spec_catalog import get_specs, normalize

DOC_PATH = Path(r"E:\SU26\WDP301\WDP301-SE1816-GROUP4_Document.docx")
BACKUP_PATH = DOC_PATH.with_name("WDP301-SE1816-GROUP4_Document.before-class-specs.docx")


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "CMU Sans Serif"
            run.font.size = Pt(11)


def fill_table(table, class_name: str) -> int:
    specs = get_specs(class_name)
    # keep header row
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for idx, (method, desc) in enumerate(specs, start=1):
        row = table.add_row()
        set_cell_text(row.cells[0], f"{idx:02d}")
        set_cell_text(row.cells[1], method)
        set_cell_text(row.cells[2], desc)
    return len(specs)


def iter_placeholder_tables(doc: Document):
    in_design = False
    section = ""
    pending_class = None
    in_spec = False
    tbl_i = 0

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
            if "II. Requirement Specifications" in text:
                in_design = True
                continue
            if not in_design:
                continue
            if re.match(r"^\d+\.\d+\s+\S", text) and "Figure" not in text:
                section = text
                in_spec = False
            elif text == "b. Class Specifications":
                in_spec = True
                pending_class = None
            elif text.startswith("c. Sequence"):
                in_spec = False
                pending_class = None
            elif in_spec and text and text not in ("No", "Method", "Description"):
                if not text.startswith("Figure") and "Class Diagram" not in text:
                    pending_class = text.strip()
        elif tag == "tbl" and in_design:
            table = doc.tables[tbl_i]
            hdr = [c.text.strip() for c in table.rows[0].cells]
            if hdr[:3] == ["No", "Method", "Description"] and len(table.rows) > 1:
                row1 = [c.text.strip() for c in table.rows[1].cells]
                if row1 and "see class diagram" in row1[1].lower():
                    yield {
                        "table_index": tbl_i,
                        "table": table,
                        "section": section,
                        "class": pending_class or "?",
                    }
                    pending_class = None
            tbl_i += 1


def main() -> None:
    if not DOC_PATH.exists():
        raise SystemExit(f"Missing document: {DOC_PATH}")

    shutil.copy2(DOC_PATH, BACKUP_PATH)
    print(f"Backup: {BACKUP_PATH}")

    doc = Document(DOC_PATH)
    total = 0
    methods = 0
    missing = []

    for item in iter_placeholder_tables(doc):
        cls = item["class"]
        count = fill_table(item["table"], cls)
        total += 1
        methods += count
        if count <= 1 and normalize(cls) not in {
            "AppHeader",
            "AuthContext",
        }:
            missing.append((item["section"], cls, item["table_index"]))

    doc.save(DOC_PATH)
    print(f"Updated tables: {total}")
    print(f"Total method rows written: {methods}")
    if missing:
        print("Tables with minimal/generic fallback:", len(missing))
        for sec, cls, ti in missing[:15]:
            print(f"  - {sec} | {cls} | table {ti}")


if __name__ == "__main__":
    main()
