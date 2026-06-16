"""
Build / merge SDS into WDP301-SE1816-GROUP4_Document.docx.

- Standalone: WDP301-SE1816-GROUP4_SDSDocument.docx
- Merge: inserts section IV. Software Design Specifications before Appendix (renamed V)
- Figure captions use Word SEQ fields continuing from the main document (after Figure 46)
"""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

from build_design_specs import SECTIONS  # noqa: E402
from fix_figure_fields import fix_figure_captions  # noqa: E402
from sds_functions_data import FEATURE_KEYS  # noqa: E402

MAIN_DOC = Path(r"e:/SU26/WDP301/WDP301-SE1816-GROUP4_Document.docx")
STANDALONE_OUT = Path(r"e:/SU26/WDP301/WDP301-SE1816-GROUP4_SDSDocument.docx")
SDS_EXPORT = ROOT / "docs" / "sds" / "export"
SDS_PUML = ROOT / "docs" / "sds"
DB_EXPORT = ROOT / "docs" / "database" / "export"
HEADER_BG = "FCE4D6"

SDS_SECTION_TITLE = "V. Software Design Specifications"
APPENDIX_TITLE = "IV. Appendix"
LEGACY_SDS_TITLES = (
    "IV. Software Design Specifications",
    "V. Software Design Specifications",
)

CAPTION_RE = re.compile(r"^Figure\s+(\d+)", re.I)
CLASS_RE = re.compile(r"^class\s+(\w+)\s*\{", re.M)
METHOD_RE = re.compile(r"^\s*\+(\w+)\([^)]*\)\s*:\s*(.+)$", re.M)


def shade(cell, hex_color):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _field_run(p, fld_type):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = p.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), fld_type)
    run._r.append(el)


def _instr_run(p, instr):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = p.add_run()
    el = OxmlElement("w:instrText")
    el.set(qn("xml:space"), "preserve")
    el.text = instr
    run._r.append(el)


def detect_last_figure_number(doc: Document) -> int:
    max_n = 0
    for p in doc.paragraphs:
        if p.style.name != "Caption":
            continue
        m = CAPTION_RE.match(p.text.strip())
        if m:
            max_n = max(max_n, int(m.group(1)))
    return max_n


def add_caption(w: DocWriter, text: str, fig: dict) -> None:
    fig["n"] += 1
    p = w.doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Figure ")
    _field_run(p, "begin")
    _instr_run(p, r" SEQ Figure \* ARABIC ")
    _field_run(p, "separate")
    p.add_run(str(fig["n"]))
    _field_run(p, "end")
    p.add_run(f": {text}")
    for run in p.runs:
        run.font.size = Pt(10)
    w._commit(p._element)


def add_image(w: DocWriter, path: Path, fig: dict, width=6.0, caption=None):
    if not path.exists():
        p = w.doc.add_paragraph()
        run = p.add_run(f"[ Missing diagram: {path.name} ]")
        run.italic = True
        w._commit(p._element)
        return
    p = w.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path.resolve()), width=Inches(width))
    w._commit(p._element)
    if caption:
        add_caption(w, caption, fig)


def add_method_table(w: DocWriter, class_name: str, methods: list[tuple[str, str]]):
    p = w.doc.add_paragraph()
    run = p.add_run(class_name)
    run.bold = True
    w._commit(p._element)
    t = w.doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    for i, h in enumerate(["No", "Method", "Description"]):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        shade(cell, HEADER_BG)
    for i, (method, desc) in enumerate(methods, start=1):
        cells = t.add_row().cells
        for j, v in enumerate([f"{i:02d}", method, desc]):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(v)
            r.font.size = Pt(10)
    for row in t.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(2.7)
        row.cells[2].width = Inches(3.3)
    w._commit(t._tbl)


def add_code(w: DocWriter, text: str):
    for line in text.strip("\n").split("\n"):
        p = w.doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        w._commit(p._element)


def parse_puml_classes(key: str) -> list[tuple[str, list[tuple[str, str]]]]:
    path = SDS_PUML / f"CD_{key}.puml"
    if not path.exists():
        return []
    text = path.read_text(encoding="utf-8")
    classes = []
    for m in CLASS_RE.finditer(text):
        name = m.group(1)
        start = m.end()
        end = text.find("}", start)
        body = text[start:end] if end != -1 else ""
        methods = []
        for mm in METHOD_RE.finditer(body):
            sig = f"{mm.group(1)}(...)"
            ret = mm.group(2).strip()
            methods.append((sig, f"Returns {ret}."))
        if not methods:
            methods.append(("(see class diagram)", "Public operations shown in the class diagram."))
        label = name
        if "Model" in name and "(Mongoose)" not in label:
            label += " (Mongoose)"
        classes.append((label, methods))
    return classes


def build_functions_list() -> list[dict]:
    functions = []
    for sec in SECTIONS:
        feature = sec["feature"]
        key = FEATURE_KEYS.get(feature)
        if not key:
            continue
        cd = SDS_PUML / f"CD_{key}.puml"
        sq = SDS_PUML / f"SQ_{key}.puml"
        if not cd.exists() or not sq.exists():
            continue
        classes = parse_puml_classes(key)
        functions.append(
            {
                "name": sec.get("name") or feature,
                "key": key,
                "uc": sec.get("uc", ""),
                "classes": classes,
                "queries": sec.get("queries", "").strip(),
            }
        )
    return functions


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "CMU Sans Serif"
    normal.font.size = Pt(12)


class DocWriter:
    """Append to doc, or insert each block before *anchor* (keeps image relationships valid)."""

    def __init__(self, doc: Document, anchor_paragraph: Paragraph | None = None):
        self.doc = doc
        self.anchor = anchor_paragraph._element if anchor_paragraph is not None else None

    def _commit(self, element):
        if self.anchor is None:
            return
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)
        self.anchor.addprevious(element)
        self.anchor = element  # preserve top-to-bottom order

    def add_heading(self, text: str, level: int):
        p = self.doc.add_heading(text, level=level)
        self._commit(p._element)

    def add_paragraph(self, text: str = "", style: str | None = None):
        p = self.doc.add_paragraph(text, style=style) if style else self.doc.add_paragraph(text)
        self._commit(p._element)
        return p

    def add_empty_paragraph(self):
        p = self.doc.add_paragraph()
        self._commit(p._element)
        return p


def add_heading(w: DocWriter, text: str, level: int):
    w.add_heading(text, level)


def build_standalone():
    doc = Document()
    setup_styles(doc)
    w = DocWriter(doc)
    fig = {"n": 0}
    add_heading(w, "I. Overview", 1)
    _add_overview(w, fig, h_packages=2, h_schema=2, h_table=3, h_schema_sub=3)
    add_heading(w, "II. Code Designs", 1)
    functions = _add_code_designs(w, fig, h_fn=2, h_sub=3)
    try:
        doc.save(STANDALONE_OUT)
        print(f"Standalone SDS: {STANDALONE_OUT} | functions: {len(functions)}")
    except PermissionError:
        alt = STANDALONE_OUT.with_name(STANDALONE_OUT.stem + "_new.docx")
        doc.save(alt)
        print(f"Standalone locked. Saved: {alt}")


def _add_overview(w: DocWriter, fig, h_packages=2, h_schema=2, h_table=3, h_schema_sub=3):
    add_heading(w, "1. Code Packages", h_packages)
    p = w.doc.add_paragraph(
        "OrcaXCare follows the MERN architecture. The client package (React) "
        "communicates with the server package (Express) through a REST API; the "
        "server is organized in layers (routes \u2192 controllers \u2192 services \u2192 "
        "models) on top of MongoDB."
    )
    w._commit(p._element)
    add_image(w, DB_EXPORT / "CodePackages.png", fig, 5.8, "Code Package Diagram of OrcaXCare")
    add_heading(w, "2. Database Design", h_schema)
    add_heading(w, "a. Database Schema", h_schema_sub)
    add_image(w, DB_EXPORT / "DB_Schema_NoPkg.png", fig, 6.2, "Database Schema of OrcaXCare")
    add_heading(w, "b. Table Description", h_schema_sub)
    p = w.doc.add_paragraph(
        "The detailed description of all collections (primary keys, foreign keys, "
        "and attributes) is provided in the RDS document, section Database Design "
        "\u2013 Table Descriptions."
    )
    p.runs[0].italic = True
    w._commit(p._element)


def _add_code_designs(w: DocWriter, fig, h_fn=3, h_sub=4, prefix=""):
    functions = build_functions_list()
    for i, fn in enumerate(functions, start=1):
        title = f"{prefix}{i}. {fn['name']}" if prefix else f"2.{i} {fn['name']}"
        add_heading(w, title, h_fn)
        uc = fn.get("uc", "")
        uc_label = f"{uc}: " if uc else ""
        add_heading(w, "a. Class Diagram", h_sub)
        add_image(w, SDS_EXPORT / f"CD_{fn['key']}.png", fig, 5.4, f"Class Diagram \u2013 {uc_label}{fn['name']}")
        add_heading(w, "b. Class Specifications", h_sub)
        for class_name, methods in fn["classes"]:
            add_method_table(w, class_name, methods)
        add_heading(w, "c. Sequence Diagram(s)", h_sub)
        add_image(w, SDS_EXPORT / f"SQ_{fn['key']}.png", fig, 6.2, f"Sequence Diagram \u2013 {uc_label}{fn['name']}")
        add_heading(w, "d. Database Queries", h_sub)
        if fn["queries"]:
            add_code(w, fn["queries"])
        w.add_empty_paragraph()
    return functions


def build_sds_body(w: DocWriter, fig: dict, h_overview=2, h_designs=2, h_fn=3, h_sub=4):
    add_heading(w, "1. Overview", h_overview)
    _add_overview(w, fig, h_packages=h_sub, h_schema=h_sub, h_table=h_sub + 1, h_schema_sub=h_sub + 1)
    add_heading(w, "2. Code Designs", h_designs)
    return _add_code_designs(w, fig, h_fn=h_fn, h_sub=h_sub)


def remove_existing_sds_section(doc: Document) -> bool:
    """Remove SDS block appended at end (re-merge safe)."""
    start_el = None
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() in LEGACY_SDS_TITLES:
            start_el = p._element
            break
    if start_el is None:
        return False
    body = doc.element.body
    to_remove = []
    removing = False
    for child in list(body):
        if child is start_el:
            removing = True
        if removing:
            if child.tag.endswith("sectPr"):
                continue
            to_remove.append(child)
    for child in to_remove:
        body.remove(child)
    return True


def restore_appendix_heading(doc: Document) -> None:
    """Undo mistaken rename IV Appendix -> V Appendix."""
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "V. Appendix":
            p.text = APPENDIX_TITLE
            return


def insert_elements_before(ref_element, elements):
    body = ref_element.getparent()
    idx = body.index(ref_element)
    for el in elements:
        body.insert(idx, el)
        idx += 1


def count_image_rels(doc: Document) -> int:
    return sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)


def merge_into_main():
    if not MAIN_DOC.exists():
        print(f"Main document not found: {MAIN_DOC}")
        return

    backup = MAIN_DOC.with_suffix(".docx.bak2")
    shutil.copy2(MAIN_DOC, backup)
    print(f"Backup: {backup}")

    main = Document(MAIN_DOC)
    setup_styles(main)

    restore_appendix_heading(main)
    removed = remove_existing_sds_section(main)
    if removed:
        print("Removed previous SDS section.")

    last_fig = detect_last_figure_number(main)
    print(f"Last figure in main doc: {last_fig} -> SDS continues at {last_fig + 1}")

    imgs_before = count_image_rels(main)

    # Append section V at end (after IV. Appendix) — do NOT insert before Appendix
    pb = main.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)

    w = DocWriter(main)
    fig = {"n": last_fig}

    add_heading(w, SDS_SECTION_TITLE, 1)
    functions = build_sds_body(w, fig, h_overview=2, h_designs=2, h_fn=3, h_sub=4)

    imgs_after = count_image_rels(main)
    print(f"Embedded images: {imgs_after - imgs_before} new (total {imgs_after})")

    converted, reset_fixed = fix_figure_captions(main)
    if converted or reset_fixed:
        print(
            f"Figure fields fixed: {len(converted)} plain captions converted, "
            f"{reset_fixed} reset flags removed"
        )

    try:
        main.save(MAIN_DOC)
        print(f"Merged SDS into: {MAIN_DOC}")
        print(f"Section order: ... III Design Specs -> IV Appendix -> V SDS")
        print(f"Functions: {len(functions)} | Figures added: {len(functions)*2 + 2}")
        print(f"Next figure number after merge: {fig['n']} (update fields in Word with F9)")
    except PermissionError:
        alt = MAIN_DOC.with_name(MAIN_DOC.stem + "_merged.docx")
        main.save(alt)
        print(f"Main doc locked. Saved: {alt}")


def main():
    import sys
    if "--render" in sys.argv:
        render_script = ROOT / "scripts" / "render_sds_diagrams.py"
        if render_script.exists():
            import subprocess
            subprocess.run([sys.executable, str(render_script)], check=True)
    merge_into_main()
    build_standalone()


if __name__ == "__main__":
    main()
