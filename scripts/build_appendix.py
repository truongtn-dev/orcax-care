"""
Build Section IV. Appendix - 1. Assumptions & Dependencies for OrcaXCare,
following the FBUS sample layout (AS-x / DE-x labels with hanging indent).
Output: e:/SU26/WDP301/WDP301-SE1816-GROUP4_Appendix.docx
"""
from __future__ import annotations

from pathlib import Path

import re

from docx import Document
from docx.shared import Inches, Pt

OUT_PATH = Path(r"e:/SU26/WDP301/WDP301-SE1816-GROUP4_Appendix.docx")
GROUP_DOC = Path(r"e:/SU26/WDP301/WDP301-SE1816-GROUP4_Document.docx")
HEADER_BG = "FCE4D6"


def collect_business_rules():
    """Collect BR rows from the per-UC 'b. Business Rules' tables in the group
    document (header 'Business Rule Description'). The leftover FBUS table uses
    'Rule Definition' as header and is skipped automatically."""
    doc = Document(GROUP_DOC)
    brs: dict[str, tuple[str, str]] = {}
    for t in doc.tables:
        if len(t.columns) != 3 or not t.rows:
            continue
        hdr = [c.text.strip().lower() for c in t.rows[0].cells]
        if hdr[0] != "id" or hdr[2] != "business rule description":
            continue
        for r in t.rows[1:]:
            bid = r.cells[0].text.strip()
            if not re.match(r"^BR-\d+$", bid):
                continue
            brs.setdefault(bid, (r.cells[1].text.strip(), r.cells[2].text.strip()))
    return [(k, *brs[k]) for k in sorted(brs, key=lambda x: int(x.split("-")[1]))]


def shade(cell, hex_color):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

ASSUMPTIONS = [
    "Systems with appropriate user interfaces will be available for Admins, Staff, Doctors, Patients, and Guests to interact effectively with the OrcaXCare system across supported devices (desktop and mobile browsers).",
    "Doctors and clinic rooms assigned in work shifts will be available as scheduled; appointment slots generated from shifts reflect the clinic's real operating capacity.",
    "Users will provide valid email addresses during registration, and the email system will be able to send and receive verification and password-reset messages successfully.",
    "Patients will arrive at the clinic and check in at the reception desk before their appointment time; reception staff verify patient identity before issuing queue tickets.",
    "Doctors are licensed professionals responsible for the clinical content they record (encounter notes, diagnoses, prescriptions); the system only supports, not replaces, clinical judgment.",
    "The ICD-10 diagnosis catalog and the medicines formulary are seeded into the database before the consultation and prescription features are used.",
    "Patients who select online payment own a valid PayOS or Momo account/app with sufficient balance to complete the transaction.",
    "Insurance cards registered by patients belong to supported providers, and the discount policies configured in the system match the providers' actual agreements with the clinic.",
    "All users operate in the Asia/Ho_Chi_Minh timezone; all displayed dates and times follow this timezone.",
    "The product backlog represents the agreed scope; features are developed iteration by iteration following the backlog priorities.",
]

DEPENDENCIES = [
    "PayOS payment gateway APIs must be available and reachable for appointment fee payment and wallet top-up; sandbox credentials are used during development.",
    "Momo payment gateway APIs must be available as the alternative payment method.",
    "A third-party email delivery service (SMTP provider) is required for account verification, password reset, booking confirmation, and notification emails.",
    "Cloud media storage (e.g., Cloudinary) is required for storing avatars, insurance card images, and medical images uploaded by doctors.",
    "MongoDB Atlas (or an equivalent MongoDB hosting service) is required as the primary database of the system.",
    "An OCR service is required for the insurance card scanning feature (UC-13.2.1); when unavailable, manual input remains the fallback.",
    "The deployment environment (Node.js hosting for the Express API and static hosting for the React client) must support HTTPS for secure token-based authentication.",
]


LIMITATIONS = [
    "The system does not support booking appointment slots in the past or slots whose start time has already passed.",
    "No dedicated mobile application is provided; the platform is accessible only via supported web browsers (responsive design).",
    "OrcaXCare handles only online appointment booking through the platform; walk-in patients are registered manually by Admin/Staff at the clinic.",
    "Payment gateways (PayOS, Momo) run in sandbox mode during this course project; no real financial transactions are processed.",
    "Insurance card OCR is best-effort assistance for data entry only; the system does not connect to any national insurance database for card validation.",
    "The EMR module is internal to OrcaXCare; interoperability standards (HL7/FHIR) and data exchange with external hospitals are not supported.",
    "The prescription module does not perform drug\u2013drug interaction checking; clinical correctness remains the prescribing doctor's responsibility.",
    "No telemedicine features (video consultation, remote diagnosis) are included in the current scope.",
    "Notifications are delivered in-app and via email only; SMS and mobile push notifications are not supported.",
    "The system supports image uploads (avatars, insurance cards, medical images) but lacks advanced features such as content moderation or automatic image analysis.",
    "Multi-language support is not implemented; all interfaces and notifications are in English only.",
    "Reporting is limited to the revenue dashboard with date/doctor filters and Excel export; no advanced business-intelligence or forecasting features are included.",
    "No AI-based recommendation features (doctor suggestions, diagnosis assistance) are included in the current scope.",
]


GLOSSARY = [
    ("Appointment", "A booking record linking a patient with a doctor at a specific time slot and clinic room, including reason, fee, and status (confirmed, checked-in, completed, cancelled)."),
    ("Appointment Slot", "A bookable time interval generated from a doctor's work shift; can be available, held, booked, or blocked."),
    ("Work Shift", "A scheduled working period of a doctor in a clinic room, created by Admin; the source for appointment slot generation."),
    ("Queue Session", "An active serving session of a clinic room opened by the doctor; queue tickets are issued into it and called in order."),
    ("Queue Ticket", "A numbered ticket issued by reception staff for a checked-in appointment, determining the patient's serving order."),
    ("Check-in", "The reception step confirming the patient has arrived for the appointment; required before a queue ticket is issued."),
    ("Encounter", "A clinical visit record created when a consultation starts; contains notes, diagnoses, images, and prescriptions."),
    ("EMR", "Electronic Medical Record - the patient's medical history consisting of signed-off encounters, diagnoses, prescriptions, and medical images."),
    ("Sign-off", "The doctor's final confirmation that locks an encounter; signed-off encounters become read-only."),
    ("Diagnosis (ICD-10)", "A medical conclusion recorded in an encounter, referencing a code from the WHO ICD-10 classification catalog."),
    ("Prescription", "A list of medicines with quantity, dosage, frequency, and duration issued by a doctor for an encounter; exportable as PDF/QR."),
    ("Formulary", "The clinic's catalog of medicines managed by staff, including unit, price, stock quantity, and reorder level."),
    ("Wallet", "A patient's internal balance used to pay consultation fees; topped up via payment gateways and credited with refunds."),
    ("Transaction", "An immutable financial record of a top-up, payment, or refund, including method, amount, and status."),
    ("Insurance Card", "A patient-registered health insurance card from a supported provider; valid cards yield booking discounts."),
    ("OCR", "Optical Character Recognition - a service that reads insurance card images to auto-fill card fields."),
    ("Complaint", "A service feedback record submitted by a patient, handled by staff with replies and status (Open, In Progress, Resolved)."),
    ("Specialty", "A medical field (e.g., Cardiology) assigned to doctors and used as a public search filter."),
    ("Department", "An organizational unit of the clinic grouping doctors and clinic rooms."),
    ("Clinic Room", "A physical examination room with a unique code, used in work shifts, queue sessions, and printed on tickets."),
    ("Guest", "An unregistered visitor who can browse public portal content and register a patient account."),
    ("Patient", "A registered user who books appointments and manages their own health-related data."),
    ("Doctor", "A clinical user who performs consultations, records EMR data, and manages queue sessions."),
    ("Staff", "An operational user handling reception check-in, pharmacy inventory, and complaint responses."),
    ("Admin", "A system administrator managing accounts, master data, doctors, and patients."),
    ("PayOS / Momo", "Third-party payment gateways used for appointment fee payment and wallet top-up."),
]

REFERENCES = [
    ("IEEE Std 830-1998 \u2013 IEEE Recommended Practice for Software Requirements Specifications.", None),
    ("MongoDB Documentation.", "https://www.mongodb.com/docs/"),
    ("Mongoose ODM Documentation.", "https://mongoosejs.com/docs/"),
    ("Express.js Documentation.", "https://expressjs.com/"),
    ("React Documentation.", "https://react.dev/"),
    ("Node.js Documentation.", "https://nodejs.org/en/docs"),
    ("JWT (JSON Web Token) Introduction.", "https://jwt.io/introduction"),
    ("bcrypt Password Hashing Function.", "https://github.com/kelektiv/node.bcrypt.js"),
    ("SMTP Email Standards \u2013 RFC 5321.", "https://datatracker.ietf.org/doc/html/rfc5321"),
    ("WHO ICD-10 Classification.", "https://icd.who.int/browse10/2019/en"),
    ("PayOS Developer Documentation.", "https://payos.vn/docs/"),
    ("MoMo Developer Documentation.", "https://developers.momo.vn/v3/"),
    ("Cloudinary Documentation.", "https://cloudinary.com/documentation"),
]


def add_hyperlink(paragraph, url, text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        p.add_run(text)


def add_items(doc, prefix, items):
    for i, text in enumerate(items, start=1):
        table = None  # use paragraph with hanging indent like the FBUS sample
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.9)
        pf.first_line_indent = Inches(-0.9)
        pf.space_after = Pt(6)
        label = p.add_run(f"{prefix}-{i}:\t")
        label.bold = True
        p.add_run(text)


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    doc.add_heading("IV. Appendix", level=1)
    doc.add_heading("1. Assumptions & Dependencies", level=2)

    p = doc.add_paragraph()
    run = p.add_run("Assumptions")
    run.bold = True
    run.italic = True
    add_items(doc, "AS", ASSUMPTIONS)

    p = doc.add_paragraph()
    run = p.add_run("Dependencies")
    run.bold = True
    run.italic = True
    add_items(doc, "DE", DEPENDENCIES)

    doc.add_heading("2. Limitations & Exclusions", level=2)
    add_bullets(doc, LIMITATIONS)

    doc.add_heading("3. Business Rules", level=2)
    rules = collect_business_rules()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [Inches(0.8), Inches(1.4), Inches(4.3)]
    for i, h in enumerate(["ID", "Category", "Rule Definition"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        shade(cell, HEADER_BG)
    for bid, cat, desc in rules:
        cells = table.add_row().cells
        for i, v in enumerate([bid, cat, desc]):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(v)
            run.font.size = Pt(11)
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
    print(f"Business rules collected: {len(rules)} ({rules[0][0]} .. {rules[-1][0]})")

    doc.add_heading("4. Glossary", level=2)
    gt = doc.add_table(rows=1, cols=2)
    gt.style = "Table Grid"
    for i, h in enumerate(["Term", "Definition"]):
        cell = gt.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        shade(cell, HEADER_BG)
    for term, definition in GLOSSARY:
        cells = gt.add_row().cells
        cells[0].text = ""
        run = cells[0].paragraphs[0].add_run(term)
        run.bold = True
        run.font.size = Pt(11)
        cells[1].text = ""
        run = cells[1].paragraphs[0].add_run(definition)
        run.font.size = Pt(11)
    for row in gt.rows:
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(4.8)

    doc.add_heading("5. References", level=2)
    for text, url in REFERENCES:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        p.add_run(text + (" " if url else ""))
        if url:
            add_hyperlink(p, url, url)

    try:
        doc.save(OUT_PATH)
        print(f"Saved: {OUT_PATH} | AS: {len(ASSUMPTIONS)}, DE: {len(DEPENDENCIES)}, LIM: {len(LIMITATIONS)}")
    except PermissionError:
        alt = OUT_PATH.with_name(OUT_PATH.stem + "_new.docx")
        doc.save(alt)
        print(f"Original locked (open in Word). Saved to: {alt}")


if __name__ == "__main__":
    main()
