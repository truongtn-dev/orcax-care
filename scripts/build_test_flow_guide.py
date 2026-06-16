"""
Generate OrcaXCare manual test flow Word document.
Output: docs/OrcaXCare_Test_Flow_Guide.docx
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "OrcaXCare_Test_Flow_Guide.docx"

AUTHOR = "TruongNTCE180140"


def set_cell_shading(cell, hex_color: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_step_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    headers = ["Bước", "Thao tác", "Kết quả mong đợi"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "E0F2FE")
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True

    for ri, (step, action, expected) in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        cells[0].text = str(step)
        cells[1].text = action
        cells[2].text = expected

    doc.add_paragraph()


def build():
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("OrcaXCare — Hướng dẫn test luồng chức năng")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(f"Scheduling + Wallet/Payment  |  Người lập: {AUTHOR}")
    s.font.size = Pt(11)
    s.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    add_para(
        doc,
        "Tài liệu này dùng để test thủ công trên UI theo đúng code hiện tại. "
        "Thanh toán: PayOS + SePay REAL (không mock). Momo đã thay bằng SePay.",
        size=10,
    )

    add_heading(doc, "1. Danh mục chức năng", 2)

    features = [
        ("Create Work Shift", "Scheduling", "Admin tạo template ca tuần: doctor, day, giờ, max patients. Chặn overlap."),
        ("Work Shifts List", "Scheduling", "Admin/doctor xem lưới tuần. Admin lọc theo bác sĩ."),
        ("Update Work Shift", "Scheduling", "Sửa giờ/sức chứa/phòng. Regenerate slot tương lai nếu có booking."),
        ("Delete Work Shift", "Scheduling", "Xóa template; xóa slot trống tương lai; giữ slot booked."),
        ("Generate Appointment Slots", "Scheduling", "Sinh slot từ template; skip holiday; không duplicate; không quá khứ."),
        ("Doctor Schedule Calendar", "Scheduling", "Lưới tuần/ngày, màu available/booked/blocked."),
        ("Block/Unlock Timeslot", "Scheduling", "Doctor khóa/mở từng slot; không block booked/past."),
        ("SePay Payment", "Wallet & Payment", "Nạp ví SePay sandbox — QR/chuyển khoản, IPN, polling."),
        ("PayOS Payment", "Wallet & Payment", "Nạp ví PayOS — VietQR, webhook, hủy pending, receipt."),
    ]

    ft = doc.add_table(rows=1 + len(features), cols=4)
    ft.style = "Table Grid"
    fh = ["Chức năng", "Module", "Mô tả (cập nhật theo code)", "Owner"]
    for i, h in enumerate(fh):
        ft.rows[0].cells[i].text = h
        set_cell_shading(ft.rows[0].cells[i], "F1F5F9")
    for ri, row in enumerate(features, start=1):
        for ci, val in enumerate(row):
            ft.rows[ri].cells[ci].text = val
        ft.rows[ri].cells[3].text = AUTHOR

    doc.add_paragraph()

    add_heading(doc, "2. Chuẩn bị", 2)
    add_bullets(
        doc,
        [
            "Terminal 1: cd server && npm run dev  (port 5000)",
            "Terminal 2: cd client && npm run dev  (http://localhost:5173)",
            "server/.env: MONGODB_URI, AUTO_SEED=true",
            "Payment REAL: PAYOS_MOCK=false, SEPAY_MOCK=false + credentials PayOS/SePay",
            "API_PUBLIC_ORIGIN=http://localhost:5000 (hoặc ngrok cho webhook/IPN)",
        ],
    )

    add_heading(doc, "3. Tài khoản test", 2)
    acc = doc.add_table(rows=5, cols=4)
    acc.style = "Table Grid"
    acc.rows[0].cells[0].text = "Vai trò"
    acc.rows[0].cells[1].text = "Email"
    acc.rows[0].cells[2].text = "Mật khẩu"
    acc.rows[0].cells[3].text = "Dùng cho"
    for c in acc.rows[0].cells:
        set_cell_shading(c, "F1F5F9")
    data = [
        ("Admin", "admin@orcaxcare.com", "Admin@123", "Shift, generate slots"),
        ("Doctor", "doctor.an@orcaxcare.com", "Doctor@123", "Calendar, block/unlock"),
        ("Patient", "patient@orcaxcare.com", "Patient@123", "Wallet PayOS/SePay"),
        ("Bác sĩ gợi ý", "Dr. Nguyen Van An", "—", "Monday 08:00–12:00, max 8"),
    ]
    for ri, row in enumerate(data, start=1):
        for ci, val in enumerate(row):
            acc.rows[ri].cells[ci].text = val

    doc.add_paragraph()

    add_para(
        doc,
        "Điều hướng Admin (scheduling): sidebar có nhóm SCHEDULING (chỉ là nhãn, không click). "
        "Bấm Work shifts hoặc Generate slots; Create shift là nút trên trang Work shifts, không có trong menu.",
        size=10,
    )
    doc.add_paragraph()

    add_heading(doc, "4. Thứ tự test khuyến nghị", 2)
    add_numbered(
        doc,
        [
            "Admin: Create Work Shift",
            "Admin: Work Shifts List (filter)",
            "Admin: Generate Appointment Slots",
            "Doctor: Work Shifts List (read-only)",
            "Doctor: Schedule Calendar",
            "Doctor: Block / Unlock slot",
            "Admin: Update Work Shift",
            "Admin: Delete Work Shift",
            "Patient: PayOS top-up + cancel pending",
            "Patient: SePay top-up",
        ],
    )

    # --- Scheduling sections ---
    add_heading(doc, "5. Scheduling — chi tiết từng luồng", 1)

    add_heading(doc, "5.1 Create Work Shift", 2)
    add_para(doc, "Actor: Admin  |  URL: /admin/work-shifts/new", bold=True)
    add_step_table(
        doc,
        [
            (1, "Login admin@orcaxcare.com / Admin@123", "Vào admin portal"),
            (2, "Sidebar → Work shifts → Create shift", "Form hiện đủ field"),
            (3, "Doctor: Dr. Nguyen Van An; Room: chọn phòng; Day: Monday", "Preview plan bên phải"),
            (4, "Start 08:00, End 12:00, Max patients 8", "Chip appointment slots hiện 8 khung giờ"),
            (5, "Create work shift", "Alert xanh — created shift Monday"),
            (6, "Tạo ca overlap 09:00–11:00 cùng doctor", "Lỗi 409 overlap"),
        ],
    )

    add_heading(doc, "5.2 Work Shifts List", 2)
    add_para(doc, "Actor: Admin  |  URL: /admin/work-shifts", bold=True)
    add_step_table(
        doc,
        [
            (1, "Mở Work shifts", "Lưới 7 cột Sun–Sat"),
            (2, "Search: Nguyen → Search", "Chỉ ca Dr. An"),
            (3, "Clear filter", "Hiện lại đủ"),
            (4, "Logout → login doctor.an@orcaxcare.com", "Doctor portal"),
            (5, "Doctor → Work shifts", "Chỉ thấy ca của mình, không Create/Edit"),
        ],
    )

    add_heading(doc, "5.3 Generate Appointment Slots", 2)
    add_para(doc, "Actor: Admin  |  URL: /admin/appointment-slots/generate", bold=True)
    add_step_table(
        doc,
        [
            (1, "Preset 14 days; One doctor → Nguyen Van An", "Meta hiện range"),
            (2, "Preview generation", "Created / Skipped preview"),
            (3, "Generate slots", "Created > 0"),
            (4, "Generate lại cùng range", "Created = 0, Skipped tăng"),
        ],
    )

    add_heading(doc, "5.4 Doctor Schedule Calendar", 2)
    add_para(doc, "Actor: Doctor  |  URL: /doctor/schedule", bold=True)
    add_step_table(
        doc,
        [
            (1, "Sau khi admin generate slots", "—"),
            (2, "My calendar → tuần có Thứ 2", "Slot Available (trắng)"),
            (3, "Toggle Week / Day; Today / mũi tên", "Điều hướng OK"),
            (4, "Click slot → panel chi tiết", "Giờ, phòng, status"),
        ],
    )

    add_heading(doc, "5.5 Block / Unlock Timeslot", 2)
    add_para(doc, "Actor: Doctor  |  URL: /doctor/schedule", bold=True)
    add_step_table(
        doc,
        [
            (1, "Click slot Available → Block slot", "Status Blocked, sọc xám"),
            (2, "Click lại → Unblock slot", "Status Available"),
            (3, "Slot quá khứ (nếu có)", "Read-only, không block"),
        ],
    )

    add_heading(doc, "5.6 Update Work Shift", 2)
    add_para(doc, "Actor: Admin  |  URL: /admin/work-shifts/:id/edit", bold=True)
    add_step_table(
        doc,
        [
            (1, "Edit ca Monday 08:00–12:00", "Form + preview bên phải"),
            (2, "Đổi End → 11:00, Max → 6", "Preview cập nhật chip slots"),
            (3, "Nếu có booked tương lai: tick Regenerate future slots", "Save OK"),
            (4, "Không tick regenerate khi đổi giờ + có booked", "Lỗi 409"),
            (5, "Save changes", "Về list, ca cập nhật"),
        ],
    )

    add_heading(doc, "5.7 Delete Work Shift", 2)
    add_para(doc, "Actor: Admin  |  URL: Edit → Delete shift", bold=True)
    add_para(
        doc,
        "Logic thực tế: luôn xóa được template. Slot available/blocked tương lai bị xóa; slot booked được giữ.",
        size=10,
    )
    add_step_table(
        doc,
        [
            (1, "Delete shift → modal OrcaXCare (không popup Windows)", "Open slots removed / Booked kept"),
            (2, "Confirm Delete shift", "Về list, ca biến mất"),
            (3, "Doctor calendar", "Slot booked vẫn còn nếu có"),
        ],
    )

    # --- Payment ---
    add_heading(doc, "6. Wallet & Payment — chi tiết từng luồng", 1)

    add_heading(doc, "6.1 PayOS Payment (REAL)", 2)
    add_para(doc, "Actor: Patient  |  URL: /patient/wallet", bold=True)
    add_step_table(
        doc,
        [
            (1, "Login patient@orcaxcare.com / Patient@123", "Patient portal"),
            (2, "Wallet → PayOS, amount 100000 → Continue", "Checkout /checkout/payos/{orderCode}"),
            (3, "Quét QR hoặc Open PayOS page", "Polling → success"),
            (4, "Wallet", "Balance +100.000 ₫, receipt success"),
            (5, "Top-up mới → Back → Cancel pending", "Status cancelled, balance không đổi"),
        ],
    )

    add_heading(doc, "6.2 SePay Payment (REAL) — thay Momo", 2)
    add_para(doc, "Actor: Patient  |  URL: /patient/wallet", bold=True)
    add_step_table(
        doc,
        [
            (1, "Wallet → SePay, amount 200000 → Continue", "Checkout SePay QR"),
            (2, "Chuyển khoản đúng memo sandbox SePay", "IPN/polling → success"),
            (3, "Wallet", "Balance tăng, transaction success"),
        ],
    )

    add_heading(doc, "7. Checklist demo (tick Pass)", 2)
    checklist = [
        "Create shift + overlap fail",
        "Work shifts list filter",
        "Generate slots + idempotent",
        "Doctor calendar slots",
        "Block / Unblock",
        "Update shift + regenerate",
        "Delete shift modal + booked preserved",
        "PayOS top-up REAL",
        "SePay top-up REAL",
    ]
    cl = doc.add_table(rows=1 + len(checklist), cols=3)
    cl.style = "Table Grid"
    cl.rows[0].cells[0].text = "#"
    cl.rows[0].cells[1].text = "Hạng mục"
    cl.rows[0].cells[2].text = "Pass?"
    for c in cl.rows[0].cells:
        set_cell_shading(c, "F1F5F9")
    for i, item in enumerate(checklist, start=1):
        cl.rows[i].cells[0].text = str(i)
        cl.rows[i].cells[1].text = item
        cl.rows[i].cells[2].text = "☐"

    doc.add_paragraph()
    add_para(doc, "OrcaXCare — Test Flow Guide (generated from codebase).", size=9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
