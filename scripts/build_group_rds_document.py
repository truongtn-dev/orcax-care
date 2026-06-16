"""
Build OrcaXCare Group RDS document (WDP301-SE1816-GROUP4_RDSDocument.docx).
Uses final draw.io use case PNGs + Iteration docx content as template.
"""
from __future__ import annotations

import shutil
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BASE = Path(__file__).resolve().parent
ASSETS = Path(r"C:/Users/nguye/.cursor/projects/e-SU26-WDP301-orcax-care/assets")
OUT_PATH = Path(r"e:/SU26/WDP301/WDP301-SE1816-GROUP4_RDSDocument.docx")
BACKUP_PATH = OUT_PATH.with_suffix(".docx.bak")
IMAGES_DIR = BASE / "rds_images"
DOWNLOADS = Path(r"c:/Users/nguye/Downloads")

USE_CASE_IMAGES = {
    "Guest & Patient (Auth)": "wdp301_g4-UseCase_Guest.drawio-4af17c27-034b-4f78-ad65-2b111c022eb8.png",
    "Patient (Features)": "wdp301_g4-UseCase_Patient.drawio-119cc865-c442-4a68-8d44-932bde0c9b41.png",
    "Doctor": "wdp301_g4-UseCase_Doctor.drawio-795d2ee7-5e11-4933-8375-b677a4c18352.png",
    "Admin": "wdp301_g4-UseCase_Admin.drawio-2c2f213b-d92b-4698-b07b-00b78a7e41be.png",
    "Staff": "wdp301_g4-UseCase_Staff.drawio-e89d3b32-cc78-4fe7-8049-88530d838d7d.png",
}

ACTORS = [
    ("Guest", "Unregistered visitor who browses the public portal, searches doctors, and registers a patient account."),
    ("Patient", "Registered user who inherits Guest capabilities and books appointments, manages wallet, insurance, and complaints."),
    ("Doctor", "Clinical user who inherits Patient portal access and performs consultations, EMR, prescriptions, and queue sessions."),
    ("Staff", "Operational user (reception/pharmacy) who manages pharmacy inventory, queue check-in, and complaint responses."),
    ("Admin", "System administrator who manages accounts, master data, doctors, and patients."),
    ("PayOS", "Secondary actor — payment gateway for appointment booking and wallet top-up."),
    ("Momo", "Secondary actor — alternative payment gateway for appointment booking."),
]

# Master use case table — continuous UC-1 .. UC-35 (+ sub-cases)
USE_CASES: list[tuple[str, str, str, str]] = [
    # Guest / Auth (Diagram 1)
    ("UC-1", "Portal", "View Portal Home", "Guest wants to open the public home page and access hospital services and CTAs."),
    ("UC-1.1", "Portal", "Search Doctors", "Guest wants to find doctors by name, specialty, or department."),
    ("UC-1.1.1", "Portal", "View Doctor Public Profile", "Guest wants to read a doctor public profile before booking."),
    ("UC-1.1.2", "Portal", "View Doctor Reviews", "Guest wants to read patient reviews on a doctor profile."),
    ("UC-1.2", "Portal", "View Featured Doctors", "Guest wants to see highlighted doctors on the home page."),
    ("UC-1.3", "Portal", "Locate Branch", "Guest wants to find clinic branch locations on a map or list."),
    ("UC-2", "Authentication", "Register Patient Account", "Guest wants to create a patient account to use the portal."),
    ("UC-2.1", "Authentication", "Verify Email Address", "Guest must verify email ownership to activate the account."),
    ("UC-2.2", "Authentication", "Resend Verification Email", "Guest wants a new verification link if the previous one expired."),
    ("UC-3", "Patient Portal", "View Patient Dashboard", "Patient wants a home screen after login with shortcuts to main tasks."),
    ("UC-3.1", "Patient Portal", "Profile Management", "Patient wants to manage personal profile information from the dashboard."),
    ("UC-3.1.1", "Patient Portal", "View Own Profile", "Patient wants to view personal and medical contact information."),
    ("UC-3.1.2", "Patient Portal", "Update Own Profile", "Patient wants to update allowed profile fields."),
    ("UC-3.1.3", "Patient Portal", "Update Profile Avatar", "Patient wants to upload or change profile photo."),
    ("UC-3.1.4", "Patient Portal", "Change Password", "Patient wants to change password while logged in."),
    ("UC-4", "Authentication", "Reset Forgotten Password", "Registered user wants to reset password when they cannot log in."),
    ("UC-5", "Authentication", "Log In", "User wants to sign in with email and password and reach the correct dashboard by role."),
    ("UC-6", "Authentication", "Log Out", "Logged-in user wants to end the session securely."),
    # Patient features (Diagram 2)
    ("UC-7", "Appointment", "Book Appointment", "Patient wants to book an appointment with a selected doctor and time slot."),
    ("UC-7.1", "Appointment", "Pay via PayOS", "Patient wants to pay appointment fee through PayOS during booking."),
    ("UC-7.2", "Appointment", "Pay via Momo", "Patient wants to pay appointment fee through Momo during booking."),
    ("UC-7.3", "Appointment", "Apply Insurance Discount", "System applies eligible insurance discount when booking."),
    ("UC-8", "Appointment", "View Appointment List", "Patient wants to view upcoming and past appointments."),
    ("UC-8.1", "Appointment", "View Appointment Detail", "Patient wants to view full details of one appointment."),
    ("UC-8.1.1", "Appointment", "Reschedule Appointment", "Patient wants to change appointment date or time when allowed."),
    ("UC-8.1.2", "Appointment", "Cancel Appointment", "Patient wants to cancel an appointment according to policy."),
    ("UC-8.1.3", "Appointment", "Rate Doctor", "Patient wants to rate the doctor after a completed visit."),
    ("UC-9", "Prescription", "View Prescription History", "Patient wants to view past prescriptions issued after visits."),
    ("UC-9.1", "Prescription", "View Prescription Detail", "Patient wants to view medication lines and instructions for one prescription."),
    ("UC-9.1.1", "Prescription", "View Prescription QR Code", "Patient wants to display QR code for pharmacy pickup."),
    ("UC-9.1.2", "Prescription", "Export Prescription PDF", "Patient wants to download prescription as PDF."),
    ("UC-10", "Notification", "View Notification Inbox", "Patient wants to read system notifications in one inbox."),
    ("UC-10.1", "Notification", "Configure Notification Preferences", "Patient wants to choose which notification channels to receive."),
    ("UC-10.2", "Notification", "Mark Notification as Read", "Patient wants to mark notifications as read."),
    ("UC-11", "Complaint", "Submit Complaint", "Patient wants to submit a service complaint to the clinic."),
    ("UC-11.1", "Complaint", "View Own Complaint List", "Patient wants to track submitted complaints."),
    ("UC-11.1.1", "Complaint", "View Own Complaint Detail", "Patient wants to read staff replies and complaint status."),
    ("UC-12", "Wallet", "View Wallet Balance", "Patient wants to view medical wallet balance."),
    ("UC-12.1", "Wallet", "View Wallet Transaction History", "Patient wants to review wallet top-up and payment history."),
    ("UC-13", "Insurance", "View Insurance Card List", "Patient wants to manage registered insurance cards."),
    ("UC-13.1", "Insurance", "Delete Insurance Card", "Patient wants to remove an insurance card from the profile."),
    ("UC-13.2", "Insurance", "Add Insurance Card", "Patient wants to add a new insurance card."),
    ("UC-13.2.1", "Insurance", "Scan Insurance Card (OCR)", "Patient wants to scan insurance card image to auto-fill fields."),
    ("UC-13.3", "Insurance", "Update Insurance Card", "Patient wants to edit insurance card information."),
    ("UC-14", "Queue", "View Queue Status", "Patient wants to see current queue number and waiting status."),
    ("UC-15", "Favorites", "View Favorite Doctors", "Patient wants to view saved favorite doctors."),
    ("UC-15.1", "Favorites", "Add Doctor to Favorites", "Patient wants to bookmark a doctor for quick access."),
    ("UC-15.2", "Favorites", "Remove Doctor from Favorites", "Patient wants to remove a doctor from favorites."),
    # Doctor (Diagram 3)
    ("UC-16", "Doctor Console", "View Doctor Dashboard", "Doctor wants an overview screen after login with today work shortcuts."),
    ("UC-16.1", "Doctor Console", "Profile Management", "Doctor wants to manage own professional profile."),
    ("UC-16.1.1", "Doctor Console", "View Own Profile", "Doctor wants to view linked doctor profile and account info."),
    ("UC-16.1.2", "Doctor Console", "Update Own Profile", "Doctor wants to update allowed profile fields such as bio and contact."),
    ("UC-16.1.3", "Doctor Console", "Change Password", "Doctor wants to change password from profile or settings."),
    ("UC-17", "Consultation", "View Today Appointments", "Doctor wants to see appointments scheduled for today."),
    ("UC-17.1", "Consultation", "Start Consultation", "Doctor wants to begin a clinical session for a checked-in patient."),
    ("UC-17.1.1", "Consultation", "Create Encounter Record", "Doctor must create an encounter record when starting consultation."),
    ("UC-17.1.1.1", "Consultation", "View Encounter Detail", "Doctor wants to view encounter notes and linked data."),
    ("UC-17.1.1.1.1", "Consultation", "Update Encounter Notes", "Doctor wants to edit clinical notes during or after visit."),
    ("UC-17.1.1.1.2", "Consultation", "Sign Off Encounter", "Doctor wants to finalize and sign off the encounter."),
    ("UC-17.1.1.2", "Consultation", "Record Diagnosis", "Doctor wants to record ICD-10 diagnosis for the encounter."),
    ("UC-17.1.1.2.1", "Consultation", "Remove Diagnosis", "Doctor wants to remove an incorrect diagnosis entry."),
    ("UC-17.1.1.2.2", "Consultation", "Search ICD-10 Catalog", "Doctor wants to search diagnosis codes from ICD-10 catalog."),
    ("UC-17.1.1.2.3", "Consultation", "Update Diagnosis", "Doctor wants to update an existing diagnosis entry."),
    ("UC-18", "Queue Session", "Open Queue Session", "Doctor wants to open a queue session for their clinic room."),
    ("UC-18.1", "Queue Session", "Call Next Patient", "Doctor calls the next patient ticket in the active queue."),
    ("UC-18.2", "Queue Session", "Recall Queue Ticket", "Doctor wants to recall a previously called ticket."),
    ("UC-18.3", "Queue Session", "Close Queue Session", "Doctor wants to close the active queue session."),
    ("UC-19", "Scheduling", "View Work Shift Schedule", "Doctor wants to view assigned work shifts."),
    ("UC-20", "Scheduling", "View Appointment Calendar", "Doctor wants to view personal appointment calendar."),
    ("UC-20.1", "Scheduling", "Block Appointment Slot", "Doctor wants to block unavailable time slots."),
    ("UC-20.2", "Scheduling", "Unblock Appointment Slot", "Doctor wants to reopen a previously blocked slot."),
    ("UC-21", "EMR", "View Patient EMR Timeline", "Doctor wants to review patient medical history timeline."),
    ("UC-23", "Medical Imaging", "Upload Medical Image", "Doctor wants to upload diagnostic images for a patient."),
    ("UC-23.1", "Medical Imaging", "View Medical Imaging Gallery", "Doctor wants to browse uploaded images for a patient."),
    ("UC-23.1.1", "Medical Imaging", "View Image Fullscreen", "Doctor wants to view an image in fullscreen mode."),
    ("UC-23.1.2", "Medical Imaging", "Delete Medical Image", "Doctor wants to delete an uploaded image when allowed."),
    ("UC-24", "Prescription", "Create Prescription", "Doctor wants to create a prescription for the current encounter."),
    ("UC-24.1", "Prescription", "Add Prescription Line Item", "Doctor wants to add a medicine line to the prescription."),
    ("UC-24.2", "Prescription", "Update Prescription Line Item", "Doctor wants to edit dosage or instructions on a line."),
    ("UC-24.3", "Prescription", "Remove Prescription Line Item", "Doctor wants to remove a line from draft prescription."),
    ("UC-24.4", "Prescription", "Save Prescription", "Doctor wants to save and issue the prescription."),
    ("UC-24.5", "Prescription", "View Prescription Detail", "Doctor wants to review saved prescription details."),
    ("UC-24.5.1", "Prescription", "Export Prescription PDF", "Doctor wants to export prescription as PDF."),
    ("UC-24.5.2", "Prescription", "View Prescription QR Code", "Doctor wants to display prescription QR code."),
    # Admin (Diagram 4) — high-level modules with note sub-features
    ("UC-25", "Administration", "Manage Accounts", "Admin wants to create, list, update, change role, deactivate or restore user accounts."),
    ("UC-26", "Administration", "Manage Specialties", "Admin wants to maintain medical specialty master data (CRUD)."),
    ("UC-27", "Administration", "Manage Departments", "Admin wants to maintain department master data including deactivate."),
    ("UC-28", "Administration", "Manage Clinic Rooms", "Admin wants to register and update clinic rooms by department."),
    ("UC-29", "Administration", "Manage Doctors", "Admin wants to manage doctor profiles including import/export Excel."),
    ("UC-30", "Administration", "Manage Patients", "Admin wants to manage patient records and deactivate accounts."),
    ("UC-31", "Administration", "View Admin Dashboard", "Admin wants a dashboard with KPIs and links to management modules."),
    # Staff (Diagram 5)
    ("UC-32", "Staff Operations", "Manage Pharmacy", "Staff wants to register medicines, track stock in/out, and view low-stock alerts."),
    ("UC-33", "Staff Operations", "Manage Queue Check-in", "Staff wants to issue queue tickets and tie them to appointments."),
    ("UC-34", "Staff Operations", "Manage Complaints", "Staff wants to view, reply to, and update status of patient complaints."),
    ("UC-35", "Staff Operations", "View Staff Dashboard", "Staff wants an operational dashboard after login."),
]

SCREEN_AUTH = [
    ("Portal Home / Search Doctors", "X", "X", "X", "X", "X"),
    ("Register / Verify Email", "X", "", "", "", ""),
    ("Log In / Log Out / Forgot Password", "X", "X", "X", "X", "X"),
    ("Patient Dashboard & Booking", "", "X", "X", "", ""),
    ("Wallet / Insurance / Queue Status", "", "X", "", "", ""),
    ("Doctor Console / EMR / Prescription", "", "", "X", "", ""),
    ("Staff Pharmacy / Queue Check-in", "", "", "", "X", ""),
    ("Admin Master Data & Accounts", "", "", "", "", "X"),
]

# Primary UC specs for Section II (summary level)
UC_SPECS = [
    ("UC-1", "View Portal Home", "Guest", [
        "Display hero, services, featured doctors, news, and login/register CTAs.",
        "Support navigation to search doctors, doctor profiles, and branch locator.",
    ], [
        "Public content is read-only for Guest.",
        "Inactive or hidden doctors must not appear in featured list.",
    ]),
    ("UC-2", "Register Patient Account", "Guest", [
        "Collect email, password, full name, phone; validate uniqueness.",
        "Send verification email; account inactive until UC-2.1 completes.",
    ], [
        "Duplicate email returns conflict.",
        "Password must meet minimum security rules.",
    ]),
    ("UC-5", "Log In", "All roles", [
        "Authenticate with email/password; issue session token.",
        "Redirect to role dashboard: Patient, Doctor, Staff, or Admin.",
    ], [
        "Deactivated accounts cannot log in.",
        "Patient must have verified email before login.",
    ]),
    ("UC-7", "Book Appointment", "Patient", [
        "Select doctor, date, time slot; confirm booking.",
        "Optional payment via PayOS/Momo; apply insurance discount when eligible.",
    ], [
        "Cannot book inactive doctors or blocked slots.",
        "Wallet balance or external payment must succeed before confirmation.",
    ]),
    ("UC-17.1", "Start Consultation", "Doctor", [
        "Open consultation from today appointment; create encounter record.",
        "Support diagnosis, notes, imaging, and prescription workflows.",
    ], [
        "Only assigned doctor can start consultation for the appointment.",
        "Encounter must be signed off before prescription is finalized.",
    ]),
    ("UC-25", "Manage Accounts", "Admin", [
        "List/search accounts; create staff accounts; update profile fields.",
        "Change role; deactivate or restore accounts with confirmation.",
    ], [
        "Admin cannot deactivate own account.",
        "Role change requires authorization audit.",
    ]),
    ("UC-32", "Manage Pharmacy", "Staff", [
        "Register medicine SKU; view inventory and detail.",
        "Record stock inbound/outbound; highlight low-stock alerts.",
    ], [
        "Stock outbound cannot exceed available quantity.",
        "Medicine code must be unique within pharmacy catalog.",
    ]),
]


def find_asset(filename: str) -> Path:
    exact = ASSETS / filename
    if exact.exists():
        return exact
    matches = list(ASSETS.glob(f"*{Path(filename).name}"))
    if matches:
        return matches[0]
    raise FileNotFoundError(f"Use case image not found: {filename}")


def prepare_images() -> dict[str, Path]:
    IMAGES_DIR.mkdir(exist_ok=True)
    result = {}
    for label, fname in USE_CASE_IMAGES.items():
        src = find_asset(fname)
        slug = label.lower().replace(" ", "_").replace("(", "").replace(")", "").replace("&", "and")
        dst = IMAGES_DIR / f"uc_{slug}.png"
        shutil.copy2(src, dst)
        result[label] = dst
    return result


def extract_iteration_image(docx_path: Path, media_name: str, out_name: str) -> Path | None:
    """Extract one image from iteration docx media folder."""
    try:
        with zipfile.ZipFile(docx_path) as z:
            key = f"word/media/{media_name}"
            if key not in z.namelist():
                return None
            out = IMAGES_DIR / out_name
            out.write_bytes(z.read(key))
            return out
    except OSError:
        return None


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths=None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri].cells[ci], val)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_image(doc: Document, path: Path, caption: str, width=Inches(6.2)) -> None:
    if path and path.exists():
        doc.add_picture(str(path), width=width)
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.style = "Caption"
    except KeyError:
        pass


def add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph("OrcaXCare Project Report")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(22)
    for line in [
        "Requirements & Design Specification (RDS)",
        "Subject: WDP301",
        "Version: 1.0",
        "",
        "Group 4 — SE1816",
        "OrcaXCare – Healthcare Management System & Patient Portal",
        "",
        f"– Can Tho, {date.today().strftime('%B %Y')} –",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_record_of_changes(doc: Document) -> None:
    add_heading(doc, "Record of Changes", 1)
    add_table(
        doc,
        ["Version", "Date", "A/M/D", "In charge", "Change Description"],
        [
            ["V1.0", date.today().strftime("%d/%m/%Y"), "A", "Group 4", "Initial OrcaXCare RDS — use cases UC-1 to UC-35, actors, diagrams, screen authorization."],
        ],
    )
    doc.add_page_break()


def add_overview(doc: Document, uc_images: dict[str, Path]) -> None:
    add_heading(doc, "I. Overview", 1)
    add_heading(doc, "1. User Requirements", 2)
    add_heading(doc, "1.1 Actors", 3)
    doc.add_paragraph(
        "Primary actors interact directly with OrcaXCare. Patient inherits Guest; Doctor inherits Patient. "
        "Staff and Admin are independent operational roles. PayOS and Momo are secondary payment actors."
    )
    add_table(doc, ["#", "Actor", "Description"], [[str(i + 1), a, d] for i, (a, d) in enumerate(ACTORS)])

    add_heading(doc, "1.2 Use Cases", 2)
    add_heading(doc, "a. Diagram(s)", 4)
    fig = 1
    for label, path in uc_images.items():
        add_image(doc, path, f"Figure {fig}: Use case diagram — {label}")
        fig += 1

    add_heading(doc, "b. Descriptions", 4)
    doc.add_paragraph(
        "The table below lists all use cases in continuous numbering (UC-1 to UC-35). "
        "Sub-cases use hierarchical IDs aligned with the diagrams above."
    )
    add_table(
        doc,
        ["ID", "Feature", "Use Case", "Use Case Description"],
        [[a, b, c, d] for a, b, c, d in USE_CASES],
    )

    add_heading(doc, "2. Overall Functionalities", 2)
    add_heading(doc, "2.1 Screens Flow", 3)
    sf = extract_iteration_image(DOWNLOADS / "Iteration 1 (3).docx", "image2.png", "screen_flow_iter1.png")
    if sf:
        add_image(doc, sf, "Figure: Screen flow — Iteration 1 (authentication & portal entry)")
    else:
        doc.add_paragraph("See Iteration 1 screen flow diagrams for Login, Register, Verify Email, and Portal Home.")

    add_heading(doc, "2.2 Screen Descriptions", 3)
    doc.add_paragraph(
        "Screen layouts and wireframes are documented in Iteration 1 deliverables: Portal Home, Patient Dashboard, "
        "Admin account/specialty/department/doctor/patient management screens, and doctor public profile."
    )

    add_heading(doc, "2.3 Screen Authorization", 3)
    add_table(
        doc,
        ["Screen/Function", "Guest", "Patient", "Doctor", "Staff", "Admin"],
        SCREEN_AUTH,
    )

    add_heading(doc, "2.4 Non-UI Functions", 3)
    doc.add_paragraph(
        "Scheduled appointment reminders, booking confirmation stubs, and email verification are implemented as "
        "system processes noted on related use cases (not separate actors)."
    )

    add_heading(doc, "3. System High Level Design", 2)
    add_heading(doc, "3.1 ERD", 3)
    erd = extract_iteration_image(DOWNLOADS / "Iteration 1 (3).docx", "image12.png", "erd_schema.png")
    if erd:
        add_image(doc, erd, "Figure: Entity Relationship Diagram — OrcaXCare Iteration 1")

    add_heading(doc, "3.2 Database Design", 3)
    add_heading(doc, "a. Database Schema", 4)
    doc.add_paragraph(
        "MongoDB collections: users, auth_tokens, patients, doctors, specialties, departments, clinic_rooms, "
        "appointments, encounters, prescriptions, wallets, insurance_cards, complaints, queue_tickets, medicines, notifications."
    )

    add_heading(doc, "3.3 Code Packages", 3)
    doc.add_paragraph(
        "MERN stack: client (React pages/components/services), server (routes/controllers/services/models), shared validation utilities."
    )

    doc.add_page_break()


def add_requirement_specs(doc: Document) -> None:
    add_heading(doc, "II. Requirement Specifications", 1)
    doc.add_paragraph(
        "Detailed specifications for primary use cases. Sub-flows (UC-x.y) inherit business rules from parent use cases."
    )
    for idx, (uc_id, name, actor, funcs, rules) in enumerate(UC_SPECS, start=1):
        add_heading(doc, f"{idx}. {name}", 2)
        add_heading(doc, f"{idx}.1 {uc_id} — {name}", 3)
        add_heading(doc, "a. Functionalities", 4)
        for f in funcs:
            doc.add_paragraph(f, style="List Bullet")
        add_heading(doc, "b. Business Rules", 4)
        for r in rules:
            doc.add_paragraph(r, style="List Bullet")
        doc.add_paragraph(f"Primary actor(s): {actor}.")
    doc.add_page_break()


def add_screen_layouts_from_iteration(doc: Document) -> None:
    """Append key screen layout sections from Iteration 1 (3)."""
    add_heading(doc, "III. Screen Layouts (Iteration 1)", 1)
    src = Document(DOWNLOADS / "Iteration 1 (3).docx")
    from docx.oxml.ns import qn

    rels = src.part.rels
    screen_sections = [
        (18, "1.1 View Accounts List"),
        (24, "1.2 View Account Detail"),
        (27, "1.3 Create Specialty"),
        (29, "1.4 Create Clinic Room"),
        (52, "1.5 View Clinic Rooms List"),
        (55, "1.6 View Doctor Public Profile"),
        (57, "1.7 Featured Doctors"),
        (60, "1.8 Upload Profile Avatar"),
        (63, "1.9 View Patient Detail (Admin)"),
    ]
    for para_idx, title in screen_sections:
        add_heading(doc, title, 3)
        if para_idx < len(src.paragraphs):
            desc_para = para_idx - 1
            if desc_para >= 0 and src.paragraphs[desc_para].text.strip():
                doc.add_paragraph(src.paragraphs[desc_para].text.strip())
        para = src.paragraphs[para_idx]
        for run in para.runs:
            blips = run._element.findall(".//" + qn("a:blip"))
            for blip in blips:
                embed = blip.get(qn("r:embed"))
                if embed:
                    rel = rels.get(embed)
                    if rel:
                        img = extract_iteration_image(
                            DOWNLOADS / "Iteration 1 (3).docx",
                            Path(rel.target_ref).name,
                            f"screen_{para_idx}_{Path(rel.target_ref).name}",
                        )
                        if img:
                            add_image(doc, img, f"Screen: {title}", width=Inches(5.8))
    doc.add_page_break()


def add_rds_sds_section(doc: Document) -> None:
    add_heading(doc, "IV. SDS — Sequence Diagrams (Iteration 1)", 1)
    seq_images = [
        ("View Account List", "image5.png"),
        ("View Account Detail", "image15.png"),
        ("Specialties CRUD", "image4.png"),
        ("View Doctor Public Profile", "image18.png"),
        ("Featured Doctors", "image10.png"),
    ]
    for title, media in seq_images:
        add_heading(doc, title, 3)
        img = extract_iteration_image(DOWNLOADS / "Iteration 1 (3).docx", media, f"seq_{media}")
        if img:
            add_image(doc, img, f"Sequence diagram: {title}", width=Inches(5.8))
    doc.add_page_break()


def copy_table_to_doc(src_table, doc: Document) -> None:
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for ri in range(rows):
        for ci in range(cols):
            table.rows[ri].cells[ci].text = src_table.rows[ri].cells[ci].text.strip()
    doc.add_paragraph()


def add_iteration1_detail_tables(doc: Document) -> None:
    """Merge step-level UC tables from Iteration 1 (1).docx."""
    iter1_path = DOWNLOADS / "Iteration 1 (1).docx"
    if not iter1_path.exists():
        return
    src = Document(iter1_path)
    add_heading(doc, "V. Detailed Use Case Flows (Iteration 1 — team deliverables)", 1)
    doc.add_paragraph(
        "Step-level flows for Iteration 1 implemented features. "
        "Legacy IDs UC-01..UC-11 map to master list Section I "
        "(e.g. UC-01 Patient Dashboard → UC-3; UC-02 View Doctors List → UC-1.1)."
    )
    sections = [
        (2, "a. Patient Dashboard (maps to UC-3)"),
        (3, "b. View Doctors List — Patient (maps to UC-1.1)"),
        (4, "c. Update Own Profile (maps to UC-3.1.2)"),
        (5, "d. Update Account — Admin (maps to UC-25)"),
        (6, "e. View Specialties List (maps to UC-26)"),
        (7, "f. Create Department (maps to UC-27)"),
        (8, "g. View Department Detail (maps to UC-27)"),
        (9, "h. Update Doctor (maps to UC-29)"),
        (10, "i. Import Doctors from Excel (maps to UC-29)"),
        (11, "j. Export Doctors to Excel (maps to UC-29)"),
        (12, "k. Update Patient Profile — Admin (maps to UC-30)"),
    ]
    for table_idx, title in sections:
        if table_idx >= len(src.tables):
            continue
        add_heading(doc, title, 3)
        copy_table_to_doc(src.tables[table_idx], doc)
    doc.add_page_break()


def add_iteration1_extra_screens(doc: Document) -> None:
    """Patient/Admin screen layouts from Iteration 1 (1).docx."""
    from docx.oxml.ns import qn

    iter1_path = DOWNLOADS / "Iteration 1 (1).docx"
    if not iter1_path.exists():
        return
    src = Document(iter1_path)
    rels = src.part.rels
    add_heading(doc, "VI. Additional Screen Layouts — Patient & Admin", 1)
    layout_items = [
        (38, "1.1 Patient Dashboard"),
        (41, "1.2 View Doctors List (Patient)"),
        (45, "1.3 Update Own Profile"),
        (51, "2.1 Update Account"),
        (55, "2.2 View Specialties List"),
    ]
    for para_idx, title in layout_items:
        add_heading(doc, title, 3)
        if para_idx - 1 < len(src.paragraphs):
            prev = src.paragraphs[para_idx - 1].text.strip()
            if prev:
                doc.add_paragraph(prev)
        if para_idx >= len(src.paragraphs):
            continue
        para = src.paragraphs[para_idx]
        for run in para.runs:
            for blip in run._element.findall(".//" + qn("a:blip")):
                embed = blip.get(qn("r:embed"))
                if not embed:
                    continue
                rel = rels.get(embed)
                if not rel:
                    continue
                name = Path(rel.target_ref).name
                img = extract_iteration_image(iter1_path, name, f"iter1_{name}")
                if img:
                    add_image(doc, img, f"Screen layout: {title}", width=Inches(5.5))


def build_document() -> Path:
    IMAGES_DIR.mkdir(exist_ok=True)
    uc_images = prepare_images()

    doc = Document()
    add_title_page(doc)
    add_record_of_changes(doc)
    add_overview(doc, uc_images)
    add_requirement_specs(doc)
    add_screen_layouts_from_iteration(doc)
    add_rds_sds_section(doc)
    add_iteration1_detail_tables(doc)
    add_iteration1_extra_screens(doc)

    if OUT_PATH.exists():
        shutil.copy2(OUT_PATH, BACKUP_PATH)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    out = build_document()
    print(f"Saved: {out}")
    if BACKUP_PATH.exists():
        print(f"Backup: {BACKUP_PATH}")
