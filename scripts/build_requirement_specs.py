"""
Build a standalone Word file with only Section II. Requirement Specifications
for OrcaXCare, following the FBUS group-document format:
  Heading 1: II. Requirement Specifications
  Heading 2: <n>. <Feature>
  Heading 3: <n>.<m> UC-x_<Use Case Name>
  Heading 4: a. Functionalities  -> 4-col spec table
  Heading 4: b. Business Rules   -> 3-col BR table
Created By values come from the backlog "In Charge" column.
Output: e:/SU26/WDP301/WDP301-SE1816-GROUP4_RequirementSpecs.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUT_PATH = Path(r"e:/SU26/WDP301/WDP301-SE1816-GROUP4_RequirementSpecs.docx")
DATE = "10/06/2026"

# ---------------------------------------------------------------- BR catalog
BR: dict[str, tuple[str, str]] = {
    "BR-01": ("Action Enabler", "Only registered and logged-in Patients can book appointments or access personal data. Guests may browse public content only."),
    "BR-02": ("Fact", "Only active (non-deactivated) doctors appear on public pages: featured section, search results, and public profiles."),
    "BR-03": ("Constraint", "Email must be unique across all accounts and is stored in lowercase."),
    "BR-04": ("Constraint", "Password must be at least 8 characters and include letters and numbers. Passwords are hashed (bcrypt) before storing."),
    "BR-05": ("Action Enabler", "A Patient account must verify its email address before it can log in."),
    "BR-06": ("Constraint", "Deactivated or locked accounts cannot log in; existing sessions of a deactivated account are revoked."),
    "BR-07": ("Fact", "All timestamps must use Asia/Ho_Chi_Minh timezone."),
    "BR-08": ("Fact", "Each role is redirected to and can only access screens mapped to its permissions (see Screen Authorization)."),
    "BR-09": ("Fact", "The system never reveals whether an email exists during password recovery (neutral response)."),
    "BR-10": ("Constraint", "After logout, the revoked token must be rejected by all protected APIs; browser back navigation must not expose protected screens."),
    "BR-11": ("Constraint", "A user can view and update only their own profile; email and role cannot be self-changed."),
    "BR-12": ("Constraint", "Avatar upload accepts image files (JPG/PNG) with size limit 5 MB."),
    "BR-13": ("Action Enabler", "An appointment is confirmed only after the consultation fee is paid successfully via Wallet, PayOS, or Momo."),
    "BR-14": ("Constraint", "A time slot is held for a maximum of 10 minutes during checkout; the hold is released if payment is not completed."),
    "BR-15": ("Constraint", "Patients may cancel or reschedule an appointment at least 24 hours before its start time; refunds are credited to the wallet per policy."),
    "BR-16": ("Constraint", "Insurance discount is applied only for valid (non-expired) cards from supported providers; one card per booking."),
    "BR-17": ("Action Enabler", "A Patient can rate a doctor only after a completed appointment, with one review per appointment."),
    "BR-18": ("Constraint", "Wallet balance can never go negative; top-ups are accepted only through PayOS or Momo gateways."),
    "BR-19": ("Action Enabler", "Prescription QR code and PDF export are available only for issued (signed) prescriptions."),
    "BR-20": ("Fact", "Notifications are retained for 90 days; booking confirmation notifications are enabled by default."),
    "BR-21": ("Constraint", "A complaint can be edited by the Patient only while its status is Open; after staff reply or resolution it becomes read-only."),
    "BR-22": ("Constraint", "A Patient can view queue status only for their own active ticket."),
    "BR-23": ("Constraint", "A Doctor can view and modify EMR data only for patients having an encounter assigned to that doctor."),
    "BR-24": ("Constraint", "An encounter can be edited only before sign-off; after sign-off it becomes immutable."),
    "BR-25": ("Constraint", "Every diagnosis must reference a valid ICD-10 code selected from the catalog."),
    "BR-26": ("Constraint", "Only one active queue session can exist per clinic room at any time."),
    "BR-27": ("Constraint", "A time slot can be blocked only when it has no confirmed appointment."),
    "BR-28": ("Constraint", "A prescription line must have quantity > 0 and reference an existing medicine; stock availability is checked when the prescription is issued."),
    "BR-29": ("Constraint", "Medical images accept JPG/PNG/DICOM up to 20 MB; an image can be deleted only before the encounter is signed off."),
    "BR-30": ("Action Enabler", "Only Admin can create staff/doctor accounts and change user roles; Admin cannot deactivate their own account."),
    "BR-31": ("Constraint", "Master data (specialty, department, clinic room) referenced by doctors or appointments cannot be hard-deleted - it must be deactivated instead."),
    "BR-32": ("Constraint", "Doctor Excel import accepts the provided .xlsx template only; invalid rows are reported while valid rows are committed."),
    "BR-33": ("Fact", "Deactivated patient/doctor accounts are soft-deleted; their historical records (appointments, encounters) are retained."),
    "BR-34": ("Constraint", "Stock outbound quantity cannot exceed current stock; a low-stock alert is raised when quantity falls below the reorder level."),
    "BR-35": ("Constraint", "A queue ticket can be issued only for a checked-in appointment of the current day."),
    "BR-36": ("Fact", "Complaint replies are visible to the Patient; status transitions follow Open -> In Progress -> Resolved."),
    "BR-37": ("Fact", "Dashboards display only data scoped to the viewer's role and permissions."),
}

# ------------------------------------------------------------- spec helpers
def spec(uc_id, name, created_by, primary, secondary, trigger, description,
         pre, post, normal, alt, exc, priority, freq, brs, other, assume):
    return {
        "id": uc_id, "name": name, "by": created_by, "primary": primary,
        "secondary": secondary, "trigger": trigger, "desc": description,
        "pre": pre, "post": post, "normal": normal, "alt": alt, "exc": exc,
        "priority": priority, "freq": freq, "brs": brs, "other": other,
        "assume": assume,
    }

# ---------------------------------------------------------------- spec data
GROUPS: list[tuple[str, list[dict]]] = [
    ("View Portal Home", [spec(
        "UC-1", "View Portal Home", "TanhNT", "Guest", "None",
        "When the Guest accesses the OrcaXCare portal URL or clicks the logo from any public page.",
        "This use case allows Guests to view the public home page including hero section, services overview, featured doctors, and news. From the home page, Guest can navigate to Search Doctors (UC-1.1), Featured Doctors (UC-1.2), and Branch Locator (UC-1.3). No login is required.",
        ["PRE-1: The system is online and accessible."],
        ["POST-1: The home page is displayed with featured doctors loaded.",
         "POST-2: Guest can continue to search doctors, view profiles, or register/login."],
        ["1. Guest accesses the portal URL.",
         "2. System loads hero section, services, news, and featured doctors (active doctors only).",
         "3. Guest clicks a doctor card or \u201cSearch Doctors\u201d to explore (UC-1.1).",
         "4. Guest optionally clicks \u201cBranches\u201d to locate a clinic (UC-1.3).",
         "5. Guest clicks \u201cLogin\u201d or \u201cRegister\u201d to proceed with an account."],
        ["2a. No featured doctors configured \u2192 System hides the featured section and shows remaining content."],
        ["2b. System fails to load content \u2192 Error: \u201cUnable to load page content. Please try again later.\u201d"],
        "High", "Frequently \u2014 entry point of all Guests.",
        ["BR-01", "BR-02", "BR-07"],
        "Featured doctors limited to 8 items; news section paginated.",
        "Guest is connected to the internet and system services are running.")]),

    ("Register", [spec(
        "UC-2", "Register Patient Account", "TruongNT", "Guest", "None",
        "When the Guest clicks the \u201cRegister\u201d button in the header or on the Login screen.",
        "This use case allows Guests to create a Patient account by filling in email, password, confirm password, full name, and phone. After submitting, the system creates an inactive account and sends a verification email (UC-2.1 Verify Email Address is included). Guest may request a new link via UC-2.2 Resend Verification Email.",
        ["PRE-1: Guest is not logged in.", "PRE-2: The email is not already registered."],
        ["POST-1: A new user account (role = patient, isEmailVerified = false) and linked patient profile are created.",
         "POST-2: A verification email containing a tokenized link is sent."],
        ["1. Guest opens the Register screen.",
         "2. Guest enters email, password, confirm password, full name, phone, and accepts terms.",
         "3. Guest clicks \u201cRegister\u201d.",
         "4. System validates inputs and email uniqueness.",
         "5. System hashes the password, creates the user and patient profile, and sends a verification email.",
         "6. System shows \u201cCheck your email\u201d notice screen.",
         "7. Guest opens the link in email \u2192 UC-2.1 validates the token and activates the account.",
         "8. Guest proceeds to Login."],
        ["4a. Email already registered \u2192 System displays: \u201cEmail already registered.\u201d and stays on the form.",
         "7a. Token expired \u2192 Guest requests a new email (UC-2.2 Resend Verification Email)."],
        ["5a. Email service fails \u2192 System keeps the account inactive and shows: \u201cCould not send verification email. Please use Resend.\u201d"],
        "High", "Moderate \u2014 once per new patient.",
        ["BR-03", "BR-04", "BR-05", "BR-07"],
        "Verification token expires after 24 hours; resend is rate-limited.",
        "Guest owns a valid email address accessible for verification.")]),

    ("Login", [spec(
        "UC-5", "Login", "TruongNT", "Guest (becoming Patient), Doctor, Staff, Admin", "None",
        "When the user clicks \u201cLogin\u201d in the header or is redirected to the Login screen from a protected page.",
        "This use case allows any registered user to sign in using email and password with a remember-me option. On success, the system issues a session token and redirects the user to the dashboard matching their role (Patient / Doctor / Staff / Admin).",
        ["PRE-1: The account exists and is active.", "PRE-2: Patient accounts must have a verified email."],
        ["POST-1: A session token is issued and stored client-side.",
         "POST-2: lastLoginAt is updated; user is redirected by role."],
        ["1. User opens the Login screen.",
         "2. User enters email, password, and optionally checks \u201cRemember me\u201d.",
         "3. User clicks \u201cLogin\u201d.",
         "4. System verifies credentials against the stored password hash.",
         "5. System checks account status (active, not locked, email verified for patient).",
         "6. System issues a session token and records last login time.",
         "7. System redirects to the role dashboard."],
        ["4a. Wrong email or password \u2192 \u201cInvalid email or password.\u201d (no detail which one).",
         "5a. Email not verified \u2192 show resend verification option (UC-2.2)."],
        ["5b. Account deactivated or locked \u2192 \u201cYour account is not available. Please contact support.\u201d"],
        "High", "Frequently \u2014 every session of every role.",
        ["BR-05", "BR-06", "BR-08"],
        "Remember-me extends token expiry; failed attempts may temporarily lock the account.",
        "User remembers registered email; system services are running.")]),

    ("Forgot Password", [spec(
        "UC-4", "Reset Forgotten Password", "TruongNT", "Patient (also Doctor, Staff, Admin)", "None",
        "When the user clicks the \u201cForgot password?\u201d link on the Login screen.",
        "This use case allows a registered user who cannot log in to reset their password. The user submits their registered email, receives a secure reset link, and sets a new password.",
        ["PRE-1: The account exists in the system."],
        ["POST-1: The password is replaced with the new hashed password.",
         "POST-2: All existing session tokens are revoked."],
        ["1. User clicks \u201cForgot password?\u201d on the Login screen.",
         "2. User enters the registered email and submits.",
         "3. System shows a neutral success message regardless of whether the email exists.",
         "4. System sends a reset link (tokenized, time-limited) to the email if registered.",
         "5. User opens the link \u2192 Reset Password screen.",
         "6. User enters new password and confirmation; system validates strength.",
         "7. System updates the password hash, revokes old sessions, and redirects to Login."],
        ["5a. Token expired or already used \u2192 \u201cThis link is no longer valid.\u201d with option to request a new one."],
        ["4a. Email service fails \u2192 log error; user may retry later (message remains neutral)."],
        "High", "Occasionally.",
        ["BR-04", "BR-09"],
        "Reset token expires after 30 minutes and is single-use.",
        "User still controls the registered email inbox.")]),

    ("Logout", [spec(
        "UC-6", "Logout", "TruongNT", "Patient, Doctor, Staff, Admin", "None",
        "When the logged-in user clicks \u201cLog Out\u201d in the avatar dropdown menu.",
        "This use case allows a logged-in user to end the session securely. The session token is revoked on the server, client storage is cleared, and the user is returned to the Login screen. Back navigation cannot reopen protected pages.",
        ["PRE-1: User is authenticated with a valid session token."],
        ["POST-1: The session token is revoked server-side.",
         "POST-2: Client session data is cleared; protected routes are blocked."],
        ["1. User clicks avatar > \u201cLog Out\u201d.",
         "2. System asks for confirmation.",
         "3. User confirms.",
         "4. System revokes the session token and clears client storage.",
         "5. System redirects to Login (replace navigation)."],
        ["3a. User cancels \u2192 remains on the current screen."],
        ["4a. Token already expired \u2192 client storage is still cleared and user is redirected to Login."],
        "High", "Frequently \u2014 end of every session.",
        ["BR-10"],
        "Applies identically to all roles.",
        "None.")]),

    ("Manage Profile", [spec(
        "UC-3.1", "Profile Management", "ThangND", "Patient (Doctor inherits)", "None",
        "When the user opens \u201cMy Profile\u201d from the dashboard avatar menu (extends UC-3 View Patient Dashboard).",
        "This use case allows the Patient to view their own profile (UC-3.1.1, included), update personal information (UC-3.1.2), upload/change avatar (UC-3.1.3), and change password (UC-3.1.4). Editable fields include name, date of birth, gender, address, phone, emergency contact, and basic medical info (blood type, allergies).",
        ["PRE-1: User is logged in (UC-5).", "PRE-2: The account has a linked patient profile."],
        ["POST-1: Updated fields are persisted.",
         "POST-2: Changes are reflected immediately on the profile screen."],
        ["1. User clicks avatar > \u201cMy Profile\u201d.",
         "2. System loads the profile of the currently logged-in user only.",
         "3. User clicks \u201cEdit\u201d and modifies allowed fields.",
         "4. System validates data (phone format, date of birth in the past).",
         "5. User clicks \u201cSave\u201d; system persists changes and shows updated profile.",
         "6. Optionally user uploads a new avatar (crop + preview) or changes password."],
        ["4a. Validation fails \u2192 field-level error messages; data is not saved.",
         "3a. User clicks \u201cCancel\u201d \u2192 original values are restored."],
        ["5a. Save fails (server error) \u2192 \u201cUnable to save changes. Please try again.\u201d"],
        "Medium", "Occasionally.",
        ["BR-11", "BR-12"],
        "Email and role are read-only; avatar accepts JPG/PNG up to 5 MB.",
        "User is the owner of the profile being edited.")]),

    ("Book Appointment", [spec(
        "UC-7", "Book Appointment", "ThangND", "Patient", "PayOS, Momo",
        "When the Patient clicks \u201cBook Appointment\u201d on a doctor profile or from the dashboard.",
        "This use case allows the Patient to book an appointment by selecting doctor, date, and an available time slot, then paying the consultation fee. Payment can be made via Wallet, PayOS (UC-7.1) or Momo (UC-7.2). An eligible insurance discount is applied automatically (UC-7.3). On success, the system confirms the appointment and sends a booking confirmation.",
        ["PRE-1: Patient is logged in with a verified account.",
         "PRE-2: The selected doctor is active and has generated appointment slots."],
        ["POST-1: An appointment record (status = confirmed) is created and the slot is marked booked.",
         "POST-2: A payment/transaction record is stored and a confirmation notification/email is sent."],
        ["1. Patient selects a doctor and clicks \u201cBook Appointment\u201d.",
         "2. System shows the doctor's calendar with available slots.",
         "3. Patient picks a date and an available time slot.",
         "4. Patient enters reason for visit and optionally selects an insurance card (UC-7.3).",
         "5. System calculates the fee (minus eligible discount) and holds the slot.",
         "6. Patient chooses payment method: Wallet / PayOS (UC-7.1) / Momo (UC-7.2).",
         "7. System processes the payment via the chosen gateway.",
         "8. System confirms the appointment, releases the hold as booked, and sends confirmation."],
        ["3a. Slot becomes unavailable while choosing \u2192 system refreshes the calendar and asks to repick.",
         "6a. Wallet balance insufficient \u2192 system suggests gateway payment or top-up.",
         "4a. No valid insurance card \u2192 full fee is charged."],
        ["7a. Gateway payment fails or times out \u2192 hold is released after 10 minutes and the appointment is not created; patient is informed."],
        "High", "Frequently \u2014 core business flow.",
        ["BR-01", "BR-13", "BR-14", "BR-16"],
        "Slot duration follows the doctor's work shift configuration; fee shown before payment.",
        "Payment gateways (PayOS/Momo) are reachable; doctor schedule is up to date.")]),

    ("Manage Appointments", [spec(
        "UC-8", "Manage Appointments", "KhoaNN", "Patient", "PayOS, Momo (refund)",
        "When the Patient opens \u201cMy Appointments\u201d from the dashboard menu.",
        "This use case allows the Patient to view upcoming and past appointments (UC-8), open one appointment detail (UC-8.1), and from the detail screen reschedule (UC-8.1.1), cancel (UC-8.1.2), or rate the doctor after completion (UC-8.1.3).",
        ["PRE-1: Patient is logged in.", "PRE-2: Patient has at least one appointment (for detail actions)."],
        ["POST-1: Appointment list reflects current statuses.",
         "POST-2: Reschedule/cancel updates the appointment and slot; rating creates a review."],
        ["1. Patient opens \u201cMy Appointments\u201d.",
         "2. System lists appointments grouped by upcoming / past with status badges.",
         "3. Patient clicks one appointment \u2192 detail screen (UC-8.1).",
         "4. [Reschedule] Patient picks a new available slot \u2265 24h before start \u2192 system moves the booking (UC-8.1.1).",
         "5. [Cancel] Patient confirms cancellation \u2265 24h before start \u2192 system cancels and refunds to wallet (UC-8.1.2).",
         "6. [Rate] For a completed visit, patient submits stars + comment (UC-8.1.3)."],
        ["4a. No alternative slot available \u2192 patient keeps or cancels the appointment.",
         "6a. Appointment already rated \u2192 rating form is hidden."],
        ["4b/5b. Action attempted < 24h before start \u2192 \u201cThis appointment can no longer be changed online. Please contact the clinic.\u201d"],
        "High", "Frequently.",
        ["BR-15", "BR-17"],
        "List is paginated; statuses: confirmed, completed, cancelled, no-show.",
        "Refund policy to wallet is accepted by the patient at booking time.")]),

    ("View Prescriptions", [spec(
        "UC-9", "View Prescription History", "TanhNT", "Patient", "None",
        "When the Patient opens \u201cMy Prescriptions\u201d from the dashboard menu.",
        "This use case allows the Patient to view a list of issued prescriptions (UC-9), open one prescription detail with medication lines and instructions (UC-9.1), display its QR code for pharmacy pickup (UC-9.1.1), and export it as PDF (UC-9.1.2).",
        ["PRE-1: Patient is logged in.", "PRE-2: At least one prescription has been issued by a doctor."],
        ["POST-1: The prescription list/detail is displayed.",
         "POST-2: QR code or PDF is generated on demand."],
        ["1. Patient opens \u201cMy Prescriptions\u201d.",
         "2. System lists prescriptions with date, doctor, and status.",
         "3. Patient opens a prescription \u2192 detail with medicine lines, dosage, and instructions.",
         "4. Patient clicks \u201cQR Code\u201d \u2192 popup shows scannable code (UC-9.1.1).",
         "5. Patient clicks \u201cExport PDF\u201d \u2192 system generates and downloads the file (UC-9.1.2)."],
        ["2a. No prescriptions yet \u2192 empty state with guidance."],
        ["5a. PDF generation fails \u2192 \u201cUnable to export file. Please try again.\u201d"],
        "Medium", "Occasionally \u2014 after each visit.",
        ["BR-19"],
        "Patient sees only their own prescriptions; QR encodes the prescription ID.",
        "Prescriptions were signed off by the doctor (UC-24.4).")]),

    ("Notifications", [spec(
        "UC-10", "View Notification Inbox", "ThangDQ", "Patient", "None",
        "When the Patient clicks the bell icon in the header.",
        "This use case allows the Patient to read system notifications (appointment reminders, results ready, booking confirmations) in one inbox (UC-10), configure which notifications to receive (UC-10.1), and mark items as read (UC-10.2).",
        ["PRE-1: Patient is logged in."],
        ["POST-1: Notifications are displayed with read/unread states.",
         "POST-2: Preference changes apply to future notifications."],
        ["1. Patient clicks the bell icon \u2192 inbox dropdown/page opens.",
         "2. System lists notifications, newest first, with unread highlighted.",
         "3. Patient clicks one notification \u2192 it is marked read and navigates to the related screen.",
         "4. Patient clicks \u201cMark all as read\u201d (UC-10.2).",
         "5. Patient opens \u201cPreferences\u201d and toggles notification types (UC-10.1)."],
        ["2a. No notifications \u2192 empty state."],
        ["5a. Saving preferences fails \u2192 previous settings are kept and an error toast is shown."],
        "Medium", "Frequently.",
        ["BR-20"],
        "Unread count badge shows up to 99+.",
        "Notification jobs (reminders, results) run on schedule.")]),

    ("Complaints (Patient)", [spec(
        "UC-11", "Submit Complaint", "ThangDQ", "Patient", "None",
        "When the Patient opens \u201cSupport / Complaints\u201d and clicks \u201cNew Complaint\u201d.",
        "This use case allows the Patient to submit a service complaint with category, subject, and description (UC-11), track submitted complaints in a list (UC-11.1), and read staff replies and status in the detail screen (UC-11.1.1).",
        ["PRE-1: Patient is logged in."],
        ["POST-1: A complaint record (status = Open) is created and visible to Staff.",
         "POST-2: Patient can follow replies and status updates."],
        ["1. Patient opens \u201cComplaints\u201d \u2192 \u201cNew Complaint\u201d.",
         "2. Patient selects category, enters subject and description, optionally attaches an image.",
         "3. Patient submits; system validates and creates the complaint with status Open.",
         "4. System notifies Staff and shows the complaint in \u201cMy Complaints\u201d list.",
         "5. Patient opens the detail later to read staff replies and status (UC-11.1.1)."],
        ["3a. Required fields missing \u2192 validation errors shown.",
         "5a. Complaint resolved \u2192 detail is read-only with the resolution note."],
        ["3b. Submit fails (server error) \u2192 draft is kept on the form for retry."],
        "Medium", "Occasionally.",
        ["BR-21", "BR-36"],
        "Attachments accept JPG/PNG up to 5 MB; one complaint may have many replies.",
        "Staff monitors the complaint queue during business hours.")]),

    ("Wallet", [spec(
        "UC-12", "View Wallet Balance", "ThangND", "Patient", "PayOS, Momo",
        "When the Patient opens \u201cMy Wallet\u201d from the dashboard menu.",
        "This use case allows the Patient to view their medical wallet balance (UC-12), review the transaction history of top-ups, payments, and refunds (UC-12.1), and top up the wallet through PayOS or Momo.",
        ["PRE-1: Patient is logged in.", "PRE-2: A wallet is auto-created with the patient account."],
        ["POST-1: Current balance and transactions are displayed accurately.",
         "POST-2: Successful top-up increases balance and adds a transaction record."],
        ["1. Patient opens \u201cMy Wallet\u201d.",
         "2. System shows the current balance and recent transactions.",
         "3. Patient opens \u201cTransaction History\u201d for the filterable full list (UC-12.1).",
         "4. Patient clicks \u201cTop Up\u201d, enters amount, and selects PayOS or Momo.",
         "5. System redirects to the gateway; patient completes the payment.",
         "6. Gateway callback verifies the payment; system credits the wallet and records the transaction."],
        ["4a. Amount below minimum top-up \u2192 validation message.",
         "6a. Patient cancels at the gateway \u2192 no balance change; transaction marked cancelled."],
        ["6b. Callback timeout \u2192 transaction stays pending; system reconciles with the gateway and updates the result."],
        "High", "Frequently \u2014 used in every paid booking.",
        ["BR-18", "BR-07"],
        "Transactions are immutable audit records; balance shown in VND.",
        "Payment gateways are reachable for top-up.")]),

    ("Insurance", [spec(
        "UC-13", "Manage Insurance Cards", "TruongNT", "Patient", "None",
        "When the Patient opens \u201cInsurance\u201d from the dashboard menu.",
        "This use case allows the Patient to view registered insurance cards (UC-13), add a new card manually or by scanning the card image with OCR auto-fill (UC-13.2, UC-13.2.1), update card information (UC-13.3), and delete a card (UC-13.1). Valid cards yield a discount at booking (UC-7.3).",
        ["PRE-1: Patient is logged in."],
        ["POST-1: The card list reflects all add/update/delete changes.",
         "POST-2: Valid cards are selectable during booking."],
        ["1. Patient opens \u201cInsurance\u201d \u2192 list of registered cards with validity badges.",
         "2. [Add] Patient clicks \u201cAdd Card\u201d, fills provider, card number, holder name, expiry \u2014 or uploads a photo for OCR auto-fill (UC-13.2.1) \u2014 then saves.",
         "3. System validates the card number format and expiry date.",
         "4. [Update] Patient edits an existing card and saves (UC-13.3).",
         "5. [Delete] Patient deletes a card after confirmation (UC-13.1)."],
        ["2a. OCR cannot read the image \u2192 patient fills the fields manually.",
         "3a. Card expired \u2192 card is saved but flagged \u201cExpired\u201d and excluded from discounts."],
        ["2b. Image upload fails \u2192 \u201cUnable to process image. Please try again or enter manually.\u201d"],
        "Medium", "Occasionally.",
        ["BR-16"],
        "Card images are stored privately; OCR is best-effort assistance only.",
        "Insurance providers list is maintained by Admin.")]),

    ("Queue Status", [spec(
        "UC-14", "View Queue Status", "TruongNT", "Patient", "None",
        "When the Patient opens \u201cQueue Status\u201d after checking in at the clinic.",
        "This use case allows the Patient to see their current queue ticket number, the number now being served, and the estimated waiting count for the clinic room of their appointment. The screen auto-refreshes while the queue session is active.",
        ["PRE-1: Patient is logged in.",
         "PRE-2: Staff has issued a queue ticket for the patient's checked-in appointment (UC-33)."],
        ["POST-1: The patient sees live ticket progress until being served."],
        ["1. Patient opens \u201cQueue Status\u201d from the dashboard or the booking confirmation.",
         "2. System shows the patient's ticket number, current called number, and people ahead.",
         "3. Screen refreshes automatically as the Doctor calls next patients (UC-18.1).",
         "4. When the patient's number is called, the screen highlights \u201cIt's your turn \u2014 Room X\u201d."],
        ["2a. No active ticket \u2192 message \u201cYou have no queue ticket today.\u201d"],
        ["3a. Connection lost \u2192 screen shows last known state with a reconnect notice."],
        "Medium", "Frequently on visit days.",
        ["BR-22"],
        "Live updates via polling/WebSocket; ticket is valid for the current day only.",
        "Patient checked in at the clinic reception.")]),

    ("Favorites", [spec(
        "UC-15", "Manage Favorite Doctors", "TanhNT", "Patient", "None",
        "When the Patient opens \u201cFavorite Doctors\u201d or clicks the heart icon on a doctor profile.",
        "This use case allows the Patient to view saved favorite doctors (UC-15), add a doctor to favorites from any doctor card or profile (UC-15.1), and remove a doctor from the list (UC-15.2). Favorites give quick access to rebooking.",
        ["PRE-1: Patient is logged in."],
        ["POST-1: The favorites list reflects add/remove actions immediately."],
        ["1. Patient clicks the heart icon on a doctor profile \u2192 doctor added to favorites (UC-15.1).",
         "2. Patient opens \u201cFavorite Doctors\u201d \u2192 grid of saved doctors with quick \u201cBook\u201d buttons.",
         "3. Patient clicks the heart icon again or \u201cRemove\u201d \u2192 doctor removed (UC-15.2)."],
        ["2a. Empty list \u2192 empty state suggesting doctor search."],
        ["1a. Action fails \u2192 toast error; icon state reverts."],
        "Low", "Occasionally.",
        ["BR-02"],
        "A deactivated doctor is hidden from the favorites list automatically.",
        "None.")]),

    ("Doctor Dashboard", [spec(
        "UC-16", "View Doctor Dashboard", "ThangND", "Doctor", "None",
        "When the Doctor logs in successfully (UC-5) or clicks \u201cDashboard\u201d in the doctor console.",
        "This use case allows the Doctor to view an overview screen after login: today's appointment count, active queue session status, pending encounters, and shortcuts to consultation, schedule, and EMR screens. The Doctor can also manage their own professional profile (UC-16.1).",
        ["PRE-1: Doctor is logged in with role = doctor."],
        ["POST-1: Dashboard widgets show live counts scoped to this doctor."],
        ["1. Doctor logs in \u2192 system redirects to the Doctor Dashboard.",
         "2. System loads today's appointments, queue session state, and pending sign-offs.",
         "3. Doctor clicks a widget to jump to Today Appointments (UC-17), Queue Session (UC-18), or Schedule (UC-19/20).",
         "4. Doctor opens \u201cMy Profile\u201d to view/update professional info (UC-16.1)."],
        ["2a. No appointments today \u2192 widgets show zero states."],
        ["2b. Widget data fails to load \u2192 the widget shows a retry action without blocking the page."],
        "High", "Frequently \u2014 every working day.",
        ["BR-08", "BR-37"],
        "Profile fields editable by doctor: bio, photo, contact; specialty/department are admin-managed.",
        "Doctor account is linked to a doctor profile record.")]),

    ("Consultation & EMR", [spec(
        "UC-17", "Start Consultation", "ThangDQ", "Doctor", "None",
        "When the Doctor opens \u201cToday Appointments\u201d and clicks \u201cStart Consultation\u201d on a checked-in patient.",
        "This use case allows the Doctor to see today's appointment list (UC-17) and start a clinical session for a checked-in patient (UC-17.1). Starting a consultation creates an encounter record (UC-17.1.1, included). Within the encounter, the Doctor views detail (UC-17.1.1.1), updates clinical notes (UC-17.1.1.1.1), records ICD-10 diagnoses (UC-17.1.1.2 with search UC-17.1.1.2.2, update UC-17.1.1.2.3, remove UC-17.1.1.2.1), and finally signs off the encounter (UC-17.1.1.1.2).",
        ["PRE-1: Doctor is logged in.",
         "PRE-2: The patient has a checked-in appointment with this doctor today."],
        ["POST-1: An encounter record linked to the appointment is created and completed.",
         "POST-2: After sign-off, the encounter becomes read-only and appears in the patient EMR timeline (UC-21)."],
        ["1. Doctor opens \u201cToday Appointments\u201d \u2192 list with check-in statuses.",
         "2. Doctor clicks \u201cStart Consultation\u201d on a checked-in patient.",
         "3. System creates an encounter (status = in-progress) and opens the encounter workspace.",
         "4. Doctor reviews patient EMR timeline (UC-21) and writes chief complaint / clinical notes.",
         "5. Doctor adds diagnoses by searching the ICD-10 catalog and selecting codes.",
         "6. Doctor optionally uploads medical images (UC-23) and creates a prescription (UC-24).",
         "7. Doctor clicks \u201cSign Off\u201d; system validates required fields and locks the encounter.",
         "8. System marks the appointment completed and notifies the patient that results are ready."],
        ["5a. ICD-10 code not found \u2192 doctor refines the search keyword.",
         "7a. Required fields missing (no diagnosis) \u2192 sign-off blocked with a checklist."],
        ["3a. Encounter already exists for the appointment \u2192 system reopens the existing encounter instead of duplicating."],
        "High", "Frequently \u2014 every consultation.",
        ["BR-23", "BR-24", "BR-25"],
        "Encounter autosaves notes as draft; sign-off requires at least one diagnosis.",
        "Patient was checked in by Staff (UC-33).")]),

    ("Queue Session", [spec(
        "UC-18", "Manage Queue Session", "TruongNT", "Doctor", "None",
        "When the Doctor opens \u201cQueue Session\u201d for their assigned clinic room.",
        "This use case allows the Doctor to open a queue session for the clinic room (UC-18), call the next patient ticket (UC-18.1), recall a previously called ticket (UC-18.2), and close the session at the end of the shift (UC-18.3). Patient queue screens update live (UC-14).",
        ["PRE-1: Doctor is logged in and assigned to a clinic room for the current shift.",
         "PRE-2: No other active session exists for the same room."],
        ["POST-1: Session state and called numbers are persisted.",
         "POST-2: Closing the session finalizes remaining tickets (no-show or transferred)."],
        ["1. Doctor opens \u201cQueue Session\u201d and clicks \u201cOpen Session\u201d for the room.",
         "2. System activates the session; Staff can now issue tickets to it (UC-33).",
         "3. Doctor clicks \u201cCall Next\u201d \u2192 the next waiting ticket is marked called and displayed (UC-18.1).",
         "4. If the patient is absent, Doctor clicks \u201cRecall\u201d on a previous ticket (UC-18.2).",
         "5. At the end of the shift, Doctor clicks \u201cClose Session\u201d and confirms (UC-18.3)."],
        ["3a. No waiting tickets \u2192 button disabled with \u201cQueue is empty\u201d.",
         "4a. Recalled patient still absent \u2192 ticket can be marked no-show."],
        ["1a. Another active session exists for the room \u2192 \u201cA session is already open for this room.\u201d"],
        "High", "Daily during clinic hours.",
        ["BR-26", "BR-35"],
        "Called numbers display on the room screen; session log kept for audit.",
        "Clinic room assignment follows the doctor work shift (UC-19).")]),

    ("Work Shift Schedule", [spec(
        "UC-19", "View Work Shift Schedule", "TruongNT", "Doctor", "None",
        "When the Doctor opens \u201cMy Shifts\u201d in the doctor console.",
        "This use case allows the Doctor to view their assigned work shifts (created by Admin, UC-29 related) in a weekly/monthly calendar, including shift time, clinic room, and status. Shifts drive appointment slot generation.",
        ["PRE-1: Doctor is logged in.", "PRE-2: Admin has created work shifts for the doctor."],
        ["POST-1: The shift calendar is displayed accurately for the selected period."],
        ["1. Doctor opens \u201cMy Shifts\u201d.",
         "2. System loads shifts for the current week with room and time details.",
         "3. Doctor navigates between weeks/months to view other periods.",
         "4. Doctor clicks a shift to see detail (room, slot duration, generated slots)."],
        ["2a. No shifts in the selected period \u2192 empty calendar state."],
        ["2b. Loading fails \u2192 retry option shown."],
        "Medium", "Weekly.",
        ["BR-08"],
        "Shift changes are managed by Admin; doctor has read-only access.",
        "Work shifts were generated before the schedule is consulted.")]),

    ("Appointment Calendar", [spec(
        "UC-20", "View Appointment Calendar", "TruongNT", "Doctor", "None",
        "When the Doctor opens \u201cMy Calendar\u201d in the doctor console.",
        "This use case allows the Doctor to view their personal appointment calendar (UC-20) with booked/available/blocked slots, block unavailable time slots (UC-20.1), and unblock previously blocked slots (UC-20.2).",
        ["PRE-1: Doctor is logged in.", "PRE-2: Appointment slots have been generated from work shifts."],
        ["POST-1: Blocked/unblocked slots are persisted and excluded/included in patient booking."],
        ["1. Doctor opens \u201cMy Calendar\u201d \u2192 day/week view of slots with statuses.",
         "2. Doctor clicks an empty slot and selects \u201cBlock\u201d with an optional reason (UC-20.1).",
         "3. System marks the slot blocked; patients can no longer book it.",
         "4. Doctor clicks a blocked slot and selects \u201cUnblock\u201d to reopen it (UC-20.2)."],
        ["2a. Slot already booked \u2192 block action unavailable; doctor contacts Staff for rescheduling."],
        ["3a. Save fails \u2192 slot state reverts with an error toast."],
        "Medium", "Weekly.",
        ["BR-27"],
        "Blocking many slots at once supported by drag selection.",
        "Slot generation job has run for the visible period.")]),

    ("Patient EMR Timeline", [spec(
        "UC-21", "View Patient EMR Timeline", "ThangDQ", "Doctor", "None",
        "When the Doctor opens the \u201cEMR\u201d tab of a patient during or before a consultation.",
        "This use case allows the Doctor to review a patient's medical history in a chronological timeline: past encounters, diagnoses, prescriptions, and medical images. Each item links to its detail view.",
        ["PRE-1: Doctor is logged in.",
         "PRE-2: Doctor has an encounter/appointment relationship with the patient (BR-23)."],
        ["POST-1: The timeline is displayed; no data is modified."],
        ["1. Doctor opens patient EMR from the encounter workspace or today list.",
         "2. System loads the timeline grouped by visit date (newest first).",
         "3. Doctor filters by type: encounters / diagnoses / prescriptions / images.",
         "4. Doctor clicks an item to open its read-only detail."],
        ["2a. New patient with no history \u2192 empty timeline state."],
        ["2b. Loading fails \u2192 retry option shown."],
        "High", "Every consultation.",
        ["BR-23"],
        "Signed-off encounters only; drafts are visible to the owning doctor.",
        "EMR data was recorded in previous encounters.")]),

    ("Medical Imaging", [spec(
        "UC-23", "Manage Medical Images", "ThangND", "Doctor", "None",
        "When the Doctor opens the \u201cImaging\u201d tab inside an encounter.",
        "This use case allows the Doctor to upload diagnostic images for a patient encounter (UC-23), browse the imaging gallery (UC-23.1), view an image fullscreen (UC-23.1.1), and delete a wrongly uploaded image (UC-23.1.2).",
        ["PRE-1: Doctor is logged in.", "PRE-2: An encounter is open (not signed off) for upload/delete actions."],
        ["POST-1: Uploaded images are linked to the encounter and visible in the gallery and EMR timeline."],
        ["1. Doctor opens the \u201cImaging\u201d tab of the encounter.",
         "2. Doctor clicks \u201cUpload\u201d, selects files, adds modality/description, and confirms.",
         "3. System validates type/size and stores the images.",
         "4. Doctor browses thumbnails in the gallery (UC-23.1) and clicks one for fullscreen view (UC-23.1.1).",
         "5. Doctor deletes an incorrect image with confirmation (UC-23.1.2)."],
        ["3a. Unsupported file type or oversize \u2192 file rejected with message; other files continue."],
        ["2a. Upload interrupted \u2192 partial files are discarded; doctor retries."],
        "Medium", "Per consultation when imaging is involved.",
        ["BR-24", "BR-29"],
        "Images stored on cloud storage; gallery paginates by 12 items.",
        "Encounter is still editable (not signed off) for modifications.")]),

    ("Prescription Management", [spec(
        "UC-24", "Create Prescription", "ThangDQ", "Doctor", "None",
        "When the Doctor clicks \u201cCreate Prescription\u201d inside an encounter.",
        "This use case allows the Doctor to create a prescription for the encounter (UC-24): add medicine line items from the formulary (UC-24.1), update dosage/instructions (UC-24.2), remove lines (UC-24.3), and save/issue the prescription (UC-24.4). Afterwards the Doctor can review the detail (UC-24.5), export PDF (UC-24.5.1), or show the QR code (UC-24.5.2).",
        ["PRE-1: Doctor is logged in with an open encounter.",
         "PRE-2: Medicines exist in the pharmacy formulary (UC-32)."],
        ["POST-1: An issued prescription is linked to the encounter and visible to the patient (UC-9).",
         "POST-2: Pharmacy stock is reserved/decremented per dispensing policy."],
        ["1. Doctor clicks \u201cCreate Prescription\u201d in the encounter workspace.",
         "2. Doctor searches a medicine and adds a line with quantity, dosage, frequency, duration, and notes (UC-24.1).",
         "3. Doctor repeats for all medicines; edits (UC-24.2) or removes (UC-24.3) lines as needed.",
         "4. System validates each line (stock, dosage format, duplicates).",
         "5. Doctor clicks \u201cSave & Issue\u201d (UC-24.4); system stores the prescription with status issued.",
         "6. System makes QR/PDF available (UC-24.5.1, UC-24.5.2) and shows it in the patient portal."],
        ["4a. Medicine out of stock \u2192 warning with suggested alternatives; doctor may keep it as external purchase.",
         "2a. Duplicate medicine line \u2192 system prompts to merge quantities."],
        ["5a. Save fails \u2192 draft retained locally for retry."],
        "High", "Most consultations.",
        ["BR-19", "BR-28"],
        "Drug-interaction warning is advisory only in this release.",
        "Formulary data is maintained by Staff (UC-32).")]),

    ("Manage Accounts", [spec(
        "UC-25", "Manage Accounts", "KhoaNN", "Admin", "None",
        "When the Admin opens \u201cAccounts\u201d in the admin console.",
        "This use case allows the Admin to view the accounts list with filters by role/status, view account detail, create staff accounts, update account info, change user roles, and deactivate/restore accounts.",
        ["PRE-1: Admin is logged in with role = admin."],
        ["POST-1: Account changes are persisted and effective immediately.",
         "POST-2: Deactivated accounts cannot log in and active sessions are revoked."],
        ["1. Admin opens \u201cAccounts\u201d \u2192 paginated list with search and role/status filters.",
         "2. Admin clicks an account \u2192 detail view with profile and activity info.",
         "3. [Create] Admin clicks \u201cCreate Account\u201d, fills email/name/role = staff, system emails initial credentials.",
         "4. [Update] Admin edits account fields and saves.",
         "5. [Change Role] Admin changes the role with confirmation.",
         "6. [Deactivate/Restore] Admin toggles account status with confirmation."],
        ["3a. Email already exists \u2192 \u201cEmail already registered.\u201d",
         "6a. Admin attempts to deactivate own account \u2192 action blocked."],
        ["3b. Credential email fails \u2192 account created; Admin can resend credentials."],
        "High", "Weekly.",
        ["BR-03", "BR-06", "BR-30", "BR-33"],
        "All account actions are written to the audit log.",
        "Admin console is restricted to the admin role.")]),

    ("Manage Specialties", [spec(
        "UC-26", "Manage Specialties", "KhoaNN", "Admin", "None",
        "When the Admin opens \u201cSpecialties\u201d under Master Data.",
        "This use case allows the Admin to create, view, update, and delete medical specialties. Specialties are referenced by doctors and used for patient search filters.",
        ["PRE-1: Admin is logged in."],
        ["POST-1: The specialties catalog reflects all changes and is used across doctor profiles and search."],
        ["1. Admin opens \u201cSpecialties\u201d \u2192 list with search.",
         "2. [Create] Admin clicks \u201cCreate\u201d, enters name + description + icon, and saves.",
         "3. [Update] Admin edits a specialty and saves.",
         "4. [Delete] Admin deletes an unused specialty after confirmation."],
        ["2a. Duplicate name \u2192 validation error."],
        ["4a. Specialty referenced by doctors \u2192 hard delete blocked; Admin deactivates instead (BR-31)."],
        "Medium", "Rarely \u2014 setup and occasional maintenance.",
        ["BR-31"],
        "Specialty icons shown on the public portal.",
        "None.")]),

    ("Manage Departments", [spec(
        "UC-27", "Manage Departments", "ThangDQ", "Admin", "None",
        "When the Admin opens \u201cDepartments\u201d under Master Data.",
        "This use case allows the Admin to create departments, view the list and detail, update information, and deactivate departments no longer in operation. Departments group doctors and clinic rooms.",
        ["PRE-1: Admin is logged in."],
        ["POST-1: The department catalog reflects the changes; deactivated departments are hidden from public filters."],
        ["1. Admin opens \u201cDepartments\u201d \u2192 list with status badges.",
         "2. [Create] Admin enters name, description, location, and head doctor, then saves.",
         "3. [View] Admin opens a department detail with linked rooms and doctors.",
         "4. [Update] Admin edits and saves changes.",
         "5. [Deactivate] Admin deactivates with confirmation."],
        ["2a. Duplicate name \u2192 validation error."],
        ["5a. Department has assigned doctors/rooms \u2192 system warns and requires reassignment before deactivation (BR-31)."],
        "Medium", "Rarely.",
        ["BR-31"],
        "Department detail shows linked clinic rooms (UC-28).",
        "None.")]),

    ("Manage Clinic Rooms", [spec(
        "UC-28", "Manage Clinic Rooms", "KhoaNN", "Admin", "None",
        "When the Admin opens \u201cClinic Rooms\u201d under Master Data.",
        "This use case allows the Admin to create clinic rooms (code, name, department, floor), view the rooms list, and update room information. Rooms are used for doctor work shifts and queue sessions.",
        ["PRE-1: Admin is logged in.", "PRE-2: The owning department exists (UC-27)."],
        ["POST-1: Room catalog reflects all changes; rooms are selectable in shift planning."],
        ["1. Admin opens \u201cClinic Rooms\u201d \u2192 list filtered by department.",
         "2. [Create] Admin enters room code, name, department, floor, and saves.",
         "3. [Update] Admin edits room info or sets the room inactive."],
        ["2a. Duplicate room code \u2192 validation error."],
        ["3a. Room has scheduled shifts \u2192 deactivation requires moving the shifts first (BR-31)."],
        "Medium", "Rarely.",
        ["BR-31"],
        "Room code printed on queue tickets.",
        "None.")]),

    ("Manage Doctors", [spec(
        "UC-29", "Manage Doctors", "ThangND", "Admin", "None",
        "When the Admin opens \u201cDoctors\u201d in the admin console.",
        "This use case allows the Admin to create doctor profiles with linked accounts, view the doctors list and detail, update profiles, deactivate doctors, and import/export doctors via Excel. It also covers work shift creation and slot generation for doctors.",
        ["PRE-1: Admin is logged in.", "PRE-2: Specialty and department catalogs exist."],
        ["POST-1: Doctor profiles and accounts reflect all changes.",
         "POST-2: Imported doctors receive accounts; exported file contains the filtered list."],
        ["1. Admin opens \u201cDoctors\u201d \u2192 list with specialty/department filters.",
         "2. [Create] Admin enters profile (name, specialty, department, degree, fee) \u2192 system creates the linked account and emails credentials.",
         "3. [Update] Admin edits profile fields and saves.",
         "4. [Deactivate] Admin deactivates a doctor \u2192 hidden from public pages, future slots disabled.",
         "5. [Import] Admin uploads the .xlsx template \u2192 system validates rows and reports results (UC-29 import).",
         "6. [Export] Admin exports the current filtered list to Excel.",
         "7. [Shifts] Admin creates work shifts and generates appointment slots for doctors."],
        ["5a. Some rows invalid \u2192 valid rows committed; error report downloadable.",
         "4a. Doctor has future confirmed appointments \u2192 system requires rescheduling/cancellation first."],
        ["2a. Email already used \u2192 creation blocked with message."],
        "High", "Weekly.",
        ["BR-02", "BR-30", "BR-31", "BR-32", "BR-33"],
        "Slot generation derives from shift templates (duration per slot).",
        "HR provides the import template data.")]),

    ("Manage Patients", [spec(
        "UC-30", "Manage Patients", "TanhNT", "Admin", "None",
        "When the Admin opens \u201cPatients\u201d in the admin console.",
        "This use case allows the Admin to view the patients list with search/filters, view patient detail (profile, appointment history), create a patient record manually (walk-in), update patient profiles, and activate/deactivate patient accounts.",
        ["PRE-1: Admin is logged in."],
        ["POST-1: Patient records reflect the changes; deactivated patients cannot log in."],
        ["1. Admin opens \u201cPatients\u201d \u2192 paginated list with search by name/phone/email.",
         "2. Admin opens a patient detail \u2192 profile, appointments, wallet summary.",
         "3. [Create] Admin creates a walk-in patient with basic info (no email verification required).",
         "4. [Update] Admin edits profile fields and saves.",
         "5. [Deactivate/Activate] Admin toggles the account with confirmation."],
        ["3a. Phone/email already exists \u2192 system suggests opening the existing record."],
        ["5a. Patient has future appointments \u2192 system warns and lists them before deactivation."],
        "Medium", "Weekly.",
        ["BR-06", "BR-33"],
        "Medical history remains read-only for Admin (clinical data belongs to doctors).",
        "None.")]),

    ("Admin Dashboard", [spec(
        "UC-31", "View Admin Dashboard", "TanhNT", "Admin", "None",
        "When the Admin logs in successfully or clicks \u201cDashboard\u201d in the admin console.",
        "This use case allows the Admin to view a dashboard with system KPIs: revenue summary with date/doctor filters and Excel export, appointment volume, new patients, active doctors, and module shortcuts.",
        ["PRE-1: Admin is logged in."],
        ["POST-1: KPI widgets and the revenue chart display correct aggregated data."],
        ["1. Admin logs in \u2192 redirected to the Admin Dashboard.",
         "2. System loads KPI widgets (revenue, appointments, patients, complaints).",
         "3. Admin filters revenue by date range and/or doctor.",
         "4. Admin clicks \u201cExport Excel\u201d to download the revenue report.",
         "5. Admin clicks a widget to jump to the related module."],
        ["3a. No data in range \u2192 chart shows empty state."],
        ["4a. Export fails \u2192 error toast with retry."],
        "High", "Daily.",
        ["BR-37", "BR-07"],
        "Revenue aggregates wallet payments and gateway transactions.",
        "Transactions are recorded consistently by the payment module.")]),

    ("Manage Pharmacy", [spec(
        "UC-32", "Manage Pharmacy", "TruongNT", "Staff", "None",
        "When the Staff opens \u201cPharmacy\u201d in the staff console.",
        "This use case allows the Staff to manage the medicines inventory: register new medicines, view the inventory list and medicine detail, update medicine info, record stock inbound/outbound, and monitor low-stock alerts.",
        ["PRE-1: Staff is logged in with role = staff."],
        ["POST-1: Inventory levels and medicine catalog reflect all changes.",
         "POST-2: Low-stock alerts update according to reorder levels."],
        ["1. Staff opens \u201cPharmacy\u201d \u2192 inventory list with stock levels and alert badges.",
         "2. [Create] Staff registers a medicine: name, unit, price, reorder level.",
         "3. [Update] Staff edits medicine info and saves.",
         "4. [Inbound] Staff records received stock with quantity, batch, and expiry date.",
         "5. [Outbound] Staff records dispensed/disposed stock with reason.",
         "6. System recalculates stock and raises low-stock alerts when below reorder level."],
        ["5a. Outbound quantity exceeds stock \u2192 blocked with message (BR-34)."],
        ["4a. Save fails \u2192 transaction is not recorded; staff retries."],
        "High", "Daily.",
        ["BR-34"],
        "Stock movements are immutable audit records; expiring batches highlighted.",
        "Physical stock counts match recorded movements.")]),

    ("Queue Check-in", [spec(
        "UC-33", "Manage Queue Check-in", "TruongNT", "Staff", "None",
        "When the Staff opens \u201cQueue Check-in\u201d at the reception desk.",
        "This use case allows the Staff to check in arriving patients and issue queue tickets tied to their appointments. The ticket joins the active queue session of the assigned clinic room (UC-18).",
        ["PRE-1: Staff is logged in.",
         "PRE-2: The patient has a confirmed appointment today.",
         "PRE-3: A queue session is open for the target clinic room (UC-18)."],
        ["POST-1: The appointment is marked checked-in and a queue ticket is issued.",
         "POST-2: The patient can follow live queue status (UC-14)."],
        ["1. Staff searches the appointment by patient name, phone, or booking code.",
         "2. System shows the appointment with doctor and room info.",
         "3. Staff confirms identity and clicks \u201cCheck In & Issue Ticket\u201d.",
         "4. System marks the appointment checked-in, issues the next ticket number, and prints/displays it.",
         "5. The ticket appears in the doctor's queue (UC-18) and on the patient's queue screen (UC-14)."],
        ["1a. Appointment not found \u2192 staff verifies booking info or directs the patient to booking support.",
         "3a. Patient arrives too early/late per policy \u2192 staff follows clinic guidance before check-in."],
        ["3b. No open queue session for the room \u2192 ticket cannot be issued; staff notifies the doctor to open the session."],
        "High", "Continuously during clinic hours.",
        ["BR-35", "BR-26"],
        "Ticket numbers reset per session per room.",
        "Reception verifies patient identity at the desk.")]),

    ("Manage Complaints (Staff)", [spec(
        "UC-34", "Manage Complaints", "KhoaNN", "Staff", "None",
        "When the Staff opens \u201cComplaints\u201d in the staff console.",
        "This use case allows the Staff to view all submitted complaints with filters, open a complaint detail, reply to the patient, and update the complaint status (Open \u2192 In Progress \u2192 Resolved).",
        ["PRE-1: Staff is logged in.", "PRE-2: At least one complaint exists."],
        ["POST-1: Replies and status changes are persisted and visible to the patient (UC-11.1.1).",
         "POST-2: The patient is notified about replies/resolution."],
        ["1. Staff opens \u201cComplaints\u201d \u2192 list filtered by status/category/date.",
         "2. Staff opens a complaint detail with full description and history.",
         "3. Staff writes a reply and sends it; system notifies the patient.",
         "4. Staff updates the status (In Progress while handling, Resolved with a resolution note).",
         "5. Resolved complaints become read-only for the patient."],
        ["3a. More information needed \u2192 staff requests details via reply; status stays In Progress."],
        ["3b. Notification delivery fails \u2192 reply is stored; notification retried by the job queue."],
        "Medium", "Daily.",
        ["BR-21", "BR-36"],
        "SLA dashboards may track time-to-resolution (future).",
        "Complaint categories are predefined.")]),

    ("Staff Dashboard", [spec(
        "UC-35", "View Staff Dashboard", "TanhNT", "Staff", "None",
        "When the Staff logs in successfully or clicks \u201cDashboard\u201d in the staff console.",
        "This use case allows the Staff to view an operational dashboard after login: today's check-in count, open complaints, low-stock alerts, and shortcuts to Queue Check-in, Pharmacy, and Complaints modules.",
        ["PRE-1: Staff is logged in with role = staff."],
        ["POST-1: Dashboard widgets show live operational counters scoped to staff permissions."],
        ["1. Staff logs in \u2192 redirected to the Staff Dashboard.",
         "2. System loads widgets: today's check-ins, waiting tickets, open complaints, low-stock alerts.",
         "3. Staff clicks a widget to jump to the related module (UC-32, UC-33, UC-34)."],
        ["2a. No activity yet today \u2192 zero states shown."],
        ["2b. A widget fails to load \u2192 retry action shown without blocking the page."],
        "Medium", "Daily.",
        ["BR-08", "BR-37"],
        "Widgets auto-refresh every 60 seconds.",
        "Staff account is provisioned by Admin (UC-25).")]),
]

# ------------------------------------------------------------ docx builders
LABEL_W = Inches(1.45)
COL2_W = Inches(1.69)
COL3_W = Inches(1.31)
COL4_W = Inches(2.04)


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    lines = text.split("\n") if isinstance(text, str) else list(text)
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)


def add_spec_table(doc, s):
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    table.autofit = False

    def full_row(label, value):
        row = table.add_row()
        set_cell_text(row.cells[0], label, bold=True)
        merged = row.cells[1].merge(row.cells[2]).merge(row.cells[3])
        set_cell_text(merged, value)

    def pair_row(l1, v1, l2, v2):
        row = table.add_row()
        set_cell_text(row.cells[0], l1, bold=True)
        set_cell_text(row.cells[1], v1)
        set_cell_text(row.cells[2], l2, bold=True)
        set_cell_text(row.cells[3], v2)

    full_row("UC ID and Name:", f"{s['id']}_{s['name']}")
    pair_row("Created By:", s["by"], "Date Created:", DATE)
    pair_row("Primary Actor:", s["primary"], "Secondary Actors:", s["secondary"])
    full_row("Trigger:", s["trigger"])
    full_row("Description:", s["desc"])
    full_row("Preconditions:", s["pre"])
    full_row("Postconditions:", s["post"])
    full_row("Normal Flow:", s["normal"])
    full_row("Alternative Flows:", s["alt"] or ["None."])
    full_row("Exceptions:", s["exc"] or ["None."])
    full_row("Priority:", s["priority"])
    full_row("Frequency of Use:", s["freq"])
    full_row("Business Rules:", ", ".join(s["brs"]))
    full_row("Other Information:", s["other"])
    full_row("Assumptions:", s["assume"])

    for row in table.rows:
        row.cells[0].width = LABEL_W
        row.cells[1].width = COL2_W
        row.cells[2].width = COL3_W
        row.cells[3].width = COL4_W
    return table


def add_br_table(doc, br_ids):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for i, h in enumerate(["ID", "Category", "Business Rule Description"]):
        set_cell_text(hdr.cells[i], h, bold=True)
    for br_id in br_ids:
        cat, desc = BR[br_id]
        row = table.add_row()
        set_cell_text(row.cells[0], br_id)
        set_cell_text(row.cells[1], cat)
        set_cell_text(row.cells[2], desc)
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(1.3)
        row.cells[2].width = Inches(4.4)
    return table


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_heading("II. Requirement Specifications", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    note = doc.add_paragraph()
    run = note.add_run(
        "Note: UC-22 is reserved. Patient inherits Guest use cases (UC-1, UC-2); "
        "Doctor inherits Patient use cases (UC-3, UC-7 ~ UC-15)."
    )
    run.italic = True
    run.font.size = Pt(10)

    for gi, (feature, specs) in enumerate(GROUPS, start=1):
        doc.add_heading(f"{gi}. {feature}", level=2)
        for si, s in enumerate(specs, start=1):
            doc.add_heading(f"{gi}.{si} {s['id']}_{s['name']}", level=3)
            doc.add_heading("a. Functionalities", level=4)
            add_spec_table(doc, s)
            doc.add_paragraph()
            doc.add_heading("b. Business Rules", level=4)
            add_br_table(doc, s["brs"])
            doc.add_paragraph()

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    print(f"Groups: {len(GROUPS)}, Specs: {sum(len(s) for _, s in GROUPS)}")


if __name__ == "__main__":
    main()
