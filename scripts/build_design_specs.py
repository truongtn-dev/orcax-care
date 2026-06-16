"""
Build Section III. Design Specifications for OrcaXCare (FBUS format).
Per function:
  Heading 2: <n>. <Feature>
  Heading 3: <n>.<m> UC-x_<Name>
  - Description / Related Use Case bullets
  UI Design        -> inserts image from docs/ui-design if found, else placeholder
  Field table      -> Field Name | Field Type | Description & Data Initialization
  Database Access  -> Collection | CRUD | Description  (MongoDB)
  MongoDB Commands -> representative queries
Output: e:/SU26/WDP301/WDP301-SE1816-GROUP4_DesignSpecs.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT_PATH = Path(r"e:/SU26/WDP301/WDP301-SE1816-GROUP4_DesignSpecs.docx")
UI_DIR = Path(r"e:/SU26/WDP301/orcax-care/docs/ui-design")
HEADER_BG = "FCE4D6"  # light header shade similar to sample


def shade(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_text(t.rows[0].cells[i], h, bold=True)
        shade(t.rows[0].cells[i], HEADER_BG)
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            set_text(cells[i], v)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    return t


def add_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(text)


def add_sub_heading(doc, text, underline=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.underline = underline
    run.font.size = Pt(12)


def add_code(doc, text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)


FIG_START = 13  # first figure number (1..12 belong to earlier document sections)
FIG = {"n": FIG_START - 1}


def _field_run(p, fld_type):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = p.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), fld_type)
    run._r.append(el)
    return run


def _instr_run(p, instr):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = p.add_run()
    el = OxmlElement("w:instrText")
    el.set(qn("xml:space"), "preserve")
    el.text = instr
    run._r.append(el)
    return run


def add_caption(doc, text):
    """Real Word caption: 'Figure { SEQ Figure \\* ARABIC }: text' with Caption
    style, so Insert > Table of Figures picks it up."""
    FIG["n"] += 1
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Figure ")
    _field_run(p, "begin")
    _instr_run(p, r" SEQ Figure \* ARABIC ")
    _field_run(p, "separate")
    p.add_run(str(FIG["n"]))  # cached number shown until fields refresh (F9)
    _field_run(p, "end")
    p.add_run(f": UI Design {text}")
    for run in p.runs:
        run.font.size = Pt(10)


def add_ui_design(doc, image_name, caption):
    add_sub_heading(doc, "UI Design")
    img = UI_DIR / image_name
    if img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Inches(5.8))
    else:
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ph.add_run(f"[ Insert UI image: docs/ui-design/{image_name} ]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    add_caption(doc, caption)


# ----------------------------------------------------------------- content
SECTIONS = [
    {
        "feature": "View Portal Home",
        "uc": "UC-1", "name": "View Portal Home",
        "description": "Allows Guest to view the public home page of OrcaXCare including hero search, services overview, featured doctors, branches, and news. From here Guest can search doctors, view a doctor public profile, locate branches, or go to Login/Register.",
        "related": "UC-1 \u2013 View Portal Home (UC-1.1 Search Doctors, UC-1.2 View Featured Doctors, UC-1.3 Locate Branch)",
        "image": "UI_PortalHome.png", "caption": "Portal Home",
        "fields": [
            ("Search keyword", "Text Input", "Guest enters doctor name to search."),
            ("Specialty", "Dropdown", "Guest filters doctors by specialty (loaded from specialties collection)."),
            ("Department", "Dropdown", "Guest filters doctors by department."),
            ("Search", "Button", "Triggers doctor search with the selected filters (UC-1.1)."),
            ("Featured Doctors", "Card List", "Shows up to 8 active doctors with name, specialty, rating, and \u201cBook\u201d button."),
            ("Branches", "Link/Map", "Opens branch list with map locations (UC-1.3)."),
            ("Login / Register", "Button", "Navigates to Login (UC-5) / Register (UC-2)."),
        ],
        "db": [
            ("doctors", "R", "Get featured/active doctors with specialty info for the home page."),
            ("specialties", "R", "Load specialty list for the search filter."),
            ("departments", "R", "Load department list for the search filter."),
            ("branches", "R", "Get branch locations for the branch section."),
            ("reviews", "R", "Get average rating per doctor displayed on doctor cards."),
        ],
        "queries": """
1/ Get featured active doctors for home page
db.doctors.aggregate([
  { $match: { isActive: true, isFeatured: true } },
  { $lookup: { from: "specialties", localField: "specialtyId", foreignField: "_id", as: "specialty" } },
  { $lookup: { from: "reviews", localField: "_id", foreignField: "doctorId", as: "reviews" } },
  { $addFields: { avgRating: { $avg: "$reviews.rating" } } },
  { $project: { fullName: 1, degree: 1, avatarUrl: 1, "specialty.name": 1, avgRating: 1 } },
  { $limit: 8 }
])

2/ Search doctors by keyword and specialty
db.doctors.find(
  { isActive: true, fullName: { $regex: keyword, $options: "i" }, specialtyId: ObjectId(specialtyId) },
  { fullName: 1, degree: 1, avatarUrl: 1, consultationFee: 1 }
).skip((page-1)*10).limit(10)
""",
    },
    {
        "feature": "Register",
        "uc": "UC-2", "name": "Register Patient Account",
        "description": "Allows Guest to create a Patient account with email, password, full name, and phone. The system validates email uniqueness, creates an inactive account with a linked patient profile, and sends a verification email. The account is activated after the email link is confirmed.",
        "related": "UC-2 \u2013 Register Patient Account (UC-2.1 Verify Email Address, UC-2.2 Resend Verification Email)",
        "image": "UI_Register.png", "caption": "Register Patient Account",
        "fields": [
            ("Full Name", "Text Input", "Guest enters full name (required, 2\u201350 characters)."),
            ("Email", "Text Input", "Guest enters email; validated for format and uniqueness."),
            ("Phone", "Text Input", "Guest enters Vietnamese phone number (10 digits)."),
            ("Password", "Password Input", "Min 8 characters with letters and numbers; strength indicator."),
            ("Confirm Password", "Password Input", "Must match Password."),
            ("Terms checkbox", "Checkbox", "Guest must accept Terms of Service before submitting."),
            ("Register", "Button", "Validates inputs and creates the account, then sends a verification email."),
            ("Login link", "Link", "Navigates to the Login screen (UC-5)."),
        ],
        "db": [
            ("users", "C/R", "Check email uniqueness; insert new user (role = patient, isEmailVerified = false)."),
            ("patients", "C", "Create the linked patient profile."),
            ("emailTokens", "C", "Store the verification token (24h expiry, single-use)."),
        ],
        "queries": """
1/ Check email uniqueness
db.users.findOne({ email: email.toLowerCase() })

2/ Create user and patient profile
db.users.insertOne({
  email: email.toLowerCase(), passwordHash: bcryptHash, fullName, phone,
  role: "patient", isEmailVerified: false, isActive: true, createdAt: new Date()
})
db.patients.insertOne({ userId: newUserId, dateOfBirth: null, gender: null, address: null })

3/ Store verification token
db.emailTokens.insertOne({ userId: newUserId, token: randomToken, type: "verify-email",
  expiresAt: new Date(Date.now() + 24*3600*1000), used: false })
""",
    },
    {
        "feature": "Login",
        "uc": "UC-5", "name": "Login",
        "description": "Allows any registered user (Patient, Doctor, Staff, Admin) to sign in with email and password. On success the system issues a JWT session token and redirects to the dashboard matching the user role. Patient accounts must have a verified email.",
        "related": "UC-5 \u2013 Login",
        "image": "UI_Login.png", "caption": "Login",
        "fields": [
            ("Email", "Text Input", "User enters the registered email."),
            ("Password", "Password Input", "User enters the password; show/hide toggle."),
            ("Remember me", "Checkbox", "Extends token expiry (7 days instead of 1 day)."),
            ("Login", "Button", "Verifies credentials and redirects by role."),
            ("Forgot password?", "Link", "Opens Reset Password flow (UC-4)."),
            ("Register link", "Link", "Navigates to Register (UC-2)."),
        ],
        "db": [
            ("users", "R/U", "Verify credentials and account status; update lastLoginAt."),
            ("sessions", "C", "Store the issued refresh token / session record."),
        ],
        "queries": """
1/ Find account by email
db.users.findOne({ email: email.toLowerCase(), isActive: true })
// then compare bcrypt(password, user.passwordHash) in service layer

2/ Update last login and create session
db.users.updateOne({ _id: user._id }, { $set: { lastLoginAt: new Date() } })
db.sessions.insertOne({ userId: user._id, refreshToken: token,
  expiresAt: rememberMe ? in7days : in1day, createdAt: new Date() })
""",
    },
    {
        "feature": "Forgot Password",
        "uc": "UC-4", "name": "Reset Forgotten Password",
        "description": "Allows a registered user who cannot log in to reset the password. The user submits the registered email, receives a time-limited reset link, and sets a new password. All old sessions are revoked after a successful reset.",
        "related": "UC-4 \u2013 Reset Forgotten Password",
        "image": "UI_ForgotPassword.png", "caption": "Forgot / Reset Password",
        "fields": [
            ("Email", "Text Input", "User enters the registered email to receive the reset link."),
            ("Send Reset Link", "Button", "Sends the tokenized link; always shows a neutral success message."),
            ("New Password", "Password Input", "On the reset screen \u2014 min 8 characters, letters + numbers."),
            ("Confirm Password", "Password Input", "Must match New Password."),
            ("Reset Password", "Button", "Updates the password hash and revokes old sessions."),
        ],
        "db": [
            ("users", "R/U", "Find the account by email; update passwordHash after reset."),
            ("emailTokens", "C/R/U", "Create reset token (30 min, single-use); validate and mark used."),
            ("sessions", "D", "Revoke all sessions of the user after the password change."),
        ],
        "queries": """
1/ Create reset token
db.emailTokens.insertOne({ userId: user._id, token: randomToken, type: "reset-password",
  expiresAt: new Date(Date.now() + 30*60*1000), used: false })

2/ Validate token and reset password
db.emailTokens.findOne({ token, type: "reset-password", used: false, expiresAt: { $gt: new Date() } })
db.users.updateOne({ _id: t.userId }, { $set: { passwordHash: newHash } })
db.emailTokens.updateOne({ _id: t._id }, { $set: { used: true } })
db.sessions.deleteMany({ userId: t.userId })
""",
    },
    {
        "feature": "Logout",
        "uc": "UC-6", "name": "Logout",
        "description": "Allows a logged-in user to end the session securely. The confirmation popup prevents accidental logout. The session token is revoked server-side, client storage is cleared, and the user returns to the Login screen.",
        "related": "UC-6 \u2013 Logout",
        "image": "UI_Logout.png", "caption": "Logout Confirmation (popup)",
        "fields": [
            ("Logout menu item", "Menu Item", "Located in the avatar dropdown; opens the confirmation popup."),
            ("Confirmation message", "Text", "\u201cAre you sure you want to log out?\u201d"),
            ("Log Out", "Button", "Revokes the token and redirects to Login."),
            ("Cancel", "Button", "Closes the popup and stays on the current screen."),
        ],
        "db": [
            ("sessions", "D", "Delete/revoke the current session token."),
        ],
        "queries": """
1/ Revoke the current session
db.sessions.deleteOne({ refreshToken: currentToken })
""",
    },
    {
        "feature": "Manage Profile",
        "uc": "UC-3.1", "name": "Profile Management",
        "description": "Allows the Patient (Doctor inherits) to view and update their own profile: personal info, contact, emergency contact, basic medical info (blood type, allergies), avatar upload, and change password. Email and role are read-only.",
        "related": "UC-3.1 \u2013 Profile Management (UC-3.1.1 View, UC-3.1.2 Update, UC-3.1.3 Avatar, UC-3.1.4 Change Password)",
        "image": "UI_Profile.png", "caption": "My Profile",
        "fields": [
            ("Avatar", "Image Upload", "JPG/PNG up to 5 MB with crop + preview (UC-3.1.3)."),
            ("Full Name", "Text Input", "Editable; required."),
            ("Email", "Text (read-only)", "Account email; cannot be self-changed."),
            ("Date of Birth", "Date Picker", "Must be a date in the past."),
            ("Gender", "Radio/Select", "Male / Female / Other."),
            ("Phone", "Text Input", "10-digit Vietnamese phone number."),
            ("Address", "Text Area", "Patient home address."),
            ("Blood Type / Allergies", "Select / Tags", "Basic medical info shown to doctors during consultation."),
            ("Save", "Button", "Validates and persists the changes."),
            ("Change Password", "Button", "Opens the change-password popup (old + new + confirm)."),
        ],
        "db": [
            ("users", "R/U", "Load account info; update fullName, phone, avatarUrl, passwordHash."),
            ("patients", "R/U", "Load and update patient profile fields (DOB, gender, address, medical info)."),
        ],
        "queries": """
1/ Load own profile
db.users.findOne({ _id: currentUserId }, { passwordHash: 0 })
db.patients.findOne({ userId: currentUserId })

2/ Update profile
db.users.updateOne({ _id: currentUserId }, { $set: { fullName, phone, avatarUrl } })
db.patients.updateOne({ userId: currentUserId },
  { $set: { dateOfBirth, gender, address, bloodType, allergies } })

3/ Change password (after verifying old password)
db.users.updateOne({ _id: currentUserId }, { $set: { passwordHash: newHash } })
""",
    },
    {
        "feature": "Book Appointment",
        "uc": "UC-7", "name": "Book Appointment",
        "description": "Allows the Patient to book an appointment by selecting a doctor, date, and available time slot, then paying the consultation fee via Wallet, PayOS, or Momo. An eligible insurance card discount is applied automatically. The slot is held for 10 minutes during checkout; on successful payment the appointment is confirmed and a confirmation is sent.",
        "related": "UC-7 \u2013 Book Appointment (UC-7.1 Pay via PayOS, UC-7.2 Pay via Momo, UC-7.3 Apply Insurance Discount)",
        "image": "UI_BookAppointment.png", "caption": "Book Appointment",
        "fields": [
            ("Doctor info", "Card", "Shows selected doctor: avatar, name, specialty, consultation fee."),
            ("Date", "Date Picker", "Patient selects the visit date; only days with generated slots are enabled."),
            ("Time Slot", "Slot Grid", "Available slots shown as chips; booked/blocked slots disabled."),
            ("Reason for visit", "Text Area", "Patient describes symptoms or reason (optional, max 500 chars)."),
            ("Insurance Card", "Dropdown", "Valid cards of the patient; discount preview applied to the fee (UC-7.3)."),
            ("Fee summary", "Text", "Consultation fee, discount amount, and final total."),
            ("Payment Method", "Radio Cards", "Wallet / PayOS / Momo (UC-7.1, UC-7.2)."),
            ("Confirm & Pay", "Button", "Holds the slot and processes the payment via the chosen method."),
        ],
        "db": [
            ("doctors", "R", "Get doctor info and consultation fee."),
            ("appointmentSlots", "R/U", "List available slots; hold and mark slot booked."),
            ("appointments", "C", "Create the appointment (status = confirmed) after payment."),
            ("insuranceCards", "R", "Get valid cards of the patient for the discount."),
            ("wallets", "R/U", "Check balance and deduct fee when paying by wallet."),
            ("transactions", "C", "Record the payment transaction (wallet/PayOS/Momo)."),
            ("notifications", "C", "Create the booking confirmation notification."),
        ],
        "queries": """
1/ Get available slots of a doctor on a date
db.appointmentSlots.find({ doctorId: ObjectId(doctorId),
  date: ISODate(date), status: "available" }).sort({ startTime: 1 })

2/ Hold a slot during checkout (atomic)
db.appointmentSlots.updateOne(
  { _id: ObjectId(slotId), status: "available" },
  { $set: { status: "held", heldBy: patientId, heldUntil: new Date(Date.now()+10*60*1000) } })

3/ Confirm appointment after successful payment
db.appointments.insertOne({ patientId, doctorId, slotId, reason,
  fee: finalFee, discount, insuranceCardId, status: "confirmed", createdAt: new Date() })
db.appointmentSlots.updateOne({ _id: slotId }, { $set: { status: "booked" } })
db.transactions.insertOne({ userId: patientId, type: "payment", method, amount: finalFee,
  refType: "appointment", refId: appointmentId, status: "success", createdAt: new Date() })
""",
    },
    {
        "feature": "Manage Appointments",
        "uc": "UC-8", "name": "Manage Appointments",
        "description": "Allows the Patient to view upcoming and past appointments, open an appointment detail, reschedule to another available slot or cancel (at least 24 hours before start, refund to wallet), and rate the doctor after a completed visit.",
        "related": "UC-8 \u2013 View Appointment List (UC-8.1 Detail, UC-8.1.1 Reschedule, UC-8.1.2 Cancel, UC-8.1.3 Rate Doctor)",
        "image": "UI_MyAppointments.png", "caption": "My Appointments",
        "fields": [
            ("Tab filter", "Tabs", "Upcoming / Past; default Upcoming."),
            ("Appointment card", "Card List", "Doctor, specialty, date-time, room, status badge (confirmed/completed/cancelled)."),
            ("View Detail", "Button", "Opens the appointment detail screen (UC-8.1)."),
            ("Reschedule", "Button", "Opens slot picker for a new time; enabled \u2265 24h before start (UC-8.1.1)."),
            ("Cancel", "Button", "Confirmation popup; refund to wallet per policy (UC-8.1.2)."),
            ("Rate Doctor", "Stars + Text Area", "1\u20135 stars and comment; shown only for completed visits not yet rated (UC-8.1.3)."),
        ],
        "db": [
            ("appointments", "R/U", "List/detail appointments; update status on reschedule/cancel."),
            ("appointmentSlots", "R/U", "Release the old slot and book the new one on reschedule."),
            ("wallets", "U", "Credit the refund on cancellation."),
            ("transactions", "C", "Record the refund transaction."),
            ("reviews", "C", "Create the doctor review after a completed visit."),
        ],
        "queries": """
1/ List my appointments
db.appointments.find({ patientId: currentPatientId })
  .sort({ startTime: -1 }).skip((page-1)*10).limit(10)

2/ Cancel with refund (>= 24h before start)
db.appointments.updateOne(
  { _id: apptId, patientId: currentPatientId, status: "confirmed",
    startTime: { $gt: new Date(Date.now() + 24*3600*1000) } },
  { $set: { status: "cancelled", cancelledAt: new Date() } })
db.appointmentSlots.updateOne({ _id: appt.slotId }, { $set: { status: "available" } })
db.wallets.updateOne({ userId: currentUserId }, { $inc: { balance: refundAmount } })

3/ Rate doctor after completed visit
db.reviews.insertOne({ appointmentId, doctorId, patientId,
  rating, comment, createdAt: new Date() })
""",
    },
    {
        "feature": "View Prescriptions",
        "uc": "UC-9", "name": "View Prescription History",
        "description": "Allows the Patient to view issued prescriptions, open a prescription detail with medicine lines, dosage, and instructions, display the QR code for pharmacy pickup, and export the prescription as a PDF file.",
        "related": "UC-9 \u2013 View Prescription History (UC-9.1 Detail, UC-9.1.1 QR Code, UC-9.1.2 Export PDF)",
        "image": "UI_Prescriptions.png", "caption": "My Prescriptions",
        "fields": [
            ("Prescription list", "Card/Table List", "Date, prescribing doctor, diagnosis summary, status."),
            ("Detail view", "Panel", "Medicine lines: name, quantity, dosage, frequency, duration, notes."),
            ("QR Code", "Button + Popup", "Displays a scannable QR encoding the prescription ID (UC-9.1.1)."),
            ("Export PDF", "Button", "Generates and downloads the prescription PDF (UC-9.1.2)."),
        ],
        "db": [
            ("prescriptions", "R", "List/detail issued prescriptions of the patient."),
            ("prescriptionItems", "R", "Get medicine lines of one prescription."),
            ("medicines", "R", "Resolve medicine names/units for the lines."),
            ("doctors", "R", "Resolve the prescribing doctor name."),
        ],
        "queries": """
1/ List my prescriptions
db.prescriptions.find({ patientId: currentPatientId, status: "issued" })
  .sort({ issuedAt: -1 }).limit(10)

2/ Get prescription detail with lines
db.prescriptionItems.aggregate([
  { $match: { prescriptionId: ObjectId(id) } },
  { $lookup: { from: "medicines", localField: "medicineId",
               foreignField: "_id", as: "medicine" } },
  { $project: { "medicine.name": 1, "medicine.unit": 1,
                quantity: 1, dosage: 1, frequency: 1, duration: 1, note: 1 } }
])
""",
    },
    {
        "feature": "Notifications",
        "uc": "UC-10", "name": "View Notification Inbox",
        "description": "Allows the Patient to read system notifications (appointment reminders, results ready, booking confirmations) in one inbox, mark items as read individually or all at once, and configure which notification types to receive.",
        "related": "UC-10 \u2013 View Notification Inbox (UC-10.1 Preferences, UC-10.2 Mark as Read)",
        "image": "UI_Notifications.png", "caption": "Notification Inbox",
        "fields": [
            ("Bell icon", "Icon + Badge", "Unread count badge (99+ max); opens the inbox dropdown."),
            ("Notification item", "List Item", "Icon by type, title, snippet, relative time; unread highlighted."),
            ("Mark all as read", "Button", "Marks all unread notifications read (UC-10.2)."),
            ("Preferences", "Toggle List", "Per-type switches: reminders, results, booking confirmations (UC-10.1)."),
        ],
        "db": [
            ("notifications", "R/U", "List notifications of the user; set isRead flags."),
            ("notificationPreferences", "R/U", "Load and update per-type preferences."),
        ],
        "queries": """
1/ List my notifications with unread count
db.notifications.find({ userId: currentUserId })
  .sort({ createdAt: -1 }).limit(20)
db.notifications.countDocuments({ userId: currentUserId, isRead: false })

2/ Mark all as read
db.notifications.updateMany({ userId: currentUserId, isRead: false },
  { $set: { isRead: true, readAt: new Date() } })

3/ Update preferences
db.notificationPreferences.updateOne({ userId: currentUserId },
  { $set: { reminder: true, resultsReady: true, bookingEmail: false } }, { upsert: true })
""",
    },
    {
        "feature": "Complaints (Patient)",
        "uc": "UC-11", "name": "Submit Complaint",
        "description": "Allows the Patient to submit a service complaint with category, subject, description, and optional image attachment; track submitted complaints in a list; and read staff replies and status updates in the complaint detail.",
        "related": "UC-11 \u2013 Submit Complaint (UC-11.1 My Complaints List, UC-11.1.1 Complaint Detail)",
        "image": "UI_Complaints.png", "caption": "Submit Complaint & My Complaints",
        "fields": [
            ("Category", "Dropdown", "Predefined categories: Service, Billing, Facility, Doctor, Other."),
            ("Subject", "Text Input", "Short summary (required, max 120 chars)."),
            ("Description", "Text Area", "Detailed description (required)."),
            ("Attachment", "Image Upload", "Optional JPG/PNG up to 5 MB."),
            ("Submit", "Button", "Creates the complaint with status Open."),
            ("My complaints list", "Table/List", "Subject, category, created date, status badge (Open/In Progress/Resolved)."),
            ("Detail thread", "Timeline", "Patient message and staff replies in chronological order."),
        ],
        "db": [
            ("complaints", "C/R", "Create the complaint; list/detail own complaints."),
            ("complaintReplies", "R", "Read staff replies of one complaint."),
            ("notifications", "C", "Notify staff about the new complaint."),
        ],
        "queries": """
1/ Create complaint
db.complaints.insertOne({ patientId: currentPatientId, category, subject,
  description, attachmentUrl, status: "open", createdAt: new Date() })

2/ List my complaints
db.complaints.find({ patientId: currentPatientId }).sort({ createdAt: -1 })

3/ Complaint detail with replies
db.complaintReplies.find({ complaintId: ObjectId(id) }).sort({ createdAt: 1 })
""",
    },
    {
        "feature": "Wallet",
        "uc": "UC-12", "name": "View Wallet Balance",
        "description": "Allows the Patient to view the medical wallet balance, top up via PayOS or Momo, and review the full transaction history (top-ups, payments, refunds) with filters by type and date range.",
        "related": "UC-12 \u2013 View Wallet Balance (UC-12.1 Transaction History)",
        "image": "UI_Wallet.png", "caption": "My Wallet",
        "fields": [
            ("Balance card", "Card", "Current balance in VND with gradient brand styling."),
            ("Top Up", "Button + Popup", "Amount input (min 50,000 VND) and gateway choice PayOS/Momo."),
            ("Quick amounts", "Chip Buttons", "100k / 200k / 500k / 1M shortcuts."),
            ("Transaction list", "Table", "Date, type (top-up/payment/refund), description, amount (+/-), status."),
            ("Filter", "Dropdown + Date Range", "Filter by type and date range (UC-12.1)."),
        ],
        "db": [
            ("wallets", "R/U", "Get balance; credit on successful top-up."),
            ("transactions", "C/R", "Create top-up transaction; list history with filters."),
        ],
        "queries": """
1/ Get my wallet
db.wallets.findOne({ userId: currentUserId })

2/ Credit wallet after gateway callback verified
db.transactions.updateOne({ _id: txId, status: "pending" },
  { $set: { status: "success", completedAt: new Date() } })
db.wallets.updateOne({ userId: tx.userId }, { $inc: { balance: tx.amount } })

3/ Transaction history with filter
db.transactions.find({ userId: currentUserId, type: "topup",
  createdAt: { $gte: ISODate(from), $lte: ISODate(to) } })
  .sort({ createdAt: -1 }).skip((page-1)*10).limit(10)
""",
    },
    {
        "feature": "Insurance",
        "uc": "UC-13", "name": "Manage Insurance Cards",
        "description": "Allows the Patient to view registered insurance cards with validity badges, add a new card manually or by scanning the card image (OCR auto-fill), update card information, and delete a card. Valid cards yield a discount when booking.",
        "related": "UC-13 \u2013 View Insurance Card List (UC-13.1 Delete, UC-13.2 Add, UC-13.2.1 Scan OCR, UC-13.3 Update)",
        "image": "UI_Insurance.png", "caption": "Insurance Cards",
        "fields": [
            ("Card list", "Card Grid", "Provider logo, masked card number, holder name, expiry, validity badge."),
            ("Add Card", "Button + Modal", "Opens the add-card form (UC-13.2)."),
            ("Scan card image", "Image Upload", "OCR auto-fills provider, number, holder, expiry (UC-13.2.1)."),
            ("Provider", "Dropdown", "Supported insurance providers maintained by Admin."),
            ("Card Number", "Text Input", "Validated format per provider."),
            ("Holder Name / Expiry", "Text / Date Picker", "Card holder and expiry date; expired cards flagged."),
            ("Save / Delete", "Button", "Persists the card / removes it after confirmation (UC-13.3, UC-13.1)."),
        ],
        "db": [
            ("insuranceCards", "C/R/U/D", "Full CRUD on the patient's own insurance cards."),
            ("insuranceProviders", "R", "Load the supported provider list."),
        ],
        "queries": """
1/ List my cards
db.insuranceCards.find({ patientId: currentPatientId }).sort({ createdAt: -1 })

2/ Add card (after OCR or manual input)
db.insuranceCards.insertOne({ patientId: currentPatientId, providerId,
  cardNumber, holderName, expiryDate: ISODate(expiry),
  imageUrl, createdAt: new Date() })

3/ Delete card
db.insuranceCards.deleteOne({ _id: ObjectId(cardId), patientId: currentPatientId })
""",
    },
    {
        "feature": "Queue Status",
        "uc": "UC-14", "name": "View Queue Status",
        "description": "Allows the Patient to see their queue ticket number, the number currently being served, and how many people are ahead for the clinic room of their appointment. The screen auto-refreshes while the queue session is active and highlights when it is the patient's turn.",
        "related": "UC-14 \u2013 View Queue Status",
        "image": "UI_QueueStatus.png", "caption": "My Queue Status",
        "fields": [
            ("My ticket number", "Big Number", "The patient's issued ticket number (e.g., A-027)."),
            ("Now serving", "Number Display", "Ticket currently called in the room; live updated."),
            ("People ahead", "Counter", "Number of waiting tickets before the patient."),
            ("Room / Doctor info", "Text", "Clinic room code and doctor name of the appointment."),
            ("Status banner", "Banner", "\u201cIt's your turn \u2014 Room 203\u201d highlighted when called."),
        ],
        "db": [
            ("queueTickets", "R", "Get the patient's active ticket and its position."),
            ("queueSessions", "R", "Get the session state and current called number."),
        ],
        "queries": """
1/ Get my active ticket today
db.queueTickets.findOne({ patientId: currentPatientId,
  status: { $in: ["waiting", "called"] },
  createdAt: { $gte: startOfToday } })

2/ Count people ahead
db.queueTickets.countDocuments({ sessionId: ticket.sessionId,
  status: "waiting", number: { $lt: ticket.number } })
""",
    },
    {
        "feature": "Favorites",
        "uc": "UC-15", "name": "Manage Favorite Doctors",
        "description": "Allows the Patient to view saved favorite doctors in a grid with quick booking access, add a doctor to favorites from any doctor card or profile via the heart icon, and remove a doctor from the favorites list.",
        "related": "UC-15 \u2013 View Favorite Doctors (UC-15.1 Add, UC-15.2 Remove)",
        "image": "UI_Favorites.png", "caption": "Favorite Doctors",
        "fields": [
            ("Favorites grid", "Card Grid", "Doctor avatar, name, specialty, rating, and quick \u201cBook\u201d button."),
            ("Heart icon", "Toggle Icon", "Filled = favorited; click toggles add/remove (UC-15.1, UC-15.2)."),
            ("Book", "Button", "Jumps directly to Book Appointment (UC-7) with the doctor preselected."),
            ("Empty state", "Illustration + Link", "Suggests searching doctors when the list is empty."),
        ],
        "db": [
            ("favorites", "C/R/D", "Add, list, and remove favorite doctors of the patient."),
            ("doctors", "R", "Resolve doctor info for the favorites grid (active only)."),
        ],
        "queries": """
1/ List my favorite doctors
db.favorites.aggregate([
  { $match: { patientId: currentPatientId } },
  { $lookup: { from: "doctors", localField: "doctorId", foreignField: "_id", as: "doctor" } },
  { $unwind: "$doctor" },
  { $match: { "doctor.isActive": true } },
  { $project: { "doctor.fullName": 1, "doctor.avatarUrl": 1, "doctor.specialtyId": 1 } }
])

2/ Toggle favorite
db.favorites.insertOne({ patientId: currentPatientId, doctorId, createdAt: new Date() })
db.favorites.deleteOne({ patientId: currentPatientId, doctorId: ObjectId(doctorId) })
""",
    },
    {
        "feature": "Doctor Dashboard",
        "uc": "UC-16", "name": "View Doctor Dashboard",
        "description": "Allows the Doctor to view an overview screen after login: today's appointment count, active queue session status, pending encounter sign-offs, and shortcuts to Today Appointments, Queue Session, Schedule, and EMR. The Doctor can also open their own professional profile from here.",
        "related": "UC-16 \u2013 View Doctor Dashboard (UC-16.1 Profile Management)",
        "image": "UI_DoctorDashboard.png", "caption": "Doctor Dashboard",
        "fields": [
            ("KPI widgets", "Stat Cards", "Today's appointments, patients seen, waiting tickets, pending sign-offs."),
            ("Today schedule", "Timeline List", "Today's appointments with check-in status and \u201cStart\u201d shortcuts (UC-17)."),
            ("Queue widget", "Card", "Active session state, now serving number, quick \u201cOpen Queue\u201d (UC-18)."),
            ("Pending sign-offs", "List", "Encounters not yet signed off with quick links."),
            ("My Profile", "Menu Item", "Opens professional profile editing (UC-16.1)."),
        ],
        "db": [
            ("appointments", "R", "Count and list today's appointments of the doctor."),
            ("queueSessions", "R", "Get the doctor's active queue session state."),
            ("encounters", "R", "Count encounters with status in-progress (pending sign-off)."),
            ("doctors", "R/U", "Load and update the doctor's own profile."),
        ],
        "queries": """
1/ Today's appointments of the doctor
db.appointments.find({ doctorId: currentDoctorId,
  startTime: { $gte: startOfToday, $lt: endOfToday } }).sort({ startTime: 1 })

2/ Pending sign-offs
db.encounters.countDocuments({ doctorId: currentDoctorId, status: "in-progress" })
""",
    },
    {
        "feature": "Consultation & EMR",
        "uc": "UC-17", "name": "Start Consultation",
        "description": "Allows the Doctor to view today's appointment list and start a clinical session for a checked-in patient. Starting a consultation creates an encounter record. Inside the encounter workspace the Doctor writes clinical notes, records ICD-10 diagnoses with catalog search, and signs off the encounter to finalize it.",
        "related": "UC-17 \u2013 Start Consultation (UC-17.1.1 Create Encounter, UC-17.1.1.1 Detail/Notes/Sign Off, UC-17.1.1.2 Record Diagnosis + ICD-10 Search)",
        "image": "UI_Consultation.png", "caption": "Consultation Workspace (Encounter)",
        "fields": [
            ("Today list", "Table", "Appointments with time, patient, check-in status, and \u201cStart Consultation\u201d button."),
            ("Patient header", "Card", "Patient name, age, gender, blood type, allergies warning chip."),
            ("Chief complaint / Notes", "Text Area", "Clinical notes; autosaved as draft."),
            ("Diagnosis search", "Autocomplete", "Searches the ICD-10 catalog by code or name (UC-17.1.1.2.2)."),
            ("Diagnosis list", "Tag/Table", "Added ICD-10 codes with type (primary/secondary) and remove icon."),
            ("Imaging / Prescription tabs", "Tabs", "Links to Medical Imaging (UC-23) and Prescription (UC-24)."),
            ("Sign Off", "Button", "Validates required fields and locks the encounter (UC-17.1.1.1.2)."),
        ],
        "db": [
            ("appointments", "R/U", "List today's checked-in appointments; mark completed after sign-off."),
            ("encounters", "C/R/U", "Create encounter on start; update notes; set signed-off status."),
            ("diagnoses", "C/R/U/D", "Add, update, and remove ICD-10 diagnosis entries."),
            ("icd10Catalog", "R", "Search diagnosis codes by keyword."),
            ("notifications", "C", "Notify the patient that results are ready."),
        ],
        "queries": """
1/ Create encounter when starting consultation (idempotent)
db.encounters.updateOne(
  { appointmentId: ObjectId(apptId) },
  { $setOnInsert: { patientId, doctorId, status: "in-progress",
      notes: "", startedAt: new Date() } },
  { upsert: true })

2/ Search ICD-10 catalog
db.icd10Catalog.find({ $or: [
    { code: { $regex: "^" + keyword, $options: "i" } },
    { name: { $regex: keyword, $options: "i" } } ] }).limit(20)

3/ Sign off encounter (requires >= 1 diagnosis)
db.encounters.updateOne(
  { _id: encounterId, doctorId: currentDoctorId, status: "in-progress" },
  { $set: { status: "signed-off", signedOffAt: new Date() } })
db.appointments.updateOne({ _id: enc.appointmentId }, { $set: { status: "completed" } })
""",
    },
    {
        "feature": "Queue Session",
        "uc": "UC-18", "name": "Manage Queue Session",
        "description": "Allows the Doctor to open a queue session for the assigned clinic room, call the next waiting ticket, recall a previously called ticket when the patient is absent, and close the session at the end of the shift. Patient queue screens update live.",
        "related": "UC-18 \u2013 Open Queue Session (UC-18.1 Call Next, UC-18.2 Recall, UC-18.3 Close Session)",
        "image": "UI_QueueSession.png", "caption": "Doctor Queue Session",
        "fields": [
            ("Session header", "Card", "Room code, session status (Active), opened time, total tickets."),
            ("Now serving", "Big Number", "Currently called ticket number, prominent display."),
            ("Call Next", "Button", "Calls the next waiting ticket (UC-18.1); disabled when queue empty."),
            ("Recall", "Button", "Re-announces a previously called ticket (UC-18.2)."),
            ("Waiting list", "Table", "Waiting tickets with number, patient, issue time, status."),
            ("Mark No-show", "Button", "Marks an absent called ticket as no-show."),
            ("Close Session", "Button", "Confirmation popup; finalizes remaining tickets (UC-18.3)."),
        ],
        "db": [
            ("queueSessions", "C/R/U", "Open, read, and close the room's queue session."),
            ("queueTickets", "R/U", "Call next ticket, recall, mark no-show."),
        ],
        "queries": """
1/ Open session (only one active per room)
db.queueSessions.updateOne(
  { roomId: ObjectId(roomId), status: "active" },
  { $setOnInsert: { doctorId: currentDoctorId, openedAt: new Date(),
      status: "active", currentNumber: 0 } },
  { upsert: true })

2/ Call next ticket (atomic)
db.queueTickets.findOneAndUpdate(
  { sessionId, status: "waiting" },
  { $set: { status: "called", calledAt: new Date() } },
  { sort: { number: 1 }, returnDocument: "after" })
db.queueSessions.updateOne({ _id: sessionId }, { $set: { currentNumber: ticket.number } })

3/ Close session
db.queueSessions.updateOne({ _id: sessionId }, { $set: { status: "closed", closedAt: new Date() } })
db.queueTickets.updateMany({ sessionId, status: "waiting" }, { $set: { status: "no-show" } })
""",
    },
    {
        "feature": "Work Shift Schedule",
        "uc": "UC-19", "name": "View Work Shift Schedule",
        "description": "Allows the Doctor to view their assigned work shifts in a weekly or monthly calendar, including shift time, clinic room, and status. Shifts are created by Admin and drive appointment slot generation; the Doctor has read-only access.",
        "related": "UC-19 \u2013 View Work Shift Schedule",
        "image": "UI_WorkShifts.png", "caption": "My Work Shifts",
        "fields": [
            ("Week/Month switch", "Toggle", "Switches the calendar view granularity."),
            ("Shift block", "Calendar Event", "Shift time range, room code, color by shift type (morning/afternoon)."),
            ("Shift detail", "Popup", "Room, time, slot duration, number of generated slots."),
            ("Navigation", "Buttons", "Previous / Today / Next period."),
        ],
        "db": [
            ("workShifts", "R", "List shifts of the doctor in the selected period."),
            ("clinicRooms", "R", "Resolve room codes for display."),
            ("appointmentSlots", "R", "Count generated slots per shift for the detail popup."),
        ],
        "queries": """
1/ Shifts of the doctor in a week
db.workShifts.find({ doctorId: currentDoctorId,
  date: { $gte: ISODate(weekStart), $lte: ISODate(weekEnd) } }).sort({ date: 1, startTime: 1 })

2/ Slots generated for one shift
db.appointmentSlots.countDocuments({ shiftId: ObjectId(shiftId) })
""",
    },
    {
        "feature": "Appointment Calendar",
        "uc": "UC-20", "name": "View Appointment Calendar",
        "description": "Allows the Doctor to view the personal appointment calendar with booked, available, and blocked slots; block unavailable time slots so patients cannot book them; and unblock previously blocked slots.",
        "related": "UC-20 \u2013 View Appointment Calendar (UC-20.1 Block Slot, UC-20.2 Unblock Slot)",
        "image": "UI_DoctorCalendar.png", "caption": "Doctor Appointment Calendar",
        "fields": [
            ("Day/Week view", "Calendar Grid", "Slots colored: booked (cyan), available (white), blocked (gray hatch)."),
            ("Slot popup", "Popup", "Booked: patient + reason; Available: \u201cBlock\u201d action; Blocked: \u201cUnblock\u201d action."),
            ("Block reason", "Text Input", "Optional reason saved with the block (UC-20.1)."),
            ("Legend", "Legend", "Color meanings for slot states."),
        ],
        "db": [
            ("appointmentSlots", "R/U", "List slots; set status blocked/available (only if not booked)."),
            ("appointments", "R", "Show patient info for booked slots."),
        ],
        "queries": """
1/ Slots of the doctor for a day
db.appointmentSlots.find({ doctorId: currentDoctorId, date: ISODate(day) })
  .sort({ startTime: 1 })

2/ Block a free slot (atomic guard)
db.appointmentSlots.updateOne(
  { _id: ObjectId(slotId), doctorId: currentDoctorId, status: "available" },
  { $set: { status: "blocked", blockReason: reason } })

3/ Unblock
db.appointmentSlots.updateOne(
  { _id: ObjectId(slotId), status: "blocked" }, { $set: { status: "available" },
  $unset: { blockReason: "" } })
""",
    },
    {
        "feature": "Patient EMR Timeline",
        "uc": "UC-21", "name": "View Patient EMR Timeline",
        "description": "Allows the Doctor to review a patient's medical history in a chronological timeline: past encounters, diagnoses, prescriptions, and medical images grouped by visit date. Each item links to its read-only detail. Access is limited to patients having an encounter with the doctor.",
        "related": "UC-21 \u2013 View Patient EMR Timeline",
        "image": "UI_EMRTimeline.png", "caption": "Patient EMR Timeline",
        "fields": [
            ("Patient summary", "Header Card", "Name, age, gender, blood type, allergies, chronic conditions."),
            ("Type filter", "Chips", "All / Encounters / Diagnoses / Prescriptions / Images."),
            ("Timeline item", "Card", "Visit date, doctor, diagnosis summary, links to detail."),
            ("Detail drawer", "Side Panel", "Read-only encounter notes, diagnosis list, prescription lines, image thumbnails."),
        ],
        "db": [
            ("encounters", "R", "Signed-off encounters of the patient, newest first."),
            ("diagnoses", "R", "Diagnoses per encounter."),
            ("prescriptions", "R", "Prescriptions linked to encounters."),
            ("medicalImages", "R", "Image thumbnails per encounter."),
        ],
        "queries": """
1/ EMR timeline of a patient
db.encounters.aggregate([
  { $match: { patientId: ObjectId(patientId), status: "signed-off" } },
  { $sort: { startedAt: -1 } },
  { $lookup: { from: "diagnoses", localField: "_id", foreignField: "encounterId", as: "diagnoses" } },
  { $lookup: { from: "prescriptions", localField: "_id", foreignField: "encounterId", as: "prescriptions" } },
  { $lookup: { from: "medicalImages", localField: "_id", foreignField: "encounterId", as: "images" } }
])
""",
    },
    {
        "feature": "Medical Imaging",
        "uc": "UC-23", "name": "Manage Medical Images",
        "description": "Allows the Doctor to upload diagnostic images (JPG/PNG/DICOM up to 20 MB) for a patient encounter, browse the imaging gallery with thumbnails, view an image fullscreen with zoom, and delete a wrongly uploaded image before the encounter is signed off.",
        "related": "UC-23 \u2013 Upload Medical Image (UC-23.1 Gallery, UC-23.1.1 Fullscreen, UC-23.1.2 Delete)",
        "image": "UI_MedicalImaging.png", "caption": "Medical Imaging Gallery",
        "fields": [
            ("Upload dropzone", "File Upload", "Drag & drop multiple files; type/size validated per file."),
            ("Modality / Description", "Select / Text", "X-Ray, CT, MRI, Ultrasound, Other + free-text description."),
            ("Gallery grid", "Thumbnail Grid", "12 per page with modality badge and upload date."),
            ("Fullscreen viewer", "Modal", "Zoom/pan, next/previous navigation (UC-23.1.1)."),
            ("Delete", "Icon Button", "Confirmation popup; only before encounter sign-off (UC-23.1.2)."),
        ],
        "db": [
            ("medicalImages", "C/R/D", "Insert metadata after cloud upload; list gallery; delete before sign-off."),
            ("encounters", "R", "Verify the encounter is open (not signed off) for upload/delete."),
        ],
        "queries": """
1/ Save image metadata after upload
db.medicalImages.insertOne({ encounterId, patientId, doctorId: currentDoctorId,
  url, modality, description, sizeBytes, uploadedAt: new Date() })

2/ Gallery of an encounter
db.medicalImages.find({ encounterId: ObjectId(encId) })
  .sort({ uploadedAt: -1 }).skip((page-1)*12).limit(12)

3/ Delete (guard: encounter still open)
db.medicalImages.deleteOne({ _id: ObjectId(imageId), encounterId: { $in: openEncounterIds } })
""",
    },
    {
        "feature": "Prescription Management",
        "uc": "UC-24", "name": "Create Prescription",
        "description": "Allows the Doctor to create a prescription inside the encounter: search medicines from the formulary, add line items with quantity, dosage, frequency, duration, and notes; edit or remove lines; and save & issue the prescription. After issuing, the detail view offers PDF export and QR code display.",
        "related": "UC-24 \u2013 Create Prescription (UC-24.1\u201324.4 Lines & Save, UC-24.5 Detail, UC-24.5.1 PDF, UC-24.5.2 QR)",
        "image": "UI_CreatePrescription.png", "caption": "Create Prescription",
        "fields": [
            ("Medicine search", "Autocomplete", "Searches formulary by name; shows unit, stock, and price."),
            ("Line item row", "Table Row", "Medicine, quantity (stepper), dosage, frequency, duration, note, remove icon."),
            ("Stock warning", "Inline Badge", "\u201cLow stock\u201d / \u201cOut of stock \u2014 external purchase\u201d warnings."),
            ("Save & Issue", "Button", "Validates lines and issues the prescription (UC-24.4)."),
            ("Detail actions", "Buttons", "Export PDF / View QR Code after issuing (UC-24.5.x)."),
        ],
        "db": [
            ("medicines", "R", "Search formulary and check stock levels."),
            ("prescriptions", "C/R/U", "Create draft, update, and issue the prescription."),
            ("prescriptionItems", "C/U/D", "Add, edit, and remove line items."),
            ("encounters", "R", "Link the prescription to the open encounter."),
        ],
        "queries": """
1/ Search medicines in formulary
db.medicines.find({ name: { $regex: keyword, $options: "i" }, isActive: true },
  { name: 1, unit: 1, price: 1, stockQuantity: 1 }).limit(10)

2/ Add line item
db.prescriptionItems.insertOne({ prescriptionId, medicineId,
  quantity, dosage, frequency, duration, note })

3/ Issue prescription
db.prescriptions.updateOne(
  { _id: prescriptionId, doctorId: currentDoctorId, status: "draft" },
  { $set: { status: "issued", issuedAt: new Date() } })
""",
    },
    {
        "feature": "Manage Accounts",
        "uc": "UC-25", "name": "Manage Accounts",
        "description": "Allows the Admin to view the accounts list with role/status filters and search, open an account detail, create staff accounts (initial credentials emailed), update account information, change user roles, and deactivate/restore accounts. Deactivating revokes all active sessions; Admin cannot deactivate their own account.",
        "related": "UC-25 \u2013 Manage Accounts",
        "image": "UI_ManageAccounts.png", "caption": "Manage Accounts (Admin)",
        "fields": [
            ("Search / Filters", "Input + Dropdowns", "Search by name/email; filter by role (patient/doctor/staff/admin) and status."),
            ("Accounts table", "Table", "Avatar, name, email, role badge, status badge, last login, actions."),
            ("Create Account", "Button + Modal", "Email, full name, role = staff; sends initial credentials email."),
            ("Edit", "Icon Button", "Opens edit form for account fields."),
            ("Change Role", "Dropdown + Confirm", "Changes role with confirmation popup."),
            ("Deactivate / Restore", "Toggle + Confirm", "Toggles account status; blocked for own account."),
        ],
        "db": [
            ("users", "C/R/U", "List/filter accounts; create staff; update info, role, and status."),
            ("sessions", "D", "Revoke all sessions when an account is deactivated."),
            ("auditLogs", "C", "Record account management actions."),
        ],
        "queries": """
1/ List accounts with filters
db.users.find({ role: "staff", isActive: true,
  $or: [ { fullName: { $regex: kw, $options: "i" } }, { email: { $regex: kw, $options: "i" } } ] })
  .sort({ createdAt: -1 }).skip((page-1)*10).limit(10)

2/ Deactivate account (not self) and revoke sessions
db.users.updateOne({ _id: ObjectId(userId), _id: { $ne: currentAdminId } },
  { $set: { isActive: false, deactivatedAt: new Date() } })
db.sessions.deleteMany({ userId: ObjectId(userId) })

3/ Change role
db.users.updateOne({ _id: ObjectId(userId) }, { $set: { role: newRole } })
db.auditLogs.insertOne({ actorId: currentAdminId, action: "change-role",
  targetId: userId, detail: { from: oldRole, to: newRole }, at: new Date() })
""",
    },
    {
        "feature": "Manage Specialties",
        "uc": "UC-26", "name": "Manage Specialties",
        "description": "Allows the Admin to create, view, update, and delete medical specialties (name, description, icon). Specialties are referenced by doctor profiles and used as public search filters; a specialty referenced by doctors cannot be hard-deleted and is deactivated instead.",
        "related": "UC-26 \u2013 Manage Specialties",
        "image": "UI_ManageSpecialties.png", "caption": "Manage Specialties (Admin)",
        "fields": [
            ("Specialties table", "Table", "Icon, name, description, linked doctors count, status, actions."),
            ("Create / Edit", "Modal Form", "Name (unique), description, icon picker."),
            ("Delete", "Icon Button + Confirm", "Hard delete only when no doctor references it; otherwise deactivate."),
            ("Search", "Text Input", "Filter the list by name."),
        ],
        "db": [
            ("specialties", "C/R/U/D", "Full CRUD on the specialties catalog."),
            ("doctors", "R", "Count references before delete (guard)."),
        ],
        "queries": """
1/ Create specialty (unique name)
db.specialties.insertOne({ name, description, icon, isActive: true, createdAt: new Date() })

2/ Guard before delete
db.doctors.countDocuments({ specialtyId: ObjectId(id) })
// > 0 -> deactivate instead:
db.specialties.updateOne({ _id: ObjectId(id) }, { $set: { isActive: false } })
""",
    },
    {
        "feature": "Manage Departments",
        "uc": "UC-27", "name": "Manage Departments",
        "description": "Allows the Admin to create departments (name, description, location, head doctor), view the list and detail with linked rooms and doctors, update information, and deactivate departments no longer in operation after reassigning their resources.",
        "related": "UC-27 \u2013 Manage Departments",
        "image": "UI_ManageDepartments.png", "caption": "Manage Departments (Admin)",
        "fields": [
            ("Departments table", "Table", "Name, location, head doctor, rooms count, doctors count, status, actions."),
            ("Create / Edit", "Modal Form", "Name (unique), description, location/floor, head doctor dropdown."),
            ("Detail panel", "Side Panel", "Linked clinic rooms and doctors of the department."),
            ("Deactivate", "Button + Confirm", "Requires reassignment of doctors/rooms first."),
        ],
        "db": [
            ("departments", "C/R/U", "Create, list, update, and deactivate departments."),
            ("clinicRooms", "R", "List rooms of one department."),
            ("doctors", "R", "List doctors assigned to the department."),
        ],
        "queries": """
1/ Department detail with linked resources
db.departments.findOne({ _id: ObjectId(id) })
db.clinicRooms.find({ departmentId: ObjectId(id) })
db.doctors.find({ departmentId: ObjectId(id), isActive: true })

2/ Deactivate after reassignment check
db.doctors.countDocuments({ departmentId: ObjectId(id), isActive: true })  // must be 0
db.departments.updateOne({ _id: ObjectId(id) }, { $set: { isActive: false } })
""",
    },
    {
        "feature": "Manage Clinic Rooms",
        "uc": "UC-28", "name": "Manage Clinic Rooms",
        "description": "Allows the Admin to create clinic rooms (unique room code, name, department, floor), view the rooms list filtered by department, and update room information or set a room inactive. Rooms are used in doctor work shifts and queue sessions; room codes are printed on queue tickets.",
        "related": "UC-28 \u2013 Manage Clinic Rooms",
        "image": "UI_ManageRooms.png", "caption": "Manage Clinic Rooms (Admin)",
        "fields": [
            ("Rooms table", "Table", "Room code, name, department, floor, status, actions."),
            ("Department filter", "Dropdown", "Filters rooms by department."),
            ("Create / Edit", "Modal Form", "Room code (unique), name, department dropdown, floor."),
            ("Deactivate", "Toggle + Confirm", "Blocked while the room has scheduled shifts."),
        ],
        "db": [
            ("clinicRooms", "C/R/U", "Create, list, update, deactivate rooms."),
            ("departments", "R", "Load department dropdown."),
            ("workShifts", "R", "Guard: check scheduled shifts before deactivation."),
        ],
        "queries": """
1/ Create room (unique code)
db.clinicRooms.insertOne({ code, name, departmentId, floor, isActive: true })

2/ Guard before deactivation
db.workShifts.countDocuments({ roomId: ObjectId(id), date: { $gte: new Date() } })  // must be 0
db.clinicRooms.updateOne({ _id: ObjectId(id) }, { $set: { isActive: false } })
""",
    },
    {
        "feature": "Manage Doctors",
        "uc": "UC-29", "name": "Manage Doctors",
        "description": "Allows the Admin to create doctor profiles with linked accounts, view the doctors list with specialty/department filters, update profiles, deactivate doctors, import doctors from an Excel template (invalid rows reported), export the filtered list to Excel, and create work shifts with appointment slot generation.",
        "related": "UC-29 \u2013 Manage Doctors",
        "image": "UI_ManageDoctors.png", "caption": "Manage Doctors (Admin)",
        "fields": [
            ("Doctors table", "Table", "Avatar, name, specialty, department, fee, status, actions."),
            ("Filters", "Dropdowns", "Specialty, department, status."),
            ("Create Doctor", "Modal Form", "Profile + account creation; credentials emailed."),
            ("Import Excel", "Button + Modal", "Uploads .xlsx template; shows valid/invalid row report."),
            ("Export Excel", "Button", "Downloads the current filtered list."),
            ("Shifts & Slots", "Sub-screen", "Create work shifts and generate appointment slots per doctor."),
        ],
        "db": [
            ("doctors", "C/R/U", "CRUD doctor profiles; import/export."),
            ("users", "C/U", "Create linked accounts; deactivate with the doctor."),
            ("workShifts", "C/R", "Create shifts for doctors."),
            ("appointmentSlots", "C", "Generate slots from shift templates."),
        ],
        "queries": """
1/ Create doctor with linked account
db.users.insertOne({ email, passwordHash: tempHash, fullName, role: "doctor", isActive: true })
db.doctors.insertOne({ userId: newUserId, specialtyId, departmentId,
  degree, consultationFee, bio, isActive: true, isFeatured: false })

2/ Generate slots from a shift
shiftSlots.forEach(s => db.appointmentSlots.insertOne({
  doctorId, shiftId, roomId, date, startTime: s.start, endTime: s.end, status: "available" }))

3/ Export filtered list
db.doctors.aggregate([
  { $match: { specialtyId: ObjectId(spec), isActive: true } },
  { $lookup: { from: "users", localField: "userId", foreignField: "_id", as: "account" } },
  { $lookup: { from: "specialties", localField: "specialtyId", foreignField: "_id", as: "specialty" } }
])
""",
    },
    {
        "feature": "Manage Patients",
        "uc": "UC-30", "name": "Manage Patients",
        "description": "Allows the Admin to view the patients list with search by name/phone/email, open a patient detail showing profile, appointment history, and wallet summary, create a walk-in patient record manually, update patient profiles, and activate/deactivate patient accounts. Clinical EMR data remains read-only for Admin.",
        "related": "UC-30 \u2013 Manage Patients",
        "image": "UI_ManagePatients.png", "caption": "Manage Patients (Admin)",
        "fields": [
            ("Search", "Text Input", "By name, phone, or email."),
            ("Patients table", "Table", "Name, phone, email, DOB, status, registered date, actions."),
            ("Detail view", "Tabs", "Profile / Appointments / Wallet summary (read-only clinical data)."),
            ("Create Patient", "Modal Form", "Walk-in patient: basic info, no email verification required."),
            ("Activate / Deactivate", "Toggle + Confirm", "Warns and lists future appointments before deactivation."),
        ],
        "db": [
            ("patients", "C/R/U", "List/search, create walk-in, update profiles."),
            ("users", "C/U", "Linked account creation and status toggling."),
            ("appointments", "R", "History tab and future-appointment warning."),
            ("wallets", "R", "Wallet summary in the detail view."),
        ],
        "queries": """
1/ Search patients
db.users.aggregate([
  { $match: { role: "patient", $or: [
      { fullName: { $regex: kw, $options: "i" } },
      { phone: { $regex: kw } }, { email: { $regex: kw, $options: "i" } } ] } },
  { $lookup: { from: "patients", localField: "_id", foreignField: "userId", as: "profile" } }
])

2/ Future appointments warning before deactivation
db.appointments.find({ patientId: ObjectId(pid), status: "confirmed",
  startTime: { $gt: new Date() } })
""",
    },
    {
        "feature": "Admin Dashboard",
        "uc": "UC-31", "name": "View Admin Dashboard",
        "description": "Allows the Admin to view a dashboard with system KPIs: revenue summary with date-range and doctor filters plus Excel export, appointment volume chart, new patient count, active doctors, and open complaints, with shortcuts to all management modules.",
        "related": "UC-31 \u2013 View Admin Dashboard",
        "image": "UI_AdminDashboard.png", "caption": "Admin Dashboard",
        "fields": [
            ("KPI cards", "Stat Cards", "Total revenue, appointments, new patients, open complaints."),
            ("Revenue chart", "Bar/Line Chart", "Revenue by day; filter by date range and doctor."),
            ("Filters", "Date Range + Dropdown", "Period and doctor selection."),
            ("Export Excel", "Button", "Downloads the revenue report for the current filter."),
            ("Module shortcuts", "Icon Grid", "Links to Accounts, Doctors, Patients, Master Data, Complaints."),
        ],
        "db": [
            ("transactions", "R", "Aggregate revenue by day/doctor for the chart and export."),
            ("appointments", "R", "Appointment volume KPIs."),
            ("users", "R", "New patients count in period."),
            ("complaints", "R", "Open complaints count."),
        ],
        "queries": """
1/ Revenue by day in range
db.transactions.aggregate([
  { $match: { type: "payment", status: "success",
      createdAt: { $gte: ISODate(from), $lte: ISODate(to) } } },
  { $group: { _id: { $dateToString: { format: "%Y-%m-%d", date: "$createdAt" } },
      revenue: { $sum: "$amount" } } },
  { $sort: { _id: 1 } }
])

2/ Revenue by doctor
db.transactions.aggregate([
  { $match: { type: "payment", status: "success", refType: "appointment" } },
  { $lookup: { from: "appointments", localField: "refId", foreignField: "_id", as: "appt" } },
  { $unwind: "$appt" },
  { $group: { _id: "$appt.doctorId", revenue: { $sum: "$amount" } } },
  { $sort: { revenue: -1 } }
])
""",
    },
    {
        "feature": "Manage Pharmacy",
        "uc": "UC-32", "name": "Manage Pharmacy",
        "description": "Allows the Staff to manage the medicines inventory: register new medicines (name, unit, price, reorder level), view the inventory list with stock levels and alert badges, update medicine information, record stock inbound with batch/expiry and outbound with reason, and monitor low-stock alerts.",
        "related": "UC-32 \u2013 Manage Pharmacy",
        "image": "UI_Pharmacy.png", "caption": "Pharmacy Inventory (Staff)",
        "fields": [
            ("Inventory table", "Table", "Medicine, unit, price, stock, reorder level, low-stock badge, actions."),
            ("Create Medicine", "Modal Form", "Name, unit, price, reorder level."),
            ("Stock Inbound", "Modal Form", "Quantity, batch number, expiry date, supplier."),
            ("Stock Outbound", "Modal Form", "Quantity (\u2264 stock), reason (dispensed/expired/damaged)."),
            ("Low Stock Alerts", "Banner/List", "Medicines below reorder level highlighted."),
        ],
        "db": [
            ("medicines", "C/R/U", "Catalog CRUD and stock level updates."),
            ("stockMovements", "C/R", "Immutable inbound/outbound records."),
        ],
        "queries": """
1/ Record inbound and update stock
db.stockMovements.insertOne({ medicineId, type: "in", quantity,
  batchNo, expiryDate, by: currentStaffId, at: new Date() })
db.medicines.updateOne({ _id: medicineId }, { $inc: { stockQuantity: quantity } })

2/ Outbound with stock guard (atomic)
db.medicines.updateOne(
  { _id: medicineId, stockQuantity: { $gte: quantity } },
  { $inc: { stockQuantity: -quantity } })
db.stockMovements.insertOne({ medicineId, type: "out", quantity, reason,
  by: currentStaffId, at: new Date() })

3/ Low stock alerts
db.medicines.find({ $expr: { $lt: ["$stockQuantity", "$reorderLevel"] }, isActive: true })
""",
    },
    {
        "feature": "Queue Check-in",
        "uc": "UC-33", "name": "Manage Queue Check-in",
        "description": "Allows the Staff at reception to find a patient's confirmed appointment of the day by name, phone, or booking code, verify identity, check the patient in, and issue a queue ticket that joins the active queue session of the assigned clinic room.",
        "related": "UC-33 \u2013 Manage Queue Check-in",
        "image": "UI_QueueCheckin.png", "caption": "Queue Check-in (Staff)",
        "fields": [
            ("Search", "Text Input", "By patient name, phone, or booking code."),
            ("Appointment card", "Card", "Patient, doctor, room, time, current status."),
            ("Check In & Issue Ticket", "Button", "Marks checked-in and issues the next ticket number."),
            ("Ticket preview", "Popup/Print", "Ticket number, room code, doctor; printable."),
            ("Today's check-ins", "Table", "Recent check-ins with ticket numbers and statuses."),
        ],
        "db": [
            ("appointments", "R/U", "Find today's confirmed appointment; mark checked-in."),
            ("queueSessions", "R", "Verify the room's session is active."),
            ("queueTickets", "C", "Issue the next ticket number in the session."),
        ],
        "queries": """
1/ Find today's appointment
db.appointments.findOne({ status: "confirmed",
  startTime: { $gte: startOfToday, $lt: endOfToday },
  $or: [ { bookingCode: code }, { patientPhone: phone } ] })

2/ Issue next ticket (atomic counter on session)
session = db.queueSessions.findOneAndUpdate(
  { roomId: appt.roomId, status: "active" },
  { $inc: { lastNumber: 1 } }, { returnDocument: "after" })
db.queueTickets.insertOne({ sessionId: session._id, appointmentId: appt._id,
  patientId: appt.patientId, number: session.lastNumber,
  status: "waiting", issuedAt: new Date() })
db.appointments.updateOne({ _id: appt._id }, { $set: { status: "checked-in" } })
""",
    },
    {
        "feature": "Manage Complaints (Staff)",
        "uc": "UC-34", "name": "Manage Complaints",
        "description": "Allows the Staff to view all submitted complaints with status/category/date filters, open a complaint detail with the full message history, reply to the patient (notification sent), and update the complaint status following Open \u2192 In Progress \u2192 Resolved with a resolution note.",
        "related": "UC-34 \u2013 Manage Complaints",
        "image": "UI_StaffComplaints.png", "caption": "Manage Complaints (Staff)",
        "fields": [
            ("Filters", "Dropdowns + Date", "Status, category, and date range."),
            ("Complaints table", "Table", "Subject, patient, category, created, status badge, assignee."),
            ("Detail thread", "Timeline", "Patient message and staff replies chronologically."),
            ("Reply box", "Text Area + Button", "Sends the reply and notifies the patient."),
            ("Status control", "Dropdown", "Open \u2192 In Progress \u2192 Resolved (+ resolution note)."),
        ],
        "db": [
            ("complaints", "R/U", "List/filter; update status and resolution note."),
            ("complaintReplies", "C/R", "Read thread; create staff replies."),
            ("notifications", "C", "Notify the patient about replies/resolution."),
        ],
        "queries": """
1/ Filter complaints
db.complaints.find({ status: "open", category,
  createdAt: { $gte: ISODate(from), $lte: ISODate(to) } }).sort({ createdAt: -1 })

2/ Reply and notify
db.complaintReplies.insertOne({ complaintId, staffId: currentStaffId,
  message, createdAt: new Date() })
db.notifications.insertOne({ userId: complaint.patientUserId, type: "complaint-reply",
  title: "New reply to your complaint", refId: complaintId, isRead: false, createdAt: new Date() })

3/ Resolve
db.complaints.updateOne({ _id: complaintId, status: "in-progress" },
  { $set: { status: "resolved", resolutionNote, resolvedAt: new Date() } })
""",
    },
    {
        "feature": "Staff Dashboard",
        "uc": "UC-35", "name": "View Staff Dashboard",
        "description": "Allows the Staff to view an operational dashboard after login: today's check-in count, waiting tickets, open complaints, and low-stock alerts, with shortcuts to Queue Check-in, Pharmacy, and Complaints modules. Widgets auto-refresh every 60 seconds.",
        "related": "UC-35 \u2013 View Staff Dashboard",
        "image": "UI_StaffDashboard.png", "caption": "Staff Dashboard",
        "fields": [
            ("KPI widgets", "Stat Cards", "Today's check-ins, waiting tickets, open complaints, low-stock count."),
            ("Quick actions", "Buttons", "Check-in Patient, Stock Inbound, View Complaints."),
            ("Recent check-ins", "Table", "Latest issued tickets with room and status."),
            ("Low stock list", "List", "Top medicines below reorder level with quick links."),
        ],
        "db": [
            ("queueTickets", "R", "Today's check-in and waiting counters."),
            ("complaints", "R", "Open complaints counter."),
            ("medicines", "R", "Low-stock list."),
        ],
        "queries": """
1/ Today's counters
db.queueTickets.countDocuments({ issuedAt: { $gte: startOfToday } })
db.queueTickets.countDocuments({ status: "waiting" })
db.complaints.countDocuments({ status: "open" })

2/ Low stock top 5
db.medicines.find({ $expr: { $lt: ["$stockQuantity", "$reorderLevel"] } })
  .sort({ stockQuantity: 1 }).limit(5)
""",
    },
]


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    doc.add_heading("III. Design Specifications", level=1)

    for i, s in enumerate(SECTIONS, start=1):
        doc.add_heading(f"{i}. {s['feature']}", level=2)
        doc.add_heading(f"{i}.1 {s['uc']}_{s['name']}", level=3)
        add_bullet(doc, "Description", s["description"])
        add_bullet(doc, "Related Use Case", s["related"])
        add_ui_design(doc, s["image"], s["caption"])
        doc.add_paragraph()
        add_table(doc,
                  ["Field Name", "Field Type", "Description & Data Initialization"],
                  s["fields"], widths=[1.6, 1.3, 3.6])
        doc.add_paragraph()
        add_sub_heading(doc, "Database Access")
        add_table(doc, ["Collection", "CRUD", "Description"],
                  s["db"], widths=[1.6, 0.9, 4.0])
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("MongoDB Commands")
        run.bold = True
        run.italic = True
        add_code(doc, s["queries"])
        doc.add_paragraph()

    try:
        doc.save(OUT_PATH)
        print(f"Saved: {OUT_PATH} | sections: {len(SECTIONS)}")
    except PermissionError:
        alt = OUT_PATH.with_name(OUT_PATH.stem + "_new.docx")
        doc.save(alt)
        print(f"Original file is locked (open in Word). Saved to: {alt}")


if __name__ == "__main__":
    main()
