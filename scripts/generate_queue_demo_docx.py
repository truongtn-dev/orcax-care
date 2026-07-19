#!/usr/bin/env python3
"""Generate detailed Word debate guide for Queue Management (with line citations)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SHOTS = ROOT / "docs" / "demo" / "queue-screenshots"
OUT = ROOT / "docs" / "demo" / "OrcaXCare_Queue_Management_Demo.docx"


def read_lines(rel_path: str, start: int, end: int) -> str:
    """1-based inclusive line range from a source file."""
    path = ROOT / rel_path
    lines = path.read_text(encoding="utf-8").splitlines()
    chunk = lines[start - 1 : end]
    numbered = [f"{start + i:>4}| {line}" for i, line in enumerate(chunk)]
    return "\n".join(numbered)


def add_heading_styled(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_code(doc, text: str, title: str | None = None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main():
    doc = Document()
    title = doc.add_heading("OrcaX Care — Queue Management: Tài liệu Debate kỹ thuật", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, f"Ngày cập nhật: {date.today().strftime('%d/%m/%Y')}")
    add_para(
        doc,
        "Mục tiêu: trình bày chi tiết thuật toán, chỉ rõ file + số dòng, trích code thật từ repo, "
        "để debate với giảng viên về thiết kế và cài đặt Queue Management (7 UC).",
    )
    add_para(
        doc,
        "Quy ước trích dẫn: File:path → dòng start–end. Đoạn code có tiền tố số dòng dạng  123| code.",
        bold=True,
    )

    # ─── 1. OVERVIEW ───────────────────────────────────────────────
    add_heading_styled(doc, "1. Tổng quan module & mô hình dữ liệu", 1)
    add_para(
        doc,
        "Queue Management quản lý hàng đợi khám theo phòng (ClinicRoom). Doctor mở ca (QueueSession), "
        "Staff check-in appointment hôm nay và cấp số (QueueTicket), Patient/TV board nhận cập nhật realtime.",
    )

    add_heading_styled(doc, "1.1. Các file lõi (map trách nhiệm)", 2)
    add_bullets(
        doc,
        [
            "server/src/services/queueSession.service.js — open/call/skip/recall/close, board, patient status (~840 dòng)",
            "server/src/services/queueCheckin.service.js — tìm appointment hôm nay, issue ticket / issue-all",
            "server/src/realtime/socket.js — Socket.IO rooms + emitQueueEvent",
            "server/src/models/QueueSession.js, QueueTicket.js, QueueAuditLog.js — schema + index",
            "server/src/utils/queueDate.js — startOfToday / endOfToday (ngày lịch, local)",
            "client: DoctorQueueSessionPage, StaffQueueCheckinPage, PatientQueueStatusPage, QueueBoardPage",
            "client/src/utils/queueBoardState.js — derive state UI board",
        ],
    )

    add_heading_styled(doc, "1.2. State machine ticket (thuật toán trạng thái)", 2)
    add_para(doc, "Enum status trong QueueTicket (model dòng 24–28):")
    add_code(doc, read_lines("server/src/models/QueueTicket.js", 23, 32), "File: server/src/models/QueueTicket.js — L23–32")

    add_para(doc, "Chuyển trạng thái hợp lệ:", bold=True)
    add_bullets(
        doc,
        [
            "issue → waiting (mặc định khi tạo ticket)",
            "call next / recall → called (đang gọi)",
            "skip → skipped (vắng mặt tạm thời, vẫn có thể recall)",
            "close session: waiting → no-show",
            "done / serving: phục vụ xong (reserved trong enum)",
        ],
    )
    add_para(
        doc,
        "Index unique (sessionId, number) L37 đảm bảo trong 1 session không trùng số. "
        "Index (sessionId, status, number) L38 tối ưu query FIFO waiting sort number ASC.",
    )
    add_code(doc, read_lines("server/src/models/QueueTicket.js", 37, 38), "File: server/src/models/QueueTicket.js — L37–38")

    add_heading_styled(doc, "1.3. Thuật toán ngày hôm nay", 2)
    add_para(
        doc,
        "Mọi nghiệp vụ 'hôm nay' dùng half-open interval [startOfToday, endOfToday). "
        "Không dùng UTC date string → tránh lệch múi giờ giữa Atlas và máy local.",
    )
    add_code(doc, read_lines("server/src/utils/queueDate.js", 1, 11), "File: server/src/utils/queueDate.js — L1–11")

    # ─── 2. UC1 ────────────────────────────────────────────────────
    add_heading_styled(doc, "2. UC1 — Open Queue Session", 1)
    add_para(doc, "API: POST /api/queue/sessions/open  |  Hàm: openSession() L257–334", bold=True)

    add_heading_styled(doc, "2.1. Thuật toán (pseudo)", 2)
    add_numbered(
        doc,
        [
            "Validate roomId là ObjectId.",
            "Resolve Doctor từ userId (isActive=true).",
            "Tìm ClinicRoom thuộc đúng department của doctor + active.",
            "Constraint A: doctor chưa có session open/paused hôm nay (ACTIVE_SESSION_STATUSES).",
            "Constraint B: room chưa có session open/paused hôm nay.",
            "Create QueueSession: status=open, currentNumber=0, lastNumber=0.",
            "Ghi QueueAuditLog action=open_session.",
            "broadcastSessionUpdate → Socket.IO push queue:update.",
        ],
    )

    add_heading_styled(doc, "2.2. Vì sao 2 constraint?", 2)
    add_para(
        doc,
        "Constraint A tránh 1 bác sĩ mở 2 ca song song → tranh số / gọi nhầm phòng. "
        "Constraint B tránh 2 bác sĩ cùng 1 phòng cấp số chồng chéo. Đây là invariant nghiệp vụ, không chỉ UI.",
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 278, 314),
        "File: server/src/services/queueSession.service.js — openSession L278–314",
    )

    add_heading_styled(doc, "2.3. Điểm debate", 2)
    add_bullets(
        doc,
        [
            "Department gắn với Doctor profile — không cho chọn khoa tự do → đúng phòng khám thực tế.",
            "409 Conflict khi trùng session — client hiện form/error rõ ràng.",
            "Audit log ngay khi open — chứng minh truy vết được.",
        ],
    )

    # ─── 3. UC2 ────────────────────────────────────────────────────
    add_heading_styled(doc, "3. UC2 — Issue Queue Ticket (Staff check-in)", 1)
    add_para(
        doc,
        "API: GET /api/staff/checkin/today | POST /api/staff/checkin/:id/issue-ticket | POST .../issue-all",
        bold=True,
    )

    add_heading_styled(doc, "3.1. Thuật toán tìm appointment hôm nay", 2)
    add_numbered(
        doc,
        [
            "Lấy mọi AppointmentSlot có date ∈ [todayStart, todayEnd).",
            "baseFilter: appointment.status='confirmed' AND slotId ∈ slotIds.",
            "Không keyword → trả list (limit 50).",
            "Có keyword → tìm User patient theo fullName/email/phone regex, OR theo reason; nếu query APT-xxx → filter postfix Id.",
        ],
    )
    add_code(
        doc,
        read_lines("server/src/services/queueCheckin.service.js", 57, 91),
        "File: server/src/services/queueCheckin.service.js — findTodayConfirmedAppointments L57–91",
    )

    add_heading_styled(doc, "3.2. Thuật toán cấp số (atomic increment)", 2)
    add_para(
        doc,
        "Điểm quan trọng: số thứ tự không lấy max(ticket)+1 ngoài DB (race condition). "
        "Dùng findOneAndUpdate + $inc lastNumber trên QueueSession — atomic trên MongoDB.",
        bold=True,
    )
    add_numbered(
        doc,
        [
            "Validate appointment confirmed + slot hôm nay + có roomId.",
            "Chặn trùng ticket active cùng appointment trong ngày.",
            "findOneAndUpdate session {roomId, date=today, status=open} với $inc: {lastNumber:1}.",
            "Nếu không tìm thấy session open → 409 (doctor chưa open).",
            "Create QueueTicket number = session.lastNumber, status=waiting.",
            "Appointment.status = checked-in.",
            "Audit issue_ticket + broadcast + emit patient-update.",
        ],
    )
    add_code(
        doc,
        read_lines("server/src/services/queueCheckin.service.js", 247, 282),
        "File: server/src/services/queueCheckin.service.js — issueQueueTicket L247–282",
    )

    add_heading_styled(doc, "3.3. Check-in all (demo / batch)", 2)
    add_para(
        doc,
        "issueAllQueueTickets L180–211: for-loop tuần tự gọi issueQueueTicket. "
        "Stop early nếu lỗi (vd. session không open) — fail-fast, tránh half-success khó giải thích.",
    )
    add_code(
        doc,
        read_lines("server/src/services/queueCheckin.service.js", 180, 210),
        "File: server/src/services/queueCheckin.service.js — issueAllQueueTickets L180–210",
    )

    add_heading_styled(doc, "3.4. Điểm debate", 2)
    add_bullets(
        doc,
        [
            "Tại sao không cho patient tự lấy số? → Reception xác thực đúng người, đúng lịch confirmed.",
            "Tại sao $inc lastNumber? → Tránh 2 staff check-in đồng thời ra trùng số (race).",
            "Check-in all là tiện demo; production vẫn nên check-in từng người tại quầy.",
        ],
    )

    # ─── 4. UC3 ────────────────────────────────────────────────────
    add_heading_styled(doc, "4. UC3 — Patient View My Queue Status", 1)
    add_para(doc, "API: GET /api/queue/my-status  |  Hàm: getPatientQueueStatus() L699–742", bold=True)

    add_heading_styled(doc, "4.1. Thuật toán peopleAhead", 2)
    add_para(
        doc,
        "peopleAhead chỉ tính khi ticket đang waiting. Công thức: "
        "COUNT(tickets WHERE sessionId=S AND status='waiting' AND number < myNumber).",
        bold=True,
    )
    add_para(
        doc,
        "Ví dụ: đang waiting #4, đã có waiting #1,#2,#3 → peopleAhead=3. "
        "Nếu bị skip / called → peopleAhead=0 (không còn 'đứng hàng').",
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 699, 740),
        "File: server/src/services/queueSession.service.js — getPatientQueueStatus L699–740",
    )

    add_heading_styled(doc, "4.2. Realtime + polling trên client", 2)
    add_para(
        doc,
        "PatientQueueStatusPage: poll mỗi 3000ms + socket event queue:patient-update. "
        "isSkipped / isCalled phân biệt UI: đang gọi vs tạm skip (không báo 'proceed now' khi đã skip).",
    )

    add_heading_styled(doc, "4.3. Điểm debate", 2)
    add_bullets(
        doc,
        [
            "peopleAhead theo số thứ tự waiting — O(1) nhờ index sessionId+status+number.",
            "Ticket skipped vẫn query được (trong $in) → patient không 'mất vé' sau khi bị skip.",
            "Fallback poll ≤3s đúng requirement cập nhật gần realtime khi mất socket.",
        ],
    )

    # ─── 5. UC4 ────────────────────────────────────────────────────
    add_heading_styled(doc, "5. UC4 — Queue Board (Waiting Room Display)", 1)
    add_para(doc, "API: GET /api/queue/board/:roomId (public)  |  Hàm: getQueueBoard() L744–835+", bold=True)

    add_heading_styled(doc, "5.1. Thuật toán ghép 3 tập ticket (parallel)", 2)
    add_numbered(
        doc,
        [
            "Load waiting: status=waiting, sort number ASC, limit 5 → Up next.",
            "Load called: status ∈ {called,serving}, sort calledAt DESC → Now serving.",
            "Load skipped: status=skipped, sort skippedAt DESC, limit 5 → cột Skipped trên board.",
            "attachPatientInfo (tên + năm sinh) song song bằng Promise.all.",
            "displayNumber = called.number (không dùng session.currentNumber cũ nếu không còn người called).",
            "state: closed > paused > empty (không current && không waiting && không skipped) > active.",
        ],
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 775, 809),
        "File: server/src/services/queueSession.service.js — getQueueBoard core L775–809",
    )

    add_heading_styled(doc, "5.2. deriveQueueBoardState (client đồng bộ REST/socket)", 2)
    add_code(
        doc,
        read_lines("client/src/utils/queueBoardState.js", 1, 10),
        "File: client/src/utils/queueBoardState.js — L1–10",
    )
    add_para(
        doc,
        "Cùng rule state ở server và client → khi socket payload thiếu field 'state', "
        "client vẫn tính đúng từ status/current/waiting/skipped.",
    )

    add_heading_styled(doc, "5.3. Điểm debate (privacy)", 2)
    add_bullets(
        doc,
        [
            "Board public: không login — phù hợp TV phòng chờ.",
            "Hiện số + tên + năm sinh (không SĐT/CMND) — phân biệt trùng tên, hạn chế PII.",
            "Skipped vẫn hiện trên board — bệnh nhân biết mình bị tạm bỏ qua, chờ recall.",
        ],
    )

    # ─── 6. UC5 ────────────────────────────────────────────────────
    add_heading_styled(doc, "6. UC5 — Call Next Patient (FIFO)", 1)
    add_para(doc, "API: POST /api/queue/sessions/:id/call-next  |  Hàm: callNextTicket() L365–439", bold=True)

    add_heading_styled(doc, "6.1. Thuật toán FIFO + mutex logic", 2)
    add_numbered(
        doc,
        [
            "Verify doctor owns session; session.status phải = open.",
            "Guard: nếu còn ticket called/serving → 409 (bắt buộc Skip trước). "
            "Tránh gọi #5 trong khi #4 vẫn 'called' → bug patient vẫn thấy proceed now.",
            "findOneAndUpdate: lấy ticket waiting có number nhỏ nhất (sort number:1), set status=called.",
            "session.currentNumber = ticket.number.",
            "Audit call_next + broadcast + patient-update.",
        ],
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 379, 406),
        "File: server/src/services/queueSession.service.js — callNextTicket L379–406",
    )
    add_para(
        doc,
        "Thuật toán xếp hàng: FIFO theo ticket.number tăng dần = đúng thứ tự cấp số tại quầy. "
        "Mongo sort + update atomic giảm race giữa 2 tab doctor.",
        bold=True,
    )

    add_heading_styled(doc, "6.2. Enrich tên + năm sinh", 2)
    add_para(
        doc,
        "attachPatientInfo L46–62: batch User.fullName + Patient.dateOfBirth → birthYear. "
        "Dùng Map theo userId tránh N+1 query.",
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 21, 61),
        "File: server/src/services/queueSession.service.js — patientInfoByUserIds + attachPatientInfo L21–61",
    )

    # ─── 7. UC6 ────────────────────────────────────────────────────
    add_heading_styled(doc, "7. UC6 — Skip & Recall", 1)

    add_heading_styled(doc, "7.1. Skip — markTicketSkipped L519–581", 2)
    add_numbered(
        doc,
        [
            "Chỉ skip ticket đang called/serving (không skip waiting).",
            "status → skipped, skippedAt = now.",
            "session.lastSkippedTicketId = ticket._id (pointer LIFO gần nhất).",
            "Audit mark_skipped + broadcast (board hiện cột Skipped) + patient-update (isSkipped=true).",
        ],
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 533, 572),
        "File: server/src/services/queueSession.service.js — markTicketSkipped L533–572",
    )

    add_heading_styled(doc, "7.2. Recall — LIFO trên skipped (recallLastSkipped L441–517)", 2)
    add_para(
        doc,
        "Thuật toán recall = LIFO (Last-In-First-Out) trên tập ticket skipped:",
        bold=True,
    )
    add_numbered(
        doc,
        [
            "Ưu tiên ticket = session.lastSkippedTicketId nếu còn status=skipped.",
            "Fallback: findOne status=skipped sort skippedAt DESC, number DESC.",
            "Đưa ticket → called, clear skippedAt, set currentNumber, clear lastSkippedTicketId.",
            "Audit action=recall + broadcast + patient-update.",
        ],
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 459, 484),
        "File: server/src/services/queueSession.service.js — recallLastSkipped L459–484",
    )
    add_para(
        doc,
        "Ví dụ: skip #2 rồi #3 → Recall gọi lại #3 (mới skip gần nhất). "
        "Người #2 vẫn nằm trong cột Skipped trên board / doctor UI cho đến khi recall tiếp hoặc close.",
    )

    add_heading_styled(doc, "7.3. UI ConfirmDialog (không window.confirm)", 2)
    add_para(
        doc,
        "DoctorQueueSessionPage dùng ConfirmDialog.jsx cho Skip và Close — popup trong app, "
        "có title/description/variant danger. Nút Recall disabled khi skippedCount=0; "
        "nút Call next disabled khi còn calledTicket (mustSkipBeforeNext).",
    )

    add_heading_styled(doc, "7.4. Điểm debate", 2)
    add_bullets(
        doc,
        [
            "Skip ≠ xóa khỏi hệ thống: ticket còn trên board (transparent với phòng chờ).",
            "Recall LIFO hợp lý: người vừa miss thường đứng gần quầy.",
            "Guard Call next khi đang có called — sửa bug trạng thái lệch giữa patient/board.",
        ],
    )

    # ─── 8. UC7 ────────────────────────────────────────────────────
    add_heading_styled(doc, "8. UC7 — Close Queue Session", 1)
    add_para(doc, "API: POST /api/queue/sessions/:id/close  |  Hàm: closeSession() L657–697", bold=True)
    add_numbered(
        doc,
        [
            "status → closed, closedAt = now.",
            "Bulk update: waiting tickets → no-show.",
            "Audit close_session + broadcast (board state=closed).",
            "issueQueueTicket yêu cầu session open → sau close không cấp số mới.",
            "Doctor UI clear session → hiện lại form Open session.",
        ],
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 671, 696),
        "File: server/src/services/queueSession.service.js — closeSession L671–696",
    )

    # ─── 9. REALTIME ───────────────────────────────────────────────
    add_heading_styled(doc, "9. Kiến trúc Realtime (Socket.IO)", 1)

    add_heading_styled(doc, "9.1. Channel naming", 2)
    add_code(doc, read_lines("server/src/realtime/socket.js", 13, 59), "File: server/src/realtime/socket.js — L13–59")
    add_bullets(
        doc,
        [
            "queue:room:{roomId} — Board TV + mọi client theo phòng",
            "queue:session:{sessionId} — Doctor console",
            "queue:patient:{patientUserId} — Patient status riêng (không lộ cho người khác)",
        ],
    )

    add_heading_styled(doc, "9.2. broadcastSessionUpdate", 2)
    add_para(
        doc,
        "Mỗi thao tác nghiệp vụ gọi loadSessionTickets (waiting + called + skipped) → serialize → emit queue:update. "
        "Đảm bảo Doctor UI, Board, Staff (nếu listen) cùng snapshot.",
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 188, 204),
        "File: server/src/services/queueSession.service.js — broadcastSessionUpdate L188–204",
    )

    add_heading_styled(doc, "9.3. Fallback polling", 2)
    add_bullets(
        doc,
        [
            "Patient: POLL_MS = 3000 → ≤ 3 giây kể cả mất socket",
            "Board: POLL_MS = 5000 → ≤ 5 giây",
            "Debate: hybrid realtime — Socket ưu tiên, poll đảm bảo availability",
        ],
    )

    # ─── 10. AUDIT ─────────────────────────────────────────────────
    add_heading_styled(doc, "10. Audit trail", 1)
    add_para(
        doc,
        "Mọi action quan trọng ghi QueueAuditLog qua writeAudit() L75–84 hoặc QueueAuditLog.create trực tiếp. "
        "Actions: open_session, issue_ticket, call_next, mark_skipped, recall, pause, resume, close_session.",
    )
    add_code(
        doc,
        read_lines("server/src/services/queueSession.service.js", 75, 84),
        "File: server/src/services/queueSession.service.js — writeAudit L75–84",
    )
    add_para(
        doc,
        "Debate: Auditor nhìn collection QueueAuditLog biết ai gọi số nào lúc nào — "
        "không chỉ dựa UI (UI có thể refresh/mất state).",
    )

    # ─── 11. FLOW END-TO-END ───────────────────────────────────────
    add_heading_styled(doc, "11. Luồng end-to-end (hình dung thuật toán tổng)", 1)
    add_code(
        doc,
        "Doctor openSession()\n"
        "  └─ QueueSession(open, lastNumber=0)\n"
        "Staff issueQueueTicket() × N\n"
        "  └─ lastNumber++ atomic → QueueTicket(#1..#N, waiting)\n"
        "       Appointment = checked-in\n"
        "Doctor callNextTicket()\n"
        "  └─ min(waiting.number) → called; board Now serving\n"
        "Doctor markTicketSkipped()  [nếu vắng]\n"
        "  └─ called → skipped; board cột Skipped; patient isSkipped\n"
        "Doctor callNextTicket()     [phải skip trước mới call được]\n"
        "Doctor recallLastSkipped()  [LIFO skipped → called]\n"
        "Doctor closeSession()\n"
        "  └─ waiting → no-show; session closed; board Clinic session ended",
    )

    # ─── 12. FAQ DEBATE ────────────────────────────────────────────
    add_heading_styled(doc, "12. Câu hỏi giảng viên & câu trả lời có dẫn chứng code", 1)

    faqs = [
        (
            "Vì sao cấp số bằng $inc lastNumber chứ không max(number)+1?",
            "Tránh race khi 2 staff issue đồng thời. Code: queueCheckin.service.js L257–264 findOneAndUpdate + $inc.",
        ),
        (
            "Call next theo thuật toán gì?",
            "FIFO: sort waiting theo number ASC, lấy phần tử nhỏ nhất. L395–399 callNextTicket.",
        ),
        (
            "Recall theo thuật toán gì?",
            "LIFO trên skipped: lastSkippedTicketId trước, fallback sort skippedAt DESC. L459–471.",
        ),
        (
            "peopleAhead tính thế nào?",
            "COUNT waiting có number < mine. L723–730. Chỉ khi status=waiting.",
        ),
        (
            "Skipped có còn hiện board không?",
            "Có — getQueueBoard load skippedTickets L784–788, UI QueueBoardPage cột Skipped.",
        ),
        (
            "Realtime dùng gì? Mất mạng thì sao?",
            "Socket.IO channels trong socket.js L14–59. Fallback poll patient 3s / board 5s.",
        ),
        (
            "Board public có lộ thông tin nhạy cảm?",
            "Chỉ số + tên + năm sinh; không phone/email. Public GET board không auth.",
        ),
        (
            "Làm sao chứng minh thao tác đã xảy ra?",
            "QueueAuditLog via writeAudit L75–84; mỗi open/call/skip/recall/close/issue đều ghi.",
        ),
        (
            "Tại sao Call next bị chặn khi đang có người called?",
            "Guard L383–393 — tránh lệch trạng thái patient #4 still 'proceed' trong khi board đã #5.",
        ),
        (
            "Test tự động?",
            "server/src/tests/queueManagement.test.js — flow open → issue → call → skip → recall → close.",
        ),
    ]
    for i, (q, a) in enumerate(faqs, 1):
        add_para(doc, f"Q{i}. {q}", bold=True)
        add_para(doc, f"A: {a}")

    # ─── 13. DEMO PREP ─────────────────────────────────────────────
    add_heading_styled(doc, "13. Chuẩn bị demo nhanh", 1)
    add_code(
        doc,
        "npm run dev:server\n"
        "npm run dev:client\n"
        "npm run demo:queue   # tạo 5 appointment confirmed hôm nay + DOB\n"
        "# Accounts:\n"
        "# Doctor  doctor.an@orcaxcare.com / Doctor@123  → /doctor/queue\n"
        "# Staff   staff@orcaxcare.com / Staff@123       → /staff/checkin\n"
        "# Patient patient@orcaxcare.com / Patient@123   → /patient/queue\n"
        "# Board   /queue-board/:roomId",
    )
    add_para(
        doc,
        "Thứ tự demo: Open session → Staff Check in all (hoặc từng người) → "
        "mở Board + Patient → Call next → Skip → Call next → Recall → Close.",
    )

    # optional screenshots appendix
    add_heading_styled(doc, "14. Phụ lục — Screenshot demo (nếu có)", 1)
    shot_files = sorted(SHOTS.glob("*.png")) if SHOTS.exists() else []
    if shot_files:
        add_para(doc, f"Có {len(shot_files)} ảnh trong docs/demo/queue-screenshots/. Một số ảnh minh họa:")
        for img in shot_files[:6]:
            doc.add_heading(img.name, level=3)
            try:
                doc.add_picture(str(img), width=Inches(5.8))
            except Exception:
                add_para(doc, f"[Không nhúng được: {img.name}]")
    else:
        add_para(doc, "Chưa có screenshot trong docs/demo/queue-screenshots/.")

    add_heading_styled(doc, "15. Checklist trả lời miệng (1 phút)", 1)
    add_numbered(
        doc,
        [
            "FIFO call next theo number tăng dần (L395–399).",
            "Atomic $inc lastNumber khi issue ticket (L257–264 checkin).",
            "Skip giữ ticket trên board; recall LIFO (L459–471).",
            "peopleAhead = COUNT waiting number nhỏ hơn (L723–730).",
            "Socket.IO 3 channel + poll fallback 3s/5s (socket.js + client).",
            "Mỗi action ghi QueueAuditLog.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Created {OUT}")
    print(f"  Size: {OUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
