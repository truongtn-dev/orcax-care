"""One-off: sprint-plan.md -> sprint-plan.docx. Run: python build_sprint_word.py"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "sprint-plan.md"
OUT_PATH = ROOT / "sprint-plan.docx"


def add_inline_bold(p, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.count("`") >= 2:
            m = re.match(r"`([^`]+)`", part)
            if m:
                run = p.add_run(m.group(1))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            else:
                p.add_run(part)
        else:
            p.add_run(part.replace("`", ""))


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j].strip() if j < len(row) else ""
            cell_text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            cell_text = cell_text.replace("`", "")
            table.rows[i].cells[j].text = cell_text


def parse_table_lines(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and is_table_row(lines[i]):
        raw = lines[i].strip()
        if re.match(r"^\|\s*[-:]+\s*\|", raw):
            i += 1
            continue
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def main() -> None:
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("##"):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            add_inline_bold(p, stripped[2:])
            i += 1
            continue

        if is_table_row(stripped):
            rows, ni = parse_table_lines(lines, i)
            add_table(doc, rows)
            i = ni
            doc.add_paragraph()
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_bold(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_bold(p, stripped)
        i += 1

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
